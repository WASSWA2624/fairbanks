#!/usr/bin/env python3
"""Build five clear, eye-catching, information-rich one-page CVs.

Photos match the original waist-up portrait (full figure in frame).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import (
    Flowable,
    HRFlowable,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
)

ROOT = Path(__file__).resolve().parent
PHOTO_SRC = ROOT.parent / "WhatsApp Image 2026-07-27 at 14.23.53.jpeg"
PHOTO_DIR = ROOT / "_assets"
PHOTO_CIRCLE = PHOTO_DIR / "photo_circle.png"
PHOTO_SQUARE = PHOTO_DIR / "photo_square.png"
PHOTO_PORTRAIT = PHOTO_DIR / "photo_portrait.png"
PREVIEW = ROOT / "_preview"

NAME = "RACHEAL NABUKEERA SEKAGIRI"
TITLE = "Senior Human Resource Executive  ·  Human Capital Strategist  ·  Organisational Development"
LOCATION = "Kampala, Uganda"
PHONE1 = "+256 772 849258"
PHONE2 = "+256 701 849258"
MOBILE = f"{PHONE1}  ·  {PHONE2}"
EMAIL = "info@fairbanksmedicalcentre.org"
WEB = "www.fairbanksmedicalcentre.org"

PROFILE = (
    "Strategic HR executive with nearly <b>30 years</b> of progressive experience and over "
    "<b>18 years</b> in executive Human Resource leadership. Aligns people, performance and "
    "organisational goals across complex institutions."
)
PROFILE2 = (
    "Founder & Managing Director, FairBanks Medical Centre Ltd. Formerly directed Group HR for "
    "<b>2,000+ employees</b> across <b>50+ departments</b>. PhD Candidate (Management) researching "
    "AI/ML prediction of occupational burnout among medical practitioners. Member, Federation of "
    "Uganda Employers (FUE)."
)

HIGHLIGHTS = [
    "30 years progressive HR & executive management",
    "18+ years executive Human Resource leadership",
    "Strategic HR for 2,000+ employees / 50+ units",
    "Founder & MD, FairBanks Medical Centre Ltd",
    "Social enterprise reach: 10,000+ beneficiaries",
    "PhD Candidate — HR analytics, AI & organisational behaviour",
    "Executive advisor on change & organisational effectiveness",
    "Workforce planning, talent & performance leadership",
]

SKILLS = [
    "HR Strategy",
    "Human Capital Management",
    "Organisational Development",
    "Talent Management",
    "Workforce Planning",
    "Succession Planning",
    "Change Management",
    "Performance Management",
    "Employee Relations",
    "Labour Law Compliance",
    "Corporate Governance",
    "Leadership Development",
    "HR Analytics / HRIS",
    "Stakeholder Engagement",
]

ROLES = [
    {
        "title": "Founder, Managing Director & Executive HR Leader",
        "org": "FairBanks Medical Centre Ltd",
        "dates": "2025 - Present",
        "points": [
            "Lead strategy, governance, HR systems and institutional growth.",
            "Founded FairBanks Social Enterprise Initiative (10,000+ beneficiaries).",
            "Build partnerships with government, insurers and development partners.",
        ],
    },
    {
        "title": "Group Human Resource & Administration Manager",
        "org": "Norvik Group",
        "dates": "Jul 2016 - Feb 2026",
        "points": [
            "Directed Group HR strategy for 2,000+ employees across 50+ units.",
            "Led workforce planning, restructuring, succession and leadership development.",
            "Strengthened HR governance, performance systems and staff retention.",
        ],
    },
    {
        "title": "Human Resource & Administration Manager",
        "org": "Norvik Hospital Ltd",
        "dates": "Sep 2013 - 2016",
        "points": [
            "Led recruitment, employee relations, appraisals and staff welfare.",
            "Improved HR policy implementation and operational efficiency.",
        ],
    },
    {
        "title": "Human Resource Manager",
        "org": "St. Catherine's Hospital",
        "dates": "2007 - 2010",
        "points": [
            "Managed HR operations, welfare programmes and labour compliance.",
            "Strengthened recruitment, retention and employee engagement.",
        ],
    },
]

EARLIER = (
    "Earlier roles: Health Management Systems Administrator, St. Catherine's Hospital "
    "(2010-2013); Human Resource Assistant, E Power Limited (1997-2000)."
)

EDUCATION = [
    ("PhD in Management (Ongoing)", "Uganda Christian University",
     "Research: Machine learning prediction of occupational burnout among medical practitioners"),
    ("Master of Social Sector Planning and Management", "Makerere University", None),
    ("Bachelor of Arts in Social Sciences", "Makerere University", None),
]

TECH = (
    "Microsoft Office  ·  HRIS  ·  ERP  ·  HR Analytics  ·  Performance systems  ·  "
    "Payroll systems  ·  Executive reporting  ·  AI applications in HR"
)
MORE = (
    "Member, Federation of Uganda Employers (FUE)  ·  English (Fluent)  ·  "
    "Luganda (Fluent)  ·  References available on request"
)

NAVY = HexColor("#0A3A52")
TEAL = HexColor("#1F6F78")
GOLD = HexColor("#B8953E")
INK = HexColor("#1F2A37")
MUTED = HexColor("#5B6B7C")
LINE = HexColor("#E4E9EF")
SOFT = HexColor("#F4F7FA")
CREAM = HexColor("#F8F5F0")
SIDEBAR = HexColor("#0A3A52")
CHARCOAL = HexColor("#18212B")
FOREST = HexColor("#1E4D52")
SLATE = HexColor("#2F4458")

DOCX_NAVY = RGBColor(0x0A, 0x3A, 0x52)
DOCX_TEAL = RGBColor(0x1F, 0x6F, 0x78)
DOCX_INK = RGBColor(0x1F, 0x2A, 0x37)
DOCX_MUTED = RGBColor(0x5B, 0x6B, 0x7C)
DOCX_GOLD = RGBColor(0xB8, 0x95, 0x3E)

VERSION_SLUGS = [
    "01_classic_airy",
    "02_sidebar_modern",
    "03_one_page",
    "04_timeline",
    "05_editorial_split",
]


class GoldRule(Flowable):
    def __init__(self, width=34 * mm, thickness=1.2):
        super().__init__()
        self._w = width
        self._t = thickness
        self.height = 5

    def wrap(self, availWidth, availHeight):
        return availWidth, self.height

    def draw(self):
        self.canv.setStrokeColor(GOLD)
        self.canv.setLineWidth(self._t)
        self.canv.setLineCap(1)
        self.canv.line(0, 2.5, self._w, 2.5)


def ensure_dirs():
    PHOTO_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW.mkdir(parents=True, exist_ok=True)
    for slug in VERSION_SLUGS:
        (ROOT / slug).mkdir(parents=True, exist_ok=True)


def prepare_photos():
    """Create CV photo assets that preserve the original complete waist-up portrait."""
    img = Image.open(PHOTO_SRC).convert("RGBA")
    w, h = img.size  # ~1066 x 1600

    # Full portrait: nearly the original (tiny trim only) — complete like source
    mx, mt, mb = int(w * 0.015), int(h * 0.008), int(h * 0.008)
    full = img.crop((mx, mt, w - mx, h - mb))
    pw = 560
    ph = int(pw * full.size[1] / full.size[0])
    portrait = full.resize((pw, ph), Image.Resampling.LANCZOS)

    # Rounded portrait with elegant navy/gold frame
    rad = 28
    mask = Image.new("L", (pw, ph), 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, pw - 1, ph - 1), radius=rad, fill=255)
    framed = Image.new("RGBA", (pw, ph), (0, 0, 0, 0))
    framed.paste(portrait, (0, 0))
    framed.putalpha(mask)
    edge = Image.new("RGBA", (pw, ph), (0, 0, 0, 0))
    ed = ImageDraw.Draw(edge)
    ed.rounded_rectangle((2, 2, pw - 3, ph - 3), radius=rad, outline=(10, 58, 82, 255), width=5)
    ed.rounded_rectangle((9, 9, pw - 10, ph - 10), radius=rad - 4, outline=(184, 149, 62, 240), width=2)
    Image.alpha_composite(framed, edge).save(PHOTO_PORTRAIT)

    # Square / circle from TOP of original so arms-crossed pose stays in frame
    # (original is waist-up; top square ≈ head through folded arms)
    side = w
    top = 0
    sq = img.crop((0, top, side, top + side)).resize((700, 700), Image.Resampling.LANCZOS)

    # Circle with inset rings so hair/face/arms are not covered
    size = 700
    inset = int(size * 0.045)
    inner = size - 2 * inset
    circ_photo = sq.resize((inner, inner), Image.Resampling.LANCZOS)
    cmask = Image.new("L", (inner, inner), 0)
    ImageDraw.Draw(cmask).ellipse((0, 0, inner - 1, inner - 1), fill=255)
    circ_photo.putalpha(cmask)
    cout = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    cout.paste(circ_photo, (inset, inset), circ_photo)
    cring = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    cd = ImageDraw.Draw(cring)
    cd.ellipse((3, 3, size - 4, size - 4), outline=(10, 58, 82, 255), width=6)
    cd.ellipse(
        (inset - 1, inset - 1, size - inset, size - inset),
        outline=(184, 149, 62, 255),
        width=2,
    )
    Image.alpha_composite(cout, cring).save(PHOTO_CIRCLE)

    # Rounded square from same complete crop
    size2 = 680
    inset2 = 8
    inner2 = size2 - 2 * inset2
    rad2 = 40
    s = sq.resize((inner2, inner2), Image.Resampling.LANCZOS)
    smask = Image.new("L", (inner2, inner2), 0)
    ImageDraw.Draw(smask).rounded_rectangle((0, 0, inner2 - 1, inner2 - 1), radius=rad2, fill=255)
    s.putalpha(smask)
    sout = Image.new("RGBA", (size2, size2), (0, 0, 0, 0))
    sout.paste(s, (inset2, inset2), s)
    sedge = Image.new("RGBA", (size2, size2), (0, 0, 0, 0))
    sd = ImageDraw.Draw(sedge)
    sd.rounded_rectangle((2, 2, size2 - 3, size2 - 3), radius=rad2 + 4, outline=(10, 58, 82, 255), width=5)
    sd.rounded_rectangle(
        (inset2, inset2, size2 - inset2 - 1, size2 - inset2 - 1),
        radius=rad2,
        outline=(184, 149, 62, 230),
        width=2,
    )
    Image.alpha_composite(sout, sedge).save(PHOTO_SQUARE)


def styles_base():
    return getSampleStyleSheet()


def add_style(styles, name, **kw):
    if name in styles.byName:
        styles.byName[name] = ParagraphStyle(name, **kw)
        return styles.byName[name]
    styles.add(ParagraphStyle(name, **kw))
    return styles[name]


def bullet(text, style):
    return Paragraph(f"•  {text}", style)


def footer_fn(label: str):
    def _draw(canv, doc):
        canv.saveState()
        canv.setStrokeColor(LINE)
        canv.setLineWidth(0.5)
        y = 8.5 * mm
        canv.line(doc.leftMargin, y, A4[0] - doc.rightMargin, y)
        canv.setFillColor(GOLD)
        canv.circle(doc.leftMargin + 1.2 * mm, y, 1.0, fill=1, stroke=0)
        canv.setFillColor(MUTED)
        canv.setFont("Helvetica", 7)
        canv.drawString(doc.leftMargin + 4.5 * mm, 4.8 * mm, label)
        canv.drawRightString(A4[0] - doc.rightMargin, 4.8 * mm, "One page")
        canv.restoreState()

    return _draw


def skill_chips(skills, width, cols=7):
    styles = styles_base()
    add_style(styles, "Chip", fontName="Helvetica", fontSize=7.2, leading=9,
              textColor=NAVY, alignment=TA_CENTER)
    cells = [Paragraph(s, styles["Chip"]) for s in skills]
    while len(cells) % cols:
        cells.append(Paragraph("", styles["Chip"]))
    rows = [cells[i:i + cols] for i in range(0, len(cells), cols)]
    t = Table(rows, colWidths=[width / cols] * cols)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), SOFT),
        ("BOX", (0, 0), (-1, -1), 0.35, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, LINE),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return t


def edu_block(styles_edu, styles_meta):
    bits = []
    for deg, uni, note in EDUCATION:
        bits.append(Paragraph(deg, styles_edu))
        meta = uni if not note else f"{uni} — {note}"
        bits.append(Paragraph(meta, styles_meta))
    return bits


# ---------------------------------------------------------------------------
# 01 Classic Airy — cream header, portrait photo, dense clear body
# ---------------------------------------------------------------------------
def build_v01_pdf(out: Path):
    styles = styles_base()
    add_style(styles, "N", fontName="Helvetica-Bold", fontSize=15.5, leading=18, textColor=NAVY, spaceAfter=1)
    add_style(styles, "T", fontName="Helvetica", fontSize=8.2, leading=10.5, textColor=TEAL, spaceAfter=3)
    add_style(styles, "C", fontName="Helvetica", fontSize=7.6, leading=10, textColor=MUTED)
    add_style(styles, "H", fontName="Helvetica-Bold", fontSize=8.8, leading=11, textColor=NAVY, spaceBefore=5, spaceAfter=1)
    add_style(styles, "B", fontName="Helvetica", fontSize=8, leading=10.6, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=2)
    add_style(styles, "J", fontName="Helvetica-Bold", fontSize=8.3, leading=10.5, textColor=NAVY, spaceBefore=3, spaceAfter=0)
    add_style(styles, "O", fontName="Helvetica", fontSize=7.6, leading=9.5, textColor=TEAL, spaceAfter=0.5)
    add_style(styles, "L", fontName="Helvetica", fontSize=7.6, leading=9.8, textColor=INK, leftIndent=6, spaceAfter=0.5)
    add_style(styles, "M", fontName="Helvetica", fontSize=7.2, leading=9.5, textColor=MUTED, spaceAfter=1.5)
    add_style(styles, "Bu", fontName="Helvetica", fontSize=7.5, leading=9.8, textColor=INK, spaceAfter=0.8)
    add_style(styles, "E", fontName="Helvetica-Bold", fontSize=7.8, leading=10, textColor=NAVY, spaceBefore=1, spaceAfter=0)
    add_style(styles, "Em", fontName="Helvetica", fontSize=7.2, leading=9.2, textColor=MUTED, spaceAfter=1.5)

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=12 * mm, rightMargin=12 * mm,
                            topMargin=8 * mm, bottomMargin=10 * mm,
                            title=f"{NAME} - Classic Airy")
    w = A4[0] - 24 * mm
    # Portrait photo — complete like original
    photo = RLImage(str(PHOTO_PORTRAIT), width=32 * mm, height=48 * mm)
    head = Table([[
        [Paragraph(NAME, styles["N"]), Paragraph(TITLE, styles["T"]),
         Paragraph(f"{LOCATION}<br/>{MOBILE}<br/>{EMAIL}  ·  {WEB}", styles["C"])],
        photo,
    ]], colWidths=[w - 36 * mm, 36 * mm])
    head.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("BACKGROUND", (0, 0), (-1, -1), CREAM),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (0, 0), 8),
        ("RIGHTPADDING", (1, 0), (1, 0), 7),
        ("LEFTPADDING", (1, 0), (1, 0), 0),
        ("RIGHTPADDING", (0, 0), (0, 0), 4),
    ]))

    story = [
        head, Spacer(1, 3),
        HRFlowable(width="100%", thickness=1.6, color=NAVY, spaceAfter=0.4),
        HRFlowable(width="26%", thickness=1.0, color=GOLD, spaceAfter=3),
        Paragraph("PROFILE", styles["H"]), GoldRule(30 * mm),
        Paragraph(PROFILE, styles["B"]),
        Paragraph(PROFILE2, styles["B"]),
        Paragraph("HIGHLIGHTS", styles["H"]), GoldRule(30 * mm),
    ]
    left_h = [bullet(h, styles["Bu"]) for h in HIGHLIGHTS[:4]]
    right_h = [bullet(h, styles["Bu"]) for h in HIGHLIGHTS[4:]]
    story.append(Table([[left_h, right_h]], colWidths=[w / 2, w / 2], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("BACKGROUND", (0, 0), (-1, -1), SOFT),
        ("BOX", (0, 0), (-1, -1), 0.3, LINE),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ])))
    story += [
        Spacer(1, 2),
        Paragraph("CORE COMPETENCIES", styles["H"]), GoldRule(30 * mm),
        skill_chips(SKILLS, w, cols=7),
        Spacer(1, 2),
        Paragraph("EXPERIENCE", styles["H"]), GoldRule(30 * mm),
    ]
    for role in ROLES:
        block = [
            Paragraph(role["title"], styles["J"]),
            Paragraph(f"{role['org']}   |   {role['dates']}", styles["O"]),
        ]
        for p in role["points"]:
            block.append(Paragraph(f"•  {p}", styles["L"]))
        story.append(KeepTogether(block))
    story.append(Paragraph(EARLIER, styles["M"]))

    # Bottom: education | more+tech in two cols to fill page
    story += [Spacer(1, 2), Paragraph("EDUCATION  ·  TECHNICAL  ·  ADDITIONAL", styles["H"]), GoldRule(50 * mm)]
    edu_bits = edu_block(styles["E"], styles["Em"])
    right_bits = [
        Paragraph("<b>Technical</b>", styles["E"]),
        Paragraph(TECH, styles["Em"]),
        Paragraph("<b>Membership & languages</b>", styles["E"]),
        Paragraph(MORE, styles["Em"]),
    ]
    story.append(Table([[edu_bits, right_bits]], colWidths=[w * 0.52, w * 0.48], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, -1), CREAM),
        ("BOX", (0, 0), (-1, -1), 0.3, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ])))
    doc.build(story, onFirstPage=footer_fn("Classic Airy"), onLaterPages=footer_fn("Classic Airy"))


def build_v01_docx(out: Path):
    _docx_standard(out, PHOTO_PORTRAIT, photo_w=1.15, portrait=True)


# ---------------------------------------------------------------------------
# 02 Sidebar — full-height portrait photo
# ---------------------------------------------------------------------------
def build_v02_pdf(out: Path):
    width, height = A4
    side_w = 58 * mm
    margin = 8 * mm
    c = pdfcanvas.Canvas(str(out), pagesize=A4)

    c.setFillColor(SIDEBAR)
    c.rect(0, 0, side_w, height, fill=1, stroke=0)

    # Full portrait in sidebar (complete waist-up)
    photo_w = 46 * mm
    photo_h = 69 * mm
    c.drawImage(
        str(PHOTO_PORTRAIT),
        (side_w - photo_w) / 2,
        height - photo_h - 8 * mm,
        width=photo_w,
        height=photo_h,
        mask="auto",
        preserveAspectRatio=True,
    )

    y = height - photo_h - 14 * mm

    def side_h(label, y0):
        c.setFillColor(GOLD)
        c.setFont("Helvetica-Bold", 7.5)
        c.drawString(6 * mm, y0, label)
        c.setStrokeColor(GOLD)
        c.setLineWidth(0.8)
        c.line(6 * mm, y0 - 1.8 * mm, side_w - 6 * mm, y0 - 1.8 * mm)
        return y0 - 5.2 * mm

    y = side_h("CONTACT", y)
    c.setFillColor(white)
    c.setFont("Helvetica", 6.9)
    for line in [LOCATION, PHONE1, PHONE2, "info@fairbanks", "medicalcentre.org", "www.fairbanks", "medicalcentre.org"]:
        c.drawString(6 * mm, y, line)
        y -= 3.3 * mm

    y -= 2 * mm
    y = side_h("STRENGTHS", y)
    c.setFillColor(white)
    c.setFont("Helvetica", 6.8)
    for sk in SKILLS[:10]:
        c.drawString(6 * mm, y, f"• {sk}")
        y -= 3.4 * mm

    y -= 2 * mm
    y = side_h("LANGUAGES", y)
    c.setFillColor(white)
    c.setFont("Helvetica", 6.9)
    c.drawString(6 * mm, y, "English — Fluent")
    y -= 3.3 * mm
    c.drawString(6 * mm, y, "Luganda — Fluent")
    y -= 4 * mm
    y = side_h("MEMBERSHIP", y)
    c.setFillColor(white)
    c.setFont("Helvetica", 6.9)
    c.drawString(6 * mm, y, "FUE Member")

    c.setFillColor(HexColor("#8FA3B5"))
    c.setFont("Helvetica", 6.5)
    c.drawString(6 * mm, 7 * mm, "Sidebar Modern · One page")

    x = side_w + margin
    max_w = width - side_w - margin - 9 * mm
    y = height - 12 * mm

    c.setFillColor(NAVY)
    c.setFont("Helvetica-Bold", 11.5)
    for line in _wrap(c, NAME, max_w, "Helvetica-Bold", 11.5):
        c.drawString(x, y, line)
        y -= 4.6 * mm
    c.setFillColor(TEAL)
    c.setFont("Helvetica", 7.6)
    for line in _wrap(c, "Senior HR Executive | Human Capital Strategist | OD", max_w, "Helvetica", 7.6):
        c.drawString(x, y, line)
        y -= 3.6 * mm

    y -= 2 * mm
    y = _sec(c, x, y, "PROFILE")
    y = _text(c, x, y, max_w, PROFILE.replace("<b>", "").replace("</b>", ""), 7.8, 10)
    y = _text(c, x, y, max_w, PROFILE2.replace("<b>", "").replace("</b>", ""), 7.8, 10)
    y -= 1.5
    y = _sec(c, x, y, "HIGHLIGHTS")
    for h in HIGHLIGHTS[:6]:
        y = _text(c, x, y, max_w, f"•  {h}", 7.4, 9.6)
    y -= 1.5
    y = _sec(c, x, y, "EXPERIENCE")
    for role in ROLES:
        c.setFillColor(NAVY)
        c.setFont("Helvetica-Bold", 7.8)
        for line in _wrap(c, role["title"], max_w, "Helvetica-Bold", 7.8):
            c.drawString(x, y, line)
            y -= 3.2 * mm
        c.setFillColor(TEAL)
        c.setFont("Helvetica", 7.2)
        c.drawString(x, y, f"{role['org']}  |  {role['dates']}")
        y -= 3.2 * mm
        for pt in role["points"][:2]:
            y = _text(c, x, y, max_w, f"•  {pt}", 7.2, 9.3)
        y -= 1.2
    y = _text(c, x, y, max_w, EARLIER, 6.8, 8.8, MUTED)
    y -= 1.5
    y = _sec(c, x, y, "EDUCATION")
    for deg, uni, note in EDUCATION:
        c.setFillColor(NAVY)
        c.setFont("Helvetica-Bold", 7.4)
        c.drawString(x, y, deg)
        y -= 3.0 * mm
        y = _text(c, x, y, max_w, uni if not note else f"{uni} — {note}", 6.9, 8.8, MUTED)
    y -= 1
    y = _sec(c, x, y, "TECHNICAL & MORE")
    y = _text(c, x, y, max_w, TECH, 6.9, 8.8)
    y = _text(c, x, y, max_w, MORE, 6.9, 8.8, MUTED)
    c.save()


def build_v02_docx(out: Path):
    doc = Document()
    _page(doc, 0.9, 0.9, 0.9, 0.9)
    table = doc.add_table(rows=1, cols=2)
    side, main = table.rows[0].cells
    side.width = Cm(5.4)
    main.width = Cm(14)
    _shade(side, "0A3A52")
    sp = side.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run().add_picture(str(PHOTO_PORTRAIT), width=Inches(1.55))

    def sh(text):
        p = side.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(text)
        _run(r, 8, True, DOCX_GOLD)

    def st(text):
        p = side.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        _run(r, 7.5, False, RGBColor(255, 255, 255))

    sh("CONTACT")
    for line in [LOCATION, MOBILE, EMAIL, WEB]:
        st(line)
    sh("STRENGTHS")
    for sk in SKILLS[:10]:
        st(f"• {sk}")
    sh("LANGUAGES")
    st("English — Fluent")
    st("Luganda — Fluent")
    sh("MEMBERSHIP")
    st("FUE Member")

    p = main.paragraphs[0]
    r = p.add_run(NAME)
    _run(r, 12, True, DOCX_NAVY)
    p = main.add_paragraph()
    r = p.add_run(TITLE)
    _run(r, 8.5, False, DOCX_TEAL)
    _h(doc, "PROFILE", cell=main, size=8.5)
    _p_in(main, PROFILE.replace("<b>", "").replace("</b>", ""), size=8)
    _p_in(main, PROFILE2.replace("<b>", "").replace("</b>", ""), size=8)
    _h(doc, "HIGHLIGHTS", cell=main, size=8.5)
    for h in HIGHLIGHTS[:6]:
        _b_in(main, h, size=7.5)
    _h(doc, "EXPERIENCE", cell=main, size=8.5)
    for role in ROLES:
        p = main.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run(role["title"])
        _run(r, 8.5, True, DOCX_NAVY)
        p = main.add_paragraph()
        r = p.add_run(f"{role['org']} | {role['dates']}")
        _run(r, 7.5, False, DOCX_TEAL)
        for pt in role["points"][:2]:
            _b_in(main, pt, size=7.5)
    _p_in(main, EARLIER, muted=True, size=7)
    _h(doc, "EDUCATION", cell=main, size=8.5)
    for deg, uni, note in EDUCATION:
        _p_in(main, f"{deg} — {uni}" + (f" ({note})" if note else ""), size=7.5)
    _h(doc, "TECHNICAL & MORE", cell=main, size=8.5)
    _p_in(main, TECH, size=7.5)
    _p_in(main, MORE, muted=True, size=7.5)
    doc.save(str(out))


# ---------------------------------------------------------------------------
# 03 Compact one-page — square photo, max information density
# ---------------------------------------------------------------------------
def build_v03_pdf(out: Path):
    styles = styles_base()
    add_style(styles, "N", fontName="Helvetica-Bold", fontSize=13.5, leading=15.5, textColor=CHARCOAL, spaceAfter=1)
    add_style(styles, "T", fontName="Helvetica", fontSize=7.6, leading=9.5, textColor=FOREST, spaceAfter=1)
    add_style(styles, "C", fontName="Helvetica", fontSize=7.2, leading=9.2, textColor=MUTED)
    add_style(styles, "H", fontName="Helvetica-Bold", fontSize=7.8, leading=9.5, textColor=FOREST, spaceBefore=4, spaceAfter=1)
    add_style(styles, "B", fontName="Helvetica", fontSize=7.4, leading=9.6, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=2)
    add_style(styles, "J", fontName="Helvetica-Bold", fontSize=7.6, leading=9.5, textColor=CHARCOAL, spaceBefore=2, spaceAfter=0)
    add_style(styles, "O", fontName="Helvetica", fontSize=7.1, leading=9, textColor=FOREST, spaceAfter=0.3)
    add_style(styles, "L", fontName="Helvetica", fontSize=7.1, leading=9.2, textColor=INK, leftIndent=5, spaceAfter=0.3)
    add_style(styles, "M", fontName="Helvetica", fontSize=6.8, leading=8.8, textColor=MUTED, spaceAfter=1)
    add_style(styles, "Bu", fontName="Helvetica", fontSize=7.1, leading=9.2, textColor=INK, spaceAfter=0.5)
    add_style(styles, "E", fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=CHARCOAL, spaceAfter=0)
    add_style(styles, "Em", fontName="Helvetica", fontSize=6.8, leading=8.6, textColor=MUTED, spaceAfter=1)

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=10 * mm, rightMargin=10 * mm,
                            topMargin=7 * mm, bottomMargin=9 * mm,
                            title=f"{NAME} - One Page")
    w = A4[0] - 20 * mm
    photo = RLImage(str(PHOTO_SQUARE), width=26 * mm, height=26 * mm)
    head = Table([[
        [Paragraph(NAME, styles["N"]), Paragraph(TITLE, styles["T"]),
         Paragraph(f"{LOCATION}  ·  {MOBILE}  ·  {EMAIL}  ·  {WEB}", styles["C"])],
        photo,
    ]], colWidths=[w - 30 * mm, 30 * mm])
    head.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ]))
    story = [
        head, Spacer(1, 2),
        HRFlowable(width="100%", thickness=1.3, color=FOREST, spaceAfter=0.4),
        HRFlowable(width="20%", thickness=0.8, color=GOLD, spaceAfter=2),
        Paragraph("PROFILE", styles["H"]), GoldRule(26 * mm, 1.0),
        Paragraph(PROFILE + " " + PROFILE2, styles["B"]),
        Paragraph("HIGHLIGHTS & COMPETENCIES", styles["H"]), GoldRule(40 * mm, 1.0),
    ]
    left = [bullet(h, styles["Bu"]) for h in HIGHLIGHTS]
    right = [Paragraph(f"•  {s}", styles["Bu"]) for s in SKILLS]
    story.append(Table([[left, right]], colWidths=[w * 0.5, w * 0.5], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, -1), SOFT),
        ("BOX", (0, 0), (-1, -1), 0.3, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ])))
    story += [Spacer(1, 2), Paragraph("EXPERIENCE", styles["H"]), GoldRule(26 * mm, 1.0)]
    for role in ROLES:
        story.append(Paragraph(role["title"], styles["J"]))
        story.append(Paragraph(f"{role['org']}  |  {role['dates']}", styles["O"]))
        for pt in role["points"][:2]:
            story.append(Paragraph(f"•  {pt}", styles["L"]))
    story += [
        Paragraph(EARLIER, styles["M"]),
        Spacer(1, 1),
        Paragraph("EDUCATION · TECHNICAL · MORE", styles["H"]), GoldRule(40 * mm, 1.0),
    ]
    edu = edu_block(styles["E"], styles["Em"])
    more = [
        Paragraph("<b>Technical</b>", styles["E"]),
        Paragraph(TECH, styles["Em"]),
        Paragraph("<b>Membership & languages</b>", styles["E"]),
        Paragraph(MORE, styles["Em"]),
    ]
    story.append(Table([[edu, more]], colWidths=[w * 0.5, w * 0.5], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, -1), CREAM),
        ("BOX", (0, 0), (-1, -1), 0.3, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ])))
    doc.build(story, onFirstPage=footer_fn("One Page"), onLaterPages=footer_fn("One Page"))


def build_v03_docx(out: Path):
    _docx_standard(out, PHOTO_SQUARE, photo_w=1.0, portrait=False, compact=True)


# ---------------------------------------------------------------------------
# 04 Timeline — circle photo from complete crop, packed timeline
# ---------------------------------------------------------------------------
def build_v04_pdf(out: Path):
    styles = styles_base()
    add_style(styles, "N", fontName="Helvetica-Bold", fontSize=14.5, leading=17, textColor=SLATE, alignment=TA_CENTER, spaceAfter=1)
    add_style(styles, "T", fontName="Helvetica", fontSize=8, leading=10, textColor=TEAL, alignment=TA_CENTER, spaceAfter=1)
    add_style(styles, "C", fontName="Helvetica", fontSize=7.3, leading=9.5, textColor=MUTED, alignment=TA_CENTER, spaceAfter=3)
    add_style(styles, "H", fontName="Helvetica-Bold", fontSize=8.2, leading=10, textColor=SLATE, alignment=TA_CENTER, spaceBefore=4, spaceAfter=2)
    add_style(styles, "B", fontName="Helvetica", fontSize=7.6, leading=10, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=3)
    add_style(styles, "D", fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=TEAL, alignment=TA_RIGHT)
    add_style(styles, "J", fontName="Helvetica-Bold", fontSize=7.8, leading=9.8, textColor=SLATE)
    add_style(styles, "O", fontName="Helvetica", fontSize=7.2, leading=9, textColor=MUTED, spaceAfter=0.5)
    add_style(styles, "L", fontName="Helvetica", fontSize=7.2, leading=9.2, textColor=INK, leftIndent=4, spaceAfter=0.3)
    add_style(styles, "M", fontName="Helvetica", fontSize=7, leading=9, textColor=MUTED, alignment=TA_CENTER, spaceAfter=2)
    add_style(styles, "Bu", fontName="Helvetica", fontSize=7.1, leading=9.2, textColor=INK, spaceAfter=0.5)

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=11 * mm, rightMargin=11 * mm,
                            topMargin=7 * mm, bottomMargin=9 * mm,
                            title=f"{NAME} - Timeline")
    w = A4[0] - 22 * mm
    photo = RLImage(str(PHOTO_CIRCLE), width=30 * mm, height=30 * mm)
    ph = Table([[photo]], colWidths=[w])
    ph.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")]))

    story = [
        Paragraph(NAME, styles["N"]),
        Paragraph(TITLE, styles["T"]),
        Paragraph(f"{LOCATION}  ·  {MOBILE}  ·  {EMAIL}", styles["C"]),
        ph, Spacer(1, 2),
        HRFlowable(width="22%", thickness=1.0, color=GOLD, spaceAfter=3),
        Paragraph("PROFILE", styles["H"]), GoldRule(28 * mm),
        Paragraph(PROFILE + " " + PROFILE2, styles["B"]),
        Paragraph("CAREER TIMELINE", styles["H"]), GoldRule(28 * mm),
    ]
    for role in ROLES:
        left = Paragraph(role["dates"].replace(" - ", "<br/>"), styles["D"])
        right = [Paragraph(role["title"], styles["J"]), Paragraph(role["org"], styles["O"])]
        for pt in role["points"][:2]:
            right.append(Paragraph(f"•  {pt}", styles["L"]))
        row = Table([[left, right]], colWidths=[24 * mm, w - 24 * mm])
        row.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BACKGROUND", (1, 0), (1, 0), SOFT),
            ("LINEBEFORE", (1, 0), (1, 0), 1.8, TEAL),
            ("LEFTPADDING", (0, 0), (0, 0), 0),
            ("RIGHTPADDING", (0, 0), (0, 0), 5),
            ("LEFTPADDING", (1, 0), (1, 0), 7),
            ("RIGHTPADDING", (1, 0), (1, 0), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
        ]))
        story.append(row)
        story.append(Spacer(1, 2))
    story.append(Paragraph(EARLIER, styles["M"]))

    # Highlights + skills fill remaining space
    story += [Paragraph("HIGHLIGHTS & SKILLS", styles["H"]), GoldRule(36 * mm)]
    hl = [bullet(h, styles["Bu"]) for h in HIGHLIGHTS[:6]]
    sk = [Paragraph(f"•  {s}", styles["Bu"]) for s in SKILLS[:10]]
    story.append(Table([[hl, sk]], colWidths=[w / 2, w / 2], style=TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, -1), CREAM),
        ("BOX", (0, 0), (-1, -1), 0.3, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ])))
    story += [
        Spacer(1, 2),
        Paragraph("EDUCATION · TECHNICAL · MORE", styles["H"]), GoldRule(40 * mm),
        Paragraph(
            "PhD Management (Ongoing), Uganda Christian University — AI/ML burnout research  ·  "
            "MSSPM, Makerere  ·  BA Social Sciences, Makerere<br/>"
            + TECH + "<br/>" + MORE,
            styles["M"],
        ),
    ]
    doc.build(story, onFirstPage=footer_fn("Timeline"), onLaterPages=footer_fn("Timeline"))


def build_v04_docx(out: Path):
    doc = Document()
    _page(doc, 1.0, 1.0, 1.2, 1.2)
    for text, size, bold, color in [
        (NAME, 13, True, RGBColor(0x2F, 0x44, 0x58)),
        (TITLE, 8.5, False, DOCX_TEAL),
        (f"{LOCATION}  ·  {MOBILE}  ·  {EMAIL}", 8, False, DOCX_MUTED),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        _run(r, size, bold, color)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(PHOTO_CIRCLE), width=Inches(1.15))
    _h(doc, "PROFILE", center=True, size=8.5)
    _p(doc, (PROFILE + " " + PROFILE2).replace("<b>", "").replace("</b>", ""), size=8)
    _h(doc, "CAREER TIMELINE", center=True, size=8.5)
    for role in ROLES:
        p = doc.add_paragraph()
        r = p.add_run(f"{role['dates']}  |  {role['title']}")
        _run(r, 8.5, True, DOCX_NAVY)
        p = doc.add_paragraph()
        r = p.add_run(role["org"])
        _run(r, 8, False, DOCX_TEAL)
        for pt in role["points"][:2]:
            _b(doc, pt, size=7.5)
    _p(doc, EARLIER, muted=True, size=7.5)
    _h(doc, "EDUCATION · MORE", center=True, size=8.5)
    _p(doc, MORE + "  ·  " + TECH, size=7.5)
    doc.save(str(out))


# ---------------------------------------------------------------------------
# 05 Editorial split — portrait in cream banner
# ---------------------------------------------------------------------------
def build_v05_pdf(out: Path):
    styles = styles_base()
    add_style(styles, "Tag", fontName="Helvetica-Bold", fontSize=7, leading=8.5, textColor=GOLD, spaceAfter=2)
    add_style(styles, "N", fontName="Helvetica-Bold", fontSize=14, leading=16.5, textColor=NAVY, spaceAfter=1)
    add_style(styles, "T", fontName="Helvetica", fontSize=7.8, leading=10, textColor=TEAL, spaceAfter=2)
    add_style(styles, "C", fontName="Helvetica", fontSize=7.2, leading=9.5, textColor=MUTED)
    add_style(styles, "H", fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=NAVY, spaceBefore=1, spaceAfter=1.5)
    add_style(styles, "B", fontName="Helvetica", fontSize=7.5, leading=9.8, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=2)
    add_style(styles, "J", fontName="Helvetica-Bold", fontSize=7.6, leading=9.5, textColor=NAVY, spaceBefore=2.5, spaceAfter=0)
    add_style(styles, "O", fontName="Helvetica", fontSize=7, leading=8.8, textColor=TEAL, spaceAfter=0.3)
    add_style(styles, "L", fontName="Helvetica", fontSize=7.1, leading=9, textColor=INK, leftIndent=4, spaceAfter=0.3)
    add_style(styles, "S", fontName="Helvetica", fontSize=7, leading=9.2, textColor=INK, spaceAfter=0.6)
    add_style(styles, "M", fontName="Helvetica", fontSize=6.9, leading=8.8, textColor=MUTED, spaceAfter=1.2)
    add_style(styles, "Bu", fontName="Helvetica", fontSize=7, leading=9.1, textColor=INK, spaceAfter=0.7)
    add_style(styles, "E", fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=NAVY, spaceAfter=0)

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=9 * mm, rightMargin=9 * mm,
                            topMargin=7 * mm, bottomMargin=9 * mm,
                            title=f"{NAME} - Editorial Split")
    w = A4[0] - 18 * mm
    photo = RLImage(str(PHOTO_PORTRAIT), width=30 * mm, height=45 * mm)
    banner = Table([[
        [
            Paragraph("CURRICULUM VITAE", styles["Tag"]),
            Paragraph(NAME, styles["N"]),
            Paragraph(TITLE, styles["T"]),
            Paragraph(f"{LOCATION}  ·  {MOBILE}<br/>{EMAIL}  ·  {WEB}", styles["C"]),
        ],
        photo,
    ]], colWidths=[w - 36 * mm, 36 * mm])
    banner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), CREAM),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))

    left = [Paragraph("AT A GLANCE", styles["H"]), GoldRule(80)]
    for h in HIGHLIGHTS:
        left.append(bullet(h, styles["Bu"]))
    left += [Spacer(1, 4), Paragraph("SKILLS", styles["H"]), GoldRule(80)]
    for sk in SKILLS:
        left.append(Paragraph(f"•  {sk}", styles["S"]))
    left += [
        Spacer(1, 4), Paragraph("EDUCATION", styles["H"]), GoldRule(80),
        Paragraph("<b>PhD Management</b> (Ongoing)<br/>Uganda Christian University<br/>AI/ML burnout research", styles["M"]),
        Paragraph("<b>MSSPM</b> — Makerere University", styles["M"]),
        Paragraph("<b>BA Social Sciences</b> — Makerere", styles["M"]),
        Spacer(1, 3), Paragraph("MORE", styles["H"]), GoldRule(80),
        Paragraph("FUE Member<br/>English & Luganda — Fluent<br/>References on request", styles["M"]),
        Spacer(1, 2), Paragraph("TECHNICAL", styles["H"]), GoldRule(80),
        Paragraph(TECH, styles["M"]),
    ]

    right = [
        Paragraph("PROFILE", styles["H"]), GoldRule(80),
        Paragraph(PROFILE, styles["B"]),
        Paragraph(PROFILE2, styles["B"]),
        Spacer(1, 2),
        Paragraph("EXPERIENCE", styles["H"]), GoldRule(80),
    ]
    for role in ROLES:
        right += [
            Paragraph(role["title"], styles["J"]),
            Paragraph(f"{role['org']}  |  {role['dates']}", styles["O"]),
        ]
        for pt in role["points"]:
            right.append(Paragraph(f"•  {pt}", styles["L"]))
    right.append(Paragraph(EARLIER, styles["M"]))

    split = Table([[left, right]], colWidths=[w * 0.36, w * 0.64])
    split.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (0, 0), SOFT),
        ("LEFTPADDING", (0, 0), (0, 0), 6),
        ("RIGHTPADDING", (0, 0), (0, 0), 6),
        ("TOPPADDING", (0, 0), (0, 0), 6),
        ("BOTTOMPADDING", (0, 0), (0, 0), 6),
        ("LEFTPADDING", (1, 0), (1, 0), 8),
        ("RIGHTPADDING", (1, 0), (1, 0), 2),
        ("TOPPADDING", (1, 0), (1, 0), 1),
        ("LINEAFTER", (0, 0), (0, 0), 0.5, LINE),
    ]))
    doc.build([banner, Spacer(1, 4), split],
              onFirstPage=footer_fn("Editorial Split"),
              onLaterPages=footer_fn("Editorial Split"))


def build_v05_docx(out: Path):
    doc = Document()
    _page(doc, 0.9, 0.9, 1.0, 1.0)
    head = doc.add_table(rows=1, cols=2)
    L, R = head.rows[0].cells
    _shade(L, "F8F5F0")
    _shade(R, "F8F5F0")
    p = L.paragraphs[0]
    r = p.add_run("CURRICULUM VITAE")
    _run(r, 7.5, True, DOCX_GOLD)
    p = L.add_paragraph()
    r = p.add_run(NAME)
    _run(r, 13, True, DOCX_NAVY)
    p = L.add_paragraph()
    r = p.add_run(TITLE)
    _run(r, 8, False, DOCX_TEAL)
    p = L.add_paragraph()
    r = p.add_run(f"{LOCATION} · {MOBILE} · {EMAIL}")
    _run(r, 7.5, False, DOCX_MUTED)
    R.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R.paragraphs[0].add_run().add_picture(str(PHOTO_PORTRAIT), width=Inches(1.1))

    body = doc.add_table(rows=1, cols=2)
    left, right = body.rows[0].cells
    _shade(left, "F4F7FA")
    left.paragraphs[0].clear()
    p = left.paragraphs[0]
    r = p.add_run("AT A GLANCE")
    _run(r, 8.5, True, DOCX_NAVY)
    for h in HIGHLIGHTS:
        _b_in(left, h, size=7.5)
    p = left.add_paragraph()
    r = p.add_run("SKILLS")
    _run(r, 8.5, True, DOCX_NAVY)
    for sk in SKILLS:
        _b_in(left, sk, size=7.5)
    p = left.add_paragraph()
    r = p.add_run("EDUCATION & MORE")
    _run(r, 8.5, True, DOCX_NAVY)
    for deg, uni, note in EDUCATION:
        _p_in(left, f"{deg} — {uni}", size=7.5)
    _p_in(left, MORE, size=7.5)

    right.paragraphs[0].clear()
    p = right.paragraphs[0]
    r = p.add_run("PROFILE")
    _run(r, 8.5, True, DOCX_NAVY)
    _p_in(right, PROFILE.replace("<b>", "").replace("</b>", ""), size=8)
    _p_in(right, PROFILE2.replace("<b>", "").replace("</b>", ""), size=8)
    p = right.add_paragraph()
    r = p.add_run("EXPERIENCE")
    _run(r, 8.5, True, DOCX_NAVY)
    for role in ROLES:
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(role["title"])
        _run(r, 8.5, True, DOCX_NAVY)
        p = right.add_paragraph()
        r = p.add_run(f"{role['org']} | {role['dates']}")
        _run(r, 7.5, False, DOCX_TEAL)
        for pt in role["points"][:2]:
            _b_in(right, pt, size=7.5)
    _p_in(right, EARLIER, muted=True, size=7)
    doc.save(str(out))


# ---------------------------------------------------------------------------
# Shared DOCX helper for classic/compact
# ---------------------------------------------------------------------------
def _docx_standard(out: Path, photo: Path, photo_w=1.1, portrait=False, compact=False):
    doc = Document()
    _page(doc, 1.0, 1.0, 1.2, 1.2)
    t = doc.add_table(rows=1, cols=2)
    L, R = t.rows[0].cells
    _shade(L, "F8F5F0")
    _shade(R, "F8F5F0")
    p = L.paragraphs[0]
    r = p.add_run(NAME)
    _run(r, 13 if compact else 14, True, DOCX_NAVY)
    p = L.add_paragraph()
    r = p.add_run(TITLE)
    _run(r, 8, False, DOCX_TEAL)
    for line in (LOCATION, MOBILE, f"{EMAIL}  ·  {WEB}"):
        p = L.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _run(r, 8, False, DOCX_MUTED)
    R.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R.paragraphs[0].add_run().add_picture(str(photo), width=Inches(photo_w))
    _h(doc, "PROFILE", size=9)
    _p(doc, PROFILE.replace("<b>", "").replace("</b>", ""), size=8.5)
    _p(doc, PROFILE2.replace("<b>", "").replace("</b>", ""), size=8.5)
    _h(doc, "HIGHLIGHTS", size=9)
    for h in HIGHLIGHTS:
        _b(doc, h, size=8)
    _h(doc, "COMPETENCIES", size=9)
    _p(doc, "  ·  ".join(SKILLS), size=8)
    _h(doc, "EXPERIENCE", size=9)
    for role in ROLES:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run(role["title"])
        _run(r, 9, True, DOCX_NAVY)
        p = doc.add_paragraph()
        r = p.add_run(f"{role['org']}  |  {role['dates']}")
        _run(r, 8, False, DOCX_TEAL)
        for pt in role["points"][:2]:
            _b(doc, pt, size=8)
    _p(doc, EARLIER, muted=True, size=7.5)
    _h(doc, "EDUCATION", size=9)
    for deg, uni, note in EDUCATION:
        _p(doc, f"{deg} — {uni}" + (f". {note}" if note else ""), size=8)
    _h(doc, "TECHNICAL & ADDITIONAL", size=9)
    _p(doc, TECH, size=8)
    _p(doc, MORE, muted=True, size=8)
    doc.save(str(out))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _wrap(c, text, max_w, font, size):
    c.setFont(font, size)
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip()
        if c.stringWidth(trial, font, size) <= max_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _sec(c, x, y, title):
    c.setFillColor(NAVY)
    c.setFont("Helvetica-Bold", 8.2)
    c.drawString(x, y, title)
    y -= 2 * mm
    c.setStrokeColor(GOLD)
    c.setLineWidth(1.1)
    c.line(x, y, x + 24 * mm, y)
    return y - 3.5 * mm


def _text(c, x, y, max_w, text, size, leading, color=INK):
    c.setFillColor(color)
    c.setFont("Helvetica", size)
    for line in _wrap(c, text, max_w, "Helvetica", size):
        c.drawString(x, y, line)
        y -= leading
    return y - 1.2


def _page(doc, top, bottom, left, right):
    for sec in doc.sections:
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)


def _run(run, size=10, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _h(doc, text, size=9, center=False, cell=None):
    target = cell if cell is not None else doc
    p = target.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _run(r, size, True, DOCX_NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B8953E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _p(doc, text, muted=False, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _run(r, size, False, DOCX_MUTED if muted else DOCX_INK)


def _b(doc, text, size=8.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    _run(r, size, False, DOCX_INK)


def _p_in(cell, text, muted=False, size=8.5):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    _run(r, size, False, DOCX_MUTED if muted else DOCX_INK)


def _b_in(cell, text, size=8):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    _run(r, size, False, DOCX_INK)


def assert_one_page(pdf_path: Path):
    import fitz
    doc = fitz.open(str(pdf_path))
    n = doc.page_count
    doc.close()
    if n != 1:
        raise RuntimeError(f"{pdf_path.name} has {n} pages; expected 1")


def render_previews():
    import fitz
    for slug in VERSION_SLUGS:
        pdf = ROOT / slug / "racheal_cv.pdf"
        out_dir = PREVIEW / slug
        out_dir.mkdir(parents=True, exist_ok=True)
        for old in out_dir.glob("*.png"):
            old.unlink()
        doc = fitz.open(str(pdf))
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        pix.save(str(out_dir / "page_1.png"))
        print(f"  {slug}: {doc.page_count} page(s)")
        doc.close()


def main():
    if not PHOTO_SRC.exists():
        raise SystemExit(f"Missing photo: {PHOTO_SRC}")
    ensure_dirs()
    prepare_photos()
    builders = [
        ("01_classic_airy", build_v01_pdf, build_v01_docx),
        ("02_sidebar_modern", build_v02_pdf, build_v02_docx),
        ("03_one_page", build_v03_pdf, build_v03_docx),
        ("04_timeline", build_v04_pdf, build_v04_docx),
        ("05_editorial_split", build_v05_pdf, build_v05_docx),
    ]
    for slug, pdf_fn, docx_fn in builders:
        folder = ROOT / slug
        pdf_path = folder / "racheal_cv.pdf"
        docx_path = folder / "racheal_cv.docx"
        print(f"Building {slug}...")
        pdf_fn(pdf_path)
        docx_fn(docx_path)
        assert_one_page(pdf_path)
        print("  OK -> pdf + docx (1 page)")
    print("Rendering previews...")
    render_previews()
    print("Done.")


if __name__ == "__main__":
    main()
