#!/usr/bin/env python3
"""Build five detailed CV designs from racheal_cv.md.

The Markdown file is the sole content source. All PDF and Word deliverables
are written into one detailed_versions folder.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.platypus import (
    Flowable,
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "racheal_cv.md"
PHOTO = ROOT / "_assets" / "photo_portrait.png"
SIGNATURE = ROOT / "signature.jpeg"
OUTPUT = ROOT / "detailed_versions"
BUILD_TMP = ROOT.parent / "tmp" / "detailed_cv_build"
PREVIEW = ROOT.parent / "tmp" / "detailed_cv_previews"
APPLICANT = "racheal_nabukeera"


@dataclass(frozen=True)
class Theme:
    slug: str
    label: str
    primary: str
    secondary: str
    accent: str
    soft: str
    header_bg: str
    header_light: bool
    body_font: str
    heading_font: str
    layout: str


THEMES = [
    Theme(
        "modern_navy",
        "Modern Navy",
        "#0B2942",
        "#2F7185",
        "#D2AA52",
        "#EEF4F7",
        "#0B2942",
        False,
        "Helvetica",
        "Helvetica-Bold",
        "modern",
    ),
]


def clean_inline(text: str) -> str:
    """Convert Markdown inline markup to clean printable text."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    text = text.replace("*(Ongoing)*", "(Ongoing)")
    text = text.replace("*", "")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u00a0", " ")
    return text.strip()


def parse_source() -> tuple[dict[str, str], list[tuple[str, str]]]:
    """Read header details and ordered body blocks directly from Markdown."""
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    header = {
        "name": clean_inline(lines[0].lstrip("#").strip()),
        "title": clean_inline(lines[2]),
        "location": clean_inline(lines[4]),
        "mobile": clean_inline(lines[6]).replace("Mobile:", "").strip(),
        "email": clean_inline(lines[8]).replace("Email:", "").strip(),
        "website": clean_inline(lines[10]).replace("Website:", "").strip(),
    }

    blocks: list[tuple[str, str]] = []
    for raw in lines[12:]:
        line = raw.strip()
        if not line:
            continue
        if line == "---":
            blocks.append(("separator", ""))
        elif line.startswith("# "):
            blocks.append(("section", clean_inline(line[2:])))
        elif line.startswith("## "):
            blocks.append(("role", clean_inline(line[3:])))
        elif line.startswith("### "):
            blocks.append(("subheading", clean_inline(line[4:])))
        elif line.startswith("* "):
            blocks.append(("bullet", clean_inline(line[2:])))
        else:
            blocks.append(("paragraph", clean_inline(line)))
    return header, blocks


def extract_section(
    blocks: list[tuple[str, str]], section_name: str
) -> tuple[list[tuple[str, str]], list[tuple[str, str]]]:
    """Split out one top-level section; leave remaining blocks unchanged."""
    target = section_name.upper()
    section: list[tuple[str, str]] = []
    remainder: list[tuple[str, str]] = []
    capturing = False
    for kind, text in blocks:
        if kind == "section":
            if text.upper() == target:
                capturing = True
                section.append((kind, text))
                continue
            if capturing:
                capturing = False
        if capturing:
            section.append((kind, text))
        else:
            remainder.append((kind, text))
    return section, remainder


def languages_above_summary(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Place Languages immediately above Executive Profile; keep other order."""
    languages, remainder = extract_section(blocks, "LANGUAGES")
    if not languages:
        return blocks

    # Drop a trailing separator that belonged only to Languages.
    while remainder and remainder[-1][0] == "separator":
        remainder.pop()

    rebuilt: list[tuple[str, str]] = []
    inserted = False
    for kind, text in remainder:
        if kind == "section" and text.upper() == "EXECUTIVE PROFILE" and not inserted:
            rebuilt.extend(languages)
            if rebuilt and rebuilt[-1][0] != "separator":
                rebuilt.append(("separator", ""))
            inserted = True
        rebuilt.append((kind, text))
    if not inserted:
        rebuilt = languages + [("separator", "")] + remainder
    return rebuilt


class NumberedCanvas(pdf_canvas.Canvas):
    """Defer page drawing so footer can show Page X of Y."""

    def __init__(self, *args, name_label: str = "", primary: str = "#0A3A52", **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states: list[dict] = []
        self.name_label = name_label
        self.primary = primary

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer(page_count)
            super().showPage()
        super().save()

    def _draw_footer(self, page_count: int):
        self.saveState()
        self.setStrokeColor(HexColor(self.primary))
        self.setLineWidth(0.5)
        self.line(14 * mm, 9 * mm, A4[0] - 14 * mm, 9 * mm)
        self.setFillColor(HexColor("#6B7785"))
        self.setFont("Helvetica", 7.3)
        self.drawString(14 * mm, 5.2 * mm, self.name_label)
        self.drawRightString(
            A4[0] - 14 * mm,
            5.2 * mm,
            f"Page {self._pageNumber} of {page_count}",
        )
        self.restoreState()


class SignatureBlock(Flowable):
    """Compact proper signing block that always stays on the current page."""

    def __init__(self, theme: Theme, signer_name: str):
        super().__init__()
        with PILImage.open(SIGNATURE) as img:
            width_px, height_px = img.size
        self.aspect = height_px / max(width_px, 1)
        self.theme = theme
        self.signer_name = signer_name
        # Compact target size so the block fits on page 5 with content.
        self.sig_w = 30 * mm
        self.sig_h = self.sig_w * self.aspect
        self.block_h = self.sig_h + 22  # label + image + printed name

    def wrap(self, avail_width, avail_height):
        self.width = avail_width
        # Stay on this page always. Use leftover space to sit at the bottom;
        # shrink if the leftover strip is shorter than the ideal block.
        min_h = 18 * mm
        self.height = max(avail_height, 1)
        if avail_height >= self.block_h:
            self.draw_sig_w = self.sig_w
            self.draw_sig_h = self.sig_h
        else:
            usable = max(10 * mm, min(avail_height, self.block_h) - 18)
            self.draw_sig_w = min(self.sig_w, usable / max(self.aspect, 0.01))
            self.draw_sig_h = self.draw_sig_w * self.aspect
            if avail_height < min_h:
                self.draw_sig_w = 16 * mm
                self.draw_sig_h = self.draw_sig_w * self.aspect
        return self.width, self.height

    def draw(self):
        muted = HexColor("#607080")
        primary = HexColor(self.theme.primary)
        left = 0
        # Flowable y=0 is the bottom of the remaining page space.
        y_name = 2
        y_img = y_name + 10
        y_label = y_img + self.draw_sig_h + 3

        self.canv.setFillColor(primary)
        self.canv.setFont(self.theme.heading_font, 8)
        self.canv.drawString(left, y_name, self.signer_name)

        self.canv.drawImage(
            str(SIGNATURE),
            left,
            y_img,
            width=self.draw_sig_w,
            height=self.draw_sig_h,
            mask="auto",
            preserveAspectRatio=True,
        )

        self.canv.setFillColor(muted)
        self.canv.setFont(self.theme.body_font, 7.5)
        self.canv.drawString(left, y_label, "Signed")
        self.canv.setStrokeColor(HexColor(self.theme.accent))
        self.canv.setLineWidth(0.8)
        self.canv.line(left, y_label - 1.5, left + self.draw_sig_w, y_label - 1.5)


def signature_flowables(theme: Theme, signer_name: str) -> list:
    """Proper signing block kept on the final content page."""
    return [Spacer(1, 6), SignatureBlock(theme, signer_name)]


class SectionBand(Flowable):
    def __init__(self, text: str, theme: Theme):
        super().__init__()
        self.text = text
        self.theme = theme
        self.height = 22

    def wrap(self, avail_width, avail_height):
        self.width = avail_width
        return avail_width, self.height

    def draw(self):
        primary = HexColor(self.theme.primary)
        accent = HexColor(self.theme.accent)
        if self.theme.layout in {"modern", "editorial"}:
            self.canv.setFillColor(primary)
            self.canv.roundRect(0, 1, self.width, 19, 3, fill=1, stroke=0)
            self.canv.setFillColor(white)
            self.canv.setFont(self.theme.heading_font, 10)
            self.canv.drawString(8, 7, self.text.upper())
        else:
            self.canv.setFillColor(primary)
            self.canv.setFont(self.theme.heading_font, 11)
            self.canv.drawString(0, 8, self.text.upper())
            self.canv.setStrokeColor(accent)
            self.canv.setLineWidth(1.6)
            self.canv.line(0, 3, min(self.width, 105), 3)


def pdf_styles(theme: Theme):
    styles = getSampleStyleSheet()
    primary = HexColor(theme.primary)
    secondary = HexColor(theme.secondary)
    ink = HexColor("#25313D")
    muted = HexColor("#607080")

    def add(name, **kwargs):
        styles.add(ParagraphStyle(name=name, **kwargs))

    add(
        "CvName",
        fontName=theme.heading_font,
        fontSize=19,
        leading=22,
        textColor=primary,
        spaceAfter=3,
    )
    add(
        "CvNameWhite",
        fontName=theme.heading_font,
        fontSize=19,
        leading=22,
        textColor=white,
        spaceAfter=3,
    )
    add(
        "CvTitle",
        fontName=theme.body_font,
        fontSize=9.5,
        leading=12,
        textColor=secondary,
        spaceAfter=5,
    )
    add(
        "CvTitleWhite",
        fontName=theme.body_font,
        fontSize=9.5,
        leading=12,
        textColor=HexColor("#DCEAF0"),
        spaceAfter=5,
    )
    add(
        "CvContact",
        fontName=theme.body_font,
        fontSize=8.5,
        leading=11,
        textColor=muted,
    )
    add(
        "CvContactWhite",
        fontName=theme.body_font,
        fontSize=8.5,
        leading=11,
        textColor=white,
    )
    add(
        "CvBody",
        fontName=theme.body_font,
        fontSize=9.3,
        leading=12.8,
        textColor=ink,
        alignment=TA_JUSTIFY,
        spaceAfter=4.5,
    )
    add(
        "CvBullet",
        fontName=theme.body_font,
        fontSize=9.1,
        leading=12.2,
        textColor=ink,
        leftIndent=13,
        firstLineIndent=-8,
        spaceAfter=1.8,
    )
    add(
        "CvRole",
        fontName=theme.heading_font,
        fontSize=10.8,
        leading=13.2,
        textColor=primary,
        spaceBefore=5,
        spaceAfter=1.5,
    )
    add(
        "CvSub",
        fontName=theme.heading_font,
        fontSize=9.4,
        leading=11.5,
        textColor=secondary,
        spaceBefore=4,
        spaceAfter=3,
    )
    add(
        "CvSmall",
        fontName=theme.body_font,
        fontSize=8,
        leading=10,
        textColor=muted,
    )
    return styles


def pdf_header(header: dict[str, str], theme: Theme, styles):
    header_light = theme.header_light
    name_style = styles["CvName"] if header_light else styles["CvNameWhite"]
    title_style = styles["CvTitle"] if header_light else styles["CvTitleWhite"]
    contact_style = styles["CvContact"] if header_light else styles["CvContactWhite"]

    photo = RLImage(str(PHOTO), width=31 * mm, height=47 * mm)
    text = [
        Paragraph(header["name"].upper(), name_style),
        Paragraph(header["title"], title_style),
        Paragraph(
            f'{header["location"]}<br/>'
            f'Mobile: {header["mobile"]}<br/>'
            f'Email: {header["email"]}<br/>'
            f'Website: {header["website"]}',
            contact_style,
        ),
    ]
    width = A4[0] - 28 * mm
    table = Table([[text, photo]], colWidths=[width - 38 * mm, 38 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), HexColor(theme.header_bg)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                ("BOX", (0, 0), (-1, -1), 0.5, HexColor(theme.accent)),
            ]
        )
    )
    return table


def build_pdf(
    theme: Theme,
    header: dict[str, str],
    blocks: list[tuple[str, str]],
    output_dir: Path | None = None,
):
    out_dir = output_dir or OUTPUT
    output = out_dir / f"{APPLICANT}_{theme.slug}_pdf.pdf"
    styles = pdf_styles(theme)
    primary = HexColor(theme.primary)
    soft = HexColor(theme.soft)
    accent = HexColor(theme.accent)

    doc = SimpleDocTemplate(
        str(output),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=11 * mm,
        bottomMargin=14 * mm,
        title=f'{header["name"]} - {theme.label} CV',
        author=header["name"],
    )
    story = [
        pdf_header(header, theme, styles),
        Spacer(1, 6),
        HRFlowable(width="100%", thickness=1.5, color=primary, spaceAfter=4),
    ]

    in_experience = False
    pending_section: list[Flowable] | None = None

    def append_content(flowable):
        nonlocal pending_section
        if pending_section:
            story.append(KeepTogether([*pending_section, flowable]))
            pending_section = None
        else:
            story.append(flowable)

    for kind, text in blocks:
        if kind == "separator":
            if pending_section is None:
                story.append(Spacer(1, 2))
        elif kind == "section":
            in_experience = text.upper() == "PROFESSIONAL EXPERIENCE"
            pending_section = [Spacer(1, 4), SectionBand(text, theme), Spacer(1, 2)]
        elif kind == "role":
            if theme.layout == "timeline" and in_experience:
                role = Table(
                    [[Paragraph(text, styles["CvRole"])]],
                    colWidths=[doc.width],
                )
                role.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, -1), soft),
                            ("LINEBEFORE", (0, 0), (0, 0), 3, primary),
                            ("LEFTPADDING", (0, 0), (-1, -1), 9),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                            ("TOPPADDING", (0, 0), (-1, -1), 2),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                        ]
                    )
                )
                append_content(role)
            else:
                append_content(Paragraph(text, styles["CvRole"]))
        elif kind == "subheading":
            append_content(Paragraph(text, styles["CvSub"]))
        elif kind == "bullet":
            append_content(Paragraph(f"- {text}", styles["CvBullet"]))
        elif kind == "paragraph":
            # Degree/institution lines are naturally compact; narrative stays justified.
            append_content(Paragraph(text, styles["CvBody"]))

    story.extend(signature_flowables(theme, header["name"].title()))

    def canvas_maker(*args, **kwargs):
        return NumberedCanvas(
            *args,
            name_label=f'{header["name"]} | {theme.label}',
            primary=theme.primary,
            **kwargs,
        )

    doc.build(story, canvasmaker=canvas_maker)
    ensure_signature_on_final_content_page(output)
    return output


def ensure_signature_on_final_content_page(pdf_path: Path, max_pages: int | None = None):
    """Drop a trailing blank signature-only page; never delete real content pages."""
    import fitz

    doc = fitz.open(str(pdf_path))
    # Remove trailing nearly-empty pages that only carry the signature spillover.
    while doc.page_count > 1:
        last = doc[-1]
        text = last.get_text().strip()
        content_lines = [
            line
            for line in text.splitlines()
            if line.strip()
            and not line.startswith("Page ")
            and "NABUKEERA" not in line
            and line.strip() not in {"Signed", "Racheal Nabukeera"}
        ]
        if content_lines or not last.get_images():
            break
        # Signature-only leftover page: stamp onto previous page if needed, then drop.
        prev = doc[-2]
        if not prev.get_images():
            _stamp_signature(prev)
        doc.delete_page(doc.page_count - 1)

    if max_pages is not None and doc.page_count > max_pages:
        # Soft preference only when the overflow page is empty-ish.
        pass

    last = doc[-1]
    if not last.get_images():
        _stamp_signature(last)

    tmp = pdf_path.with_suffix(".tmp.pdf")
    doc.save(str(tmp))
    doc.close()
    tmp.replace(pdf_path)


def _stamp_signature(page) -> None:
    """Stamp a compact signing block at the bottom-left of a PDF page."""
    import fitz

    rect = page.rect
    sig_w = 95
    with PILImage.open(SIGNATURE) as img:
        sig_h = sig_w * (img.size[1] / max(img.size[0], 1))
    x1 = 48
    y1 = rect.height - 48 - sig_h - 18
    page.insert_image(
        fitz.Rect(x1, y1, x1 + sig_w, y1 + sig_h),
        filename=str(SIGNATURE),
        keep_proportion=True,
    )
    page.insert_text(
        (x1, y1 - 8),
        "Signed",
        fontsize=7.5,
        color=(0.38, 0.44, 0.50),
        fontname="helv",
    )
    page.insert_text(
        (x1, y1 + sig_h + 12),
        "Racheal Nabukeera",
        fontsize=8,
        color=(0.04, 0.23, 0.32),
        fontname="helv",
    )


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def shade_cell(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color.lstrip("#"))
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_run(run, font: str, size: float, color: RGBColor, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def build_docx(
    theme: Theme,
    header: dict[str, str],
    blocks: list[tuple[str, str]],
    output_dir: Path | None = None,
):
    out_dir = output_dir or OUTPUT
    output = out_dir / f"{APPLICANT}_{theme.slug}_word.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    primary = rgb(theme.primary)
    secondary = rgb(theme.secondary)
    ink = RGBColor(0x25, 0x31, 0x3D)
    muted = RGBColor(0x60, 0x70, 0x80)
    header_color = primary if theme.header_light else RGBColor(255, 255, 255)
    header_secondary = secondary if theme.header_light else RGBColor(0xDC, 0xEA, 0xF0)

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    shade_cell(left, theme.header_bg)
    shade_cell(right, theme.header_bg)
    set_cell_margins(left, 140, 140, 140, 140)
    set_cell_margins(right, 140, 70, 140, 100)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(header["name"].upper()), theme.heading_font.replace("Helvetica-", "Arial "), 18, header_color, True)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(header["title"]), "Arial", 9.5, header_secondary)
    for line in (
        header["location"],
        f'Mobile: {header["mobile"]}',
        f'Email: {header["email"]}',
        f'Website: {header["website"]}',
    ):
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(line), "Arial", 8.5, header_color if not theme.header_light else muted)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(str(PHOTO), width=Inches(1.2))

    for style_name, font, size, color, bold in (
        ("Normal", "Arial" if theme.body_font == "Helvetica" else "Times New Roman", 9.5, ink, False),
        ("Title", "Arial", 18, primary, True),
        ("Heading 1", "Arial", 12, primary, True),
        ("Heading 2", "Arial", 11, primary, True),
        ("Heading 3", "Arial", 9.5, secondary, True),
    ):
        style = doc.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold

    body_font = "Times New Roman" if theme.body_font == "Times-Roman" else "Arial"
    for kind, text in blocks:
        if kind == "separator":
            continue
        if kind == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(5)
            if theme.layout in {"modern", "editorial"}:
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), theme.primary.lstrip("#"))
                p._p.get_or_add_pPr().append(shade)
                set_run(p.add_run(text.upper()), "Arial", 10.5, RGBColor(255, 255, 255), True)
            else:
                set_run(p.add_run(text.upper()), "Arial", 11.5, primary, True)
                p_bdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "12")
                bottom.set(qn("w:color"), theme.accent.lstrip("#"))
                p_bdr.append(bottom)
                p._p.get_or_add_pPr().append(p_bdr)
        elif kind == "role":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(text), "Arial", 11, primary, True)
        elif kind == "subheading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(3)
            set_run(p.add_run(text), "Arial", 9.5, secondary, True)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(text), body_font, 9.3, ink)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.08
            set_run(p.add_run(text), body_font, 9.5, ink)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run("Signed"), "Arial", 8, muted)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(SIGNATURE), width=Inches(1.15))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(header["name"].title()), "Arial", 9, primary, True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(f'{header["name"]} | {theme.label} | Page '), "Arial", 8, muted)
    add_field(p, " PAGE ")
    set_run(p.add_run(" of "), "Arial", 8, muted)
    add_field(p, " NUMPAGES ")

    doc.save(str(output))
    return output


def render_previews(pdf_paths: list[Path]):
    import fitz

    PREVIEW.mkdir(parents=True, exist_ok=True)
    for old in PREVIEW.glob("*.png"):
        old.unlink()
    for pdf_path in pdf_paths:
        pdf = fitz.open(str(pdf_path))
        for index, page in enumerate(pdf, 1):
            pix = page.get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False)
            pix.save(str(PREVIEW / f"{pdf_path.stem}_page_{index}.png"))
        print(f"{pdf_path.name}: {pdf.page_count} pages")
        pdf.close()


def verify_source_coverage(files: list[Path], blocks: list[tuple[str, str]]):
    """Check high-value source phrases exist in every generated deliverable."""
    import fitz

    required = [
        "2,000",
        "50+",
        "FairBanks Social Enterprise Initiative",
        "Machine Learning-Based Prediction",
        "Federation of Uganda Employers",
        "Human Resource Assistant",
        "Health Management Systems Administrator",
        "Firminus Mugumya",
        "firminus.mugumya@mak.ac.ug",
    ]
    for path in files:
        if path.suffix == ".pdf":
            document = fitz.open(str(path))
            text = "\n".join(page.get_text() for page in document)
            document.close()
        else:
            document = Document(str(path))
            text = "\n".join(p.text for p in document.paragraphs)
            text += "\n" + "\n".join(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
        missing = [phrase for phrase in required if phrase not in text]
        if missing:
            raise RuntimeError(f"{path.name} missing source content: {missing}")


def main():
    if not SOURCE.exists():
        raise SystemExit(f"Missing source of truth: {SOURCE}")
    if not PHOTO.exists():
        raise SystemExit(f"Missing portrait: {PHOTO}")
    if not SIGNATURE.exists():
        raise SystemExit(f"Missing signature: {SIGNATURE}")
    import shutil

    OUTPUT.mkdir(parents=True, exist_ok=True)
    if BUILD_TMP.exists():
        shutil.rmtree(BUILD_TMP)
    BUILD_TMP.mkdir(parents=True, exist_ok=True)

    header, blocks = parse_source()
    blocks = languages_above_summary(blocks)
    built: list[Path] = []
    for theme in THEMES:
        pdf = build_pdf(theme, header, blocks, BUILD_TMP)
        docx = build_docx(theme, header, blocks, BUILD_TMP)
        built.extend([pdf, docx])

    import time

    generated: list[Path] = []
    pending = list(built)
    for attempt in range(6):
        still_pending: list[Path] = []
        for src in pending:
            dest = OUTPUT / src.name
            try:
                shutil.copy2(src, dest)
                if dest not in generated:
                    generated.append(dest)
            except PermissionError:
                still_pending.append(src)
        pending = still_pending
        if not pending:
            break
        time.sleep(1.0)
    locked = [src.name for src in pending]
    for name in locked:
        print(f"Locked (close viewer and rebuild): {name}")

    keep = {path.name for path in built}
    for old in OUTPUT.iterdir():
        if old.is_file() and old.name not in keep:
            try:
                old.unlink()
            except PermissionError:
                print(f"Skipped locked obsolete file: {old.name}")

    # Always verify/render from the freshly built temp copies.
    pdf_paths = [path for path in built if path.suffix == ".pdf"]
    verify_source_coverage(built, blocks)
    render_previews(pdf_paths)

    ready = ROOT / "detailed_versions_ready"
    if ready.exists():
        shutil.rmtree(ready)
    if locked:
        ready.mkdir(parents=True, exist_ok=True)
        for src in built:
            shutil.copy2(src, ready / src.name)
        print(
            f"Word files updated in {OUTPUT}. "
            f"PDFs are locked by an open viewer; full set saved to {ready}. "
            "Close the PDF tabs, then rerun this script to refresh detailed_versions."
        )
        raise SystemExit(1)

    print(f"Generated {len(generated)} detailed CV files in {OUTPUT}")


if __name__ == "__main__":
    main()
