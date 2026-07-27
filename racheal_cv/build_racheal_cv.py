#!/usr/bin/env python3
"""Build Racheal Nabukeera professional CV (PDF + Word)."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Flowable,
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent
PHOTO_SRC = ROOT / "WhatsApp Image 2026-07-27 at 14.23.53.jpeg"
PHOTO_CIRCLE = ROOT / "tmp" / "racheal_photo_circle.png"
OUT_PDF = ROOT / "racheal_nabukeera_cv.pdf"
OUT_DOCX = ROOT / "racheal_nabukeera_cv.docx"

NAVY = HexColor("#0B3D5C")
TEAL = HexColor("#1A6B7A")
ACCENT = HexColor("#C9A227")
INK = HexColor("#1F2933")
MUTED = HexColor("#52606D")
RULE = HexColor("#D9E2EC")
SOFT = HexColor("#F5F7FA")

DOCX_NAVY = RGBColor(0x0B, 0x3D, 0x5C)
DOCX_TEAL = RGBColor(0x1A, 0x6B, 0x7A)
DOCX_INK = RGBColor(0x1F, 0x29, 0x33)
DOCX_MUTED = RGBColor(0x52, 0x60, 0x6D)


COMPETENCIES = {
    "Human Capital Leadership": [
        "Human Resource Strategy",
        "Human Capital Management",
        "Workforce Planning",
        "Talent Acquisition",
        "Talent Management",
        "Succession Planning",
    ],
    "Organisational Effectiveness": [
        "Organisational Development",
        "Change Management",
        "Business Transformation",
        "Leadership Development",
        "Employee Engagement",
        "Organisational Culture",
    ],
    "Human Resource Operations": [
        "Employee Relations",
        "Industrial Relations",
        "Labour Law Compliance",
        "Performance Management",
        "HR Policy Development",
        "Compensation & Benefits",
        "Staff Welfare",
    ],
    "Corporate Leadership": [
        "Executive Leadership",
        "Corporate Governance",
        "Strategic Planning",
        "Risk Management",
        "Stakeholder Engagement",
        "Corporate Administration",
    ],
    "Digital & Analytical Capability": [
        "Human Resource Information Systems (HRIS)",
        "Human Resource Analytics",
        "Management Reporting",
        "Data Analysis",
        "Digital Transformation",
        "AI in Human Resource Management",
    ],
}

EXPERIENCE = [
    {
        "title": "Founder, Managing Director & Executive Human Resource Leader",
        "org": "FairBanks Medical Centre Ltd",
        "dates": "2025 - Present",
        "summary": (
            "Provide executive leadership for a growing healthcare organisation, "
            "overseeing strategic planning, Human Resource management, governance, "
            "organisational development, institutional sustainability and business growth."
        ),
        "responsibilities": [
            "Provide strategic leadership and institutional direction.",
            "Lead organisational development and workforce planning.",
            "Develop Human Resource policies and organisational structures.",
            "Oversee performance management and staff development.",
            "Champion organisational culture and employee engagement.",
            "Direct corporate governance and compliance.",
            "Build strategic partnerships with government, insurers and development partners.",
            "Lead resource mobilisation and organisational sustainability initiatives.",
        ],
        "achievements": [
            "Founded FairBanks Medical Centre Ltd.",
            "Established the FairBanks Social Enterprise Initiative.",
            "Expanded specialist healthcare services.",
            "Led community health programmes benefiting over 10,000 people.",
            "Built strategic partnerships supporting institutional growth.",
            "Strengthened governance systems and organisational performance.",
        ],
    },
    {
        "title": "Group Human Resource & Administration Manager",
        "org": "Norvik Group",
        "dates": "July 2016 - February 2026",
        "summary": (
            "Provided strategic leadership for Human Resource and Administration functions "
            "across the Group, supporting executive management in driving organisational "
            "performance, workforce effectiveness and business transformation."
        ),
        "responsibilities": [
            "Developed and implemented Group Human Resource Strategy.",
            "Led workforce planning and organisational restructuring.",
            "Directed leadership development and succession planning initiatives.",
            "Managed organisational development and change management programmes.",
            "Oversaw recruitment and talent management.",
            "Strengthened employee engagement and staff retention.",
            "Developed Human Resource policies and governance frameworks.",
            "Ensured compliance with labour legislation and organisational policies.",
            "Supervised Human Resource teams across multiple operational units.",
            "Advised Executive Management on Human Capital Strategy.",
        ],
        "achievements": [
            "Provided strategic Human Resource leadership supporting more than 2,000 employees.",
            "Oversaw Human Resource operations across 50+ operational departments and business units.",
            "Led major organisational development and transformation initiatives.",
            "Improved Human Resource governance and organisational effectiveness.",
            "Strengthened performance management and leadership development systems.",
            "Enhanced workforce productivity through strategic Human Resource initiatives.",
        ],
    },
    {
        "title": "Human Resource & Administration Manager",
        "org": "Norvik Hospital Ltd",
        "dates": "September 2013 - 2016",
        "summary": (
            "Provided leadership for the Human Resource and Administration Department "
            "while strengthening Human Resource systems, employee engagement and "
            "organisational performance."
        ),
        "responsibilities": [
            "Managed recruitment and selection.",
            "Coordinated employee relations and disciplinary processes.",
            "Supervised performance appraisal systems.",
            "Developed and implemented Human Resource policies.",
            "Managed staff welfare initiatives.",
            "Oversaw administration and facilities management.",
            "Ensured compliance with labour legislation.",
        ],
        "achievements": [
            "Improved Human Resource operational efficiency.",
            "Strengthened organisational policy implementation.",
            "Enhanced staff performance management systems.",
            "Improved recruitment and retention practices.",
        ],
    },
    {
        "title": "Health Management Systems Administrator",
        "org": "St. Catherine's Hospital",
        "dates": "2010 - June 2013",
        "summary": (
            "Managed Hospital Information Systems while improving organisational reporting, "
            "digital information management and operational efficiency."
        ),
        "responsibilities": [
            "Managed Hospital Information Systems.",
            "Produced executive management reports.",
            "Improved administrative reporting processes.",
            "Coordinated user training.",
            "Strengthened information management.",
            "Supported digital transformation initiatives.",
        ],
        "achievements": [
            "Improved operational reporting.",
            "Enhanced information management systems.",
            "Increased administrative efficiency.",
        ],
    },
    {
        "title": "Human Resource Manager",
        "org": "St. Catherine's Hospital",
        "dates": "2007 - 2010",
        "summary": (
            "Managed Human Resource functions while supporting organisational growth through "
            "effective people management, policy implementation and employee development."
        ),
        "responsibilities": [
            "Led recruitment and selection.",
            "Managed employee relations.",
            "Coordinated performance management.",
            "Developed and implemented Human Resource policies.",
            "Managed employee welfare programmes.",
            "Ensured labour law compliance.",
            "Supervised Human Resource administration.",
        ],
        "achievements": [
            "Improved Human Resource systems and processes.",
            "Strengthened employee engagement.",
            "Enhanced recruitment and staff retention.",
            "Improved Human Resource service delivery.",
        ],
    },
    {
        "title": "Human Resource Assistant",
        "org": "E Power Limited",
        "dates": "1997 - 2000",
        "summary": (
            "Provided administrative and operational support to the Human Resource Department "
            "while contributing to efficient Human Resource service delivery."
        ),
        "responsibilities": [
            "Assisted with recruitment and onboarding.",
            "Maintained personnel records and employee files.",
            "Supported payroll administration.",
            "Coordinated leave administration.",
            "Assisted in organising staff training.",
            "Supported employee welfare initiatives.",
            "Prepared Human Resource reports and documentation.",
        ],
        "achievements": [
            "Maintained accurate Human Resource records.",
            "Supported efficient recruitment processes.",
            "Improved Human Resource administration.",
        ],
    },
]

HIGHLIGHTS = [
    "Nearly <b>30 years</b> of progressive Human Resource and Executive Management experience.",
    "More than <b>18 years</b> in Executive Human Resource Leadership.",
    "Strategic Human Resource leadership supporting <b>2,000+ employees</b>.",
    "Oversight of Human Resource functions across <b>50+ operational departments and business units</b>.",
    "Extensive experience in Human Capital Strategy, Organisational Development and Business Transformation.",
    "Successfully led organisation-wide workforce planning, recruitment, employee relations and performance management initiatives.",
    "Executive advisor to senior leadership on Human Capital Strategy, organisational effectiveness and change management.",
    "Founder and Managing Director of FairBanks Medical Centre Ltd.",
    "Successfully established the FairBanks Social Enterprise Initiative, reaching over <b>10,000 community beneficiaries</b>.",
    "Member of the <b>Federation of Uganda Employers (FUE)</b>.",
    "PhD Candidate in Management specialising in Human Resource Analytics, Artificial Intelligence and Organisational Behaviour.",
]

TECHNICAL = [
    "Microsoft Office Suite",
    "Human Resource Information Systems (HRIS)",
    "Enterprise Resource Planning (ERP) Systems",
    "Human Resource Analytics",
    "Performance Management Systems",
    "Payroll Administration Systems",
    "Data Analysis & Executive Reporting",
    "Artificial Intelligence Applications in Human Resource Management",
]

RESEARCH_INTERESTS = [
    "Human Resource Analytics",
    "Artificial Intelligence",
    "Machine Learning",
    "Organisational Behaviour",
    "Occupational Burnout",
    "Leadership",
    "Employee Wellbeing",
    "Healthcare Management",
    "Organisational Development",
]


def prepare_photo(size: int = 520) -> Path:
    PHOTO_CIRCLE.parent.mkdir(parents=True, exist_ok=True)
    img = Image.open(PHOTO_SRC).convert("RGBA")
    w, h = img.size
    side = min(w, h)
    left = (w - side) // 2
    top = max(0, (h - side) // 2 - int(side * 0.08))
    img = img.crop((left, top, left + side, top + side)).resize((size, size), Image.Resampling.LANCZOS)

    mask = Image.new("L", (size, size), 0)
    draw = ImageDraw.Draw(mask)
    draw.ellipse((0, 0, size - 1, size - 1), fill=255)

    output = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    output.paste(img, (0, 0))
    output.putalpha(mask)

    # Soft ring
    ring = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    ring_draw = ImageDraw.Draw(ring)
    ring_draw.ellipse((2, 2, size - 3, size - 3), outline=(11, 61, 92, 255), width=8)
    output = Image.alpha_composite(output, ring)
    output.save(PHOTO_CIRCLE)
    return PHOTO_CIRCLE


class SectionHeader(Flowable):
    def __init__(self, text: str, width: float):
        super().__init__()
        self.text = text.upper()
        self._width = width
        self.height = 16

    def wrap(self, availWidth, availHeight):
        self._width = availWidth
        return availWidth, self.height

    def draw(self):
        self.canv.setFillColor(NAVY)
        self.canv.setFont("Helvetica-Bold", 10.5)
        self.canv.drawString(0, 5, self.text)
        self.canv.setStrokeColor(ACCENT)
        self.canv.setLineWidth(2)
        text_w = self.canv.stringWidth(self.text, "Helvetica-Bold", 10.5)
        self.canv.line(0, 1.5, min(self._width, text_w + 28), 1.5)


def make_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Name",
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=21,
            textColor=NAVY,
            alignment=TA_LEFT,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Tagline",
            fontName="Helvetica",
            fontSize=8.8,
            leading=11.5,
            textColor=TEAL,
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Contact",
            fontName="Helvetica",
            fontSize=8.2,
            leading=11,
            textColor=MUTED,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyJust",
            fontName="Helvetica",
            fontSize=8.6,
            leading=11.4,
            textColor=INK,
            alignment=TA_JUSTIFY,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CvBullet",
            fontName="Helvetica",
            fontSize=8.4,
            leading=11,
            textColor=INK,
            leftIndent=10,
            bulletIndent=0,
            spaceAfter=1.2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="JobTitle",
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=12,
            textColor=NAVY,
            spaceBefore=2,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="OrgLine",
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=11,
            textColor=TEAL,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubHead",
            fontName="Helvetica-Bold",
            fontSize=8.2,
            leading=10.5,
            textColor=NAVY,
            spaceBefore=3,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CompHead",
            fontName="Helvetica-Bold",
            fontSize=7.8,
            leading=10,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CompItem",
            fontName="Helvetica",
            fontSize=7.4,
            leading=9.6,
            textColor=INK,
            alignment=TA_LEFT,
            leftIndent=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EduTitle",
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=11,
            textColor=NAVY,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EduMeta",
            fontName="Helvetica",
            fontSize=8.3,
            leading=10.5,
            textColor=MUTED,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Footer",
            fontName="Helvetica",
            fontSize=7.5,
            textColor=MUTED,
            alignment=TA_CENTER,
        )
    )
    return styles


def bullet(text: str, style) -> Paragraph:
    return Paragraph(f"• {text}", style)


def build_pdf() -> None:
    styles = make_styles()
    photo = prepare_photo()

    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title="Racheal Nabukeera - Curriculum Vitae",
        author="Racheal Nabukeera",
    )
    width = A4[0] - 28 * mm
    story = []

    name_block = [
        Paragraph("RACHEAL NABUKEERA", styles["Name"]),
        Paragraph(
            "Senior Human Resource Executive | Human Capital Strategist<br/>"
            "Organisational Development &amp; Executive Leadership",
            styles["Tagline"],
        ),
        Paragraph(
            "Kampala, Uganda<br/>"
            "Mobile: +256 772 849258 | +256 701 849258<br/>"
            "Email: info@fairbanksmedicalcentre.org | nracheal017@gmail.com<br/>"
            "Website: www.fairbanksmedicalcentre.org",
            styles["Contact"],
        ),
    ]
    photo_img = RLImage(str(photo), width=28 * mm, height=28 * mm)
    header = Table(
        [[name_block, photo_img]],
        colWidths=[width - 34 * mm, 34 * mm],
    )
    header.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header)
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=2.2, color=NAVY, spaceBefore=0, spaceAfter=1))
    story.append(HRFlowable(width="100%", thickness=0.8, color=ACCENT, spaceBefore=0, spaceAfter=8))

    # Executive profile
    story.append(SectionHeader("Executive Profile", width))
    story.append(Spacer(1, 5))
    story.append(
        Paragraph(
            "Strategic Human Resource Executive with nearly <b>30 years of progressive professional "
            "experience</b> and over <b>18 years of executive Human Resource leadership</b>, driving "
            "organisational transformation, workforce excellence and sustainable business performance "
            "across complex organisations.",
            styles["BodyJust"],
        )
    )
    story.append(
        Paragraph(
            "Recognised for developing and implementing innovative Human Capital strategies that align "
            "people, performance and organisational objectives. Demonstrated expertise in Human Resource "
            "Strategy, Organisational Development, Talent Management, Workforce Planning, Leadership "
            "Development, Performance Management, Employee Relations, Change Management, Corporate "
            "Governance and Executive Administration.",
            styles["BodyJust"],
        )
    )
    story.append(
        Paragraph(
            "Successfully provided strategic Human Resource leadership supporting <b>more than 2,000 "
            "employees</b> across <b>over 50 operational departments and business units</b>, partnering "
            "with executive leadership to strengthen organisational capability, improve workforce "
            "productivity and deliver measurable business results.",
            styles["BodyJust"],
        )
    )
    story.append(
        Paragraph(
            "Currently serving as Founder and Managing Director of FairBanks Medical Centre Ltd, "
            "providing executive leadership in organisational strategy, corporate governance, "
            "institutional development, stakeholder engagement and business sustainability.",
            styles["BodyJust"],
        )
    )
    story.append(
        Paragraph(
            "A PhD Candidate in Management at Uganda Christian University with research focusing on "
            "the application of Artificial Intelligence and Machine Learning in predicting occupational "
            "burnout among medical practitioners, reflecting a passion for innovation, evidence-based "
            "leadership and organisational excellence.",
            styles["BodyJust"],
        )
    )
    story.append(Spacer(1, 4))

    # Highlights
    story.append(SectionHeader("Key Career Highlights", width))
    story.append(Spacer(1, 5))
    for item in HIGHLIGHTS:
        story.append(bullet(item, styles["CvBullet"]))
    story.append(Spacer(1, 5))

    # Competencies
    story.append(SectionHeader("Executive Core Competencies", width))
    story.append(Spacer(1, 5))
    groups = list(COMPETENCIES.items())
    # First row: 3 cols, second row: 2 cols spanning nicely
    def comp_cell(title: str, items: list[str]):
        lines = [Paragraph(title, styles["CompHead"])]
        for it in items:
            lines.append(Paragraph(f"• {it}", styles["CompItem"]))
        return lines

    row1 = [comp_cell(*groups[i]) for i in range(3)]
    t1 = Table([row1], colWidths=[width / 3.0] * 3)
    t1.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), SOFT),
                ("BOX", (0, 0), (-1, -1), 0.4, RULE),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, RULE),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(t1)
    story.append(Spacer(1, 3))
    row2 = [comp_cell(*groups[3]), comp_cell(*groups[4])]
    t2 = Table([row2], colWidths=[width / 2.0] * 2)
    t2.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), SOFT),
                ("BOX", (0, 0), (-1, -1), 0.4, RULE),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, RULE),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(t2)
    story.append(Spacer(1, 7))

    # Experience
    for idx, role in enumerate(EXPERIENCE):
        block = []
        if idx == 0:
            block.extend(
                [
                    SectionHeader("Professional Experience", width),
                    Spacer(1, 5),
                ]
            )
        block.extend(
            [
                Paragraph(role["title"], styles["JobTitle"]),
                Paragraph(f"{role['org']}  |  {role['dates']}", styles["OrgLine"]),
                Paragraph(role["summary"], styles["BodyJust"]),
                Paragraph("Major Responsibilities", styles["SubHead"]),
            ]
        )
        for item in role["responsibilities"]:
            block.append(bullet(item, styles["CvBullet"]))
        block.append(Paragraph("Key Achievements", styles["SubHead"]))
        for item in role["achievements"]:
            block.append(bullet(item, styles["CvBullet"]))
        block.append(Spacer(1, 5))
        story.append(KeepTogether(block))

    half = (len(TECHNICAL) + 1) // 2
    left = [bullet(x, styles["CvBullet"]) for x in TECHNICAL[:half]]
    right = [bullet(x, styles["CvBullet"]) for x in TECHNICAL[half:]]
    tech = Table([[left, right]], colWidths=[width / 2.0] * 2)
    tech.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    half_r = (len(RESEARCH_INTERESTS) + 1) // 2
    left_r = [bullet(x, styles["CvBullet"]) for x in RESEARCH_INTERESTS[:half_r]]
    right_r = [bullet(x, styles["CvBullet"]) for x in RESEARCH_INTERESTS[half_r:]]
    res = Table([[left_r, right_r]], colWidths=[width / 2.0] * 2)
    res.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    story.append(
        KeepTogether(
            [
                SectionHeader("Education", width),
                Spacer(1, 5),
                Paragraph(
                    "Doctor of Philosophy (PhD) in Management (Ongoing)",
                    styles["EduTitle"],
                ),
                Paragraph("Uganda Christian University", styles["EduMeta"]),
                Paragraph(
                    "Research Topic: <b>Machine Learning-Based Prediction of Occupational "
                    "Burnout Among Medical Practitioners in Uganda</b>",
                    styles["BodyJust"],
                ),
                Paragraph(
                    "Master of Social Sector Planning and Management",
                    styles["EduTitle"],
                ),
                Paragraph("Makerere University", styles["EduMeta"]),
                Paragraph("Bachelor of Arts in Social Sciences", styles["EduTitle"]),
                Paragraph("Makerere University", styles["EduMeta"]),
                Spacer(1, 4),
                SectionHeader("Professional Membership", width),
                Spacer(1, 5),
                bullet(
                    "<b>Member, Federation of Uganda Employers (FUE)</b>",
                    styles["CvBullet"],
                ),
            ]
        )
    )
    story.append(Spacer(1, 4))
    story.append(
        KeepTogether(
            [
                SectionHeader("Technical Competencies", width),
                Spacer(1, 5),
                tech,
                Spacer(1, 5),
                SectionHeader("Academic Research", width),
                Spacer(1, 5),
                Paragraph("Doctoral Research", styles["SubHead"]),
                Paragraph(
                    "<b>Machine Learning-Based Prediction of Occupational Burnout Among Medical "
                    "Practitioners in Uganda</b>",
                    styles["BodyJust"],
                ),
                Paragraph("Research Interests", styles["SubHead"]),
                res,
                Spacer(1, 5),
                SectionHeader("Languages", width),
                Spacer(1, 5),
                bullet("English - Fluent", styles["CvBullet"]),
                bullet("Luganda - Fluent", styles["CvBullet"]),
            ]
        )
    )

    def on_page(canvas, doc_):
        canvas.saveState()
        canvas.setStrokeColor(RULE)
        canvas.setLineWidth(0.5)
        canvas.line(14 * mm, 8 * mm, A4[0] - 14 * mm, 8 * mm)
        canvas.setFont("Helvetica", 7.2)
        canvas.setFillColor(MUTED)
        canvas.drawString(14 * mm, 4.5 * mm, "Racheal Nabukeera | Curriculum Vitae")
        canvas.drawRightString(A4[0] - 14 * mm, 4.5 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def set_run_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B3D5C")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=DOCX_NAVY)
    # gold underline via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc: Document, text: str, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=DOCX_INK)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=9.2, bold=True, color=DOCX_INK)
        r2 = p.add_run(text)
        set_run_font(r2, size=9.2, color=DOCX_INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.2, color=DOCX_INK)


def set_cell_shading(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def build_docx() -> None:
    photo = prepare_photo()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # Header table: text + photo
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    left.width = Cm(14.5)
    right.width = Cm(3.8)

    p = left.paragraphs[0]
    r = p.add_run("RACHEAL NABUKEERA")
    set_run_font(r, size=18, bold=True, color=DOCX_NAVY)

    p2 = left.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(
        "Senior Human Resource Executive | Human Capital Strategist | "
        "Organisational Development & Executive Leadership"
    )
    set_run_font(r2, size=9, color=DOCX_TEAL)

    for line in [
        "Kampala, Uganda",
        "Mobile: +256 772 849258 | +256 701 849258",
        "Email: info@fairbanksmedicalcentre.org | nracheal017@gmail.com",
        "Website: www.fairbanksmedicalcentre.org",
    ]:
        cp = left.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run(line)
        set_run_font(cr, size=9, color=DOCX_MUTED)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run().add_picture(str(photo), width=Inches(1.15))

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(4)
    line.paragraph_format.space_after = Pt(2)
    add_horizontal_line(line)

    # Profile
    add_section_heading(doc, "Executive Profile")
    add_body(
        doc,
        "Strategic Human Resource Executive with nearly 30 years of progressive professional "
        "experience and over 18 years of executive Human Resource leadership, driving organisational "
        "transformation, workforce excellence and sustainable business performance across complex "
        "organisations.",
    )
    add_body(
        doc,
        "Recognised for developing and implementing innovative Human Capital strategies that align "
        "people, performance and organisational objectives. Demonstrated expertise in Human Resource "
        "Strategy, Organisational Development, Talent Management, Workforce Planning, Leadership "
        "Development, Performance Management, Employee Relations, Change Management, Corporate "
        "Governance and Executive Administration.",
    )
    add_body(
        doc,
        "Successfully provided strategic Human Resource leadership supporting more than 2,000 "
        "employees across over 50 operational departments and business units, partnering with "
        "executive leadership to strengthen organisational capability, improve workforce productivity "
        "and deliver measurable business results.",
    )
    add_body(
        doc,
        "Currently serving as Founder and Managing Director of FairBanks Medical Centre Ltd, "
        "providing executive leadership in organisational strategy, corporate governance, "
        "institutional development, stakeholder engagement and business sustainability.",
    )
    add_body(
        doc,
        "A PhD Candidate in Management at Uganda Christian University with research focusing on "
        "the application of Artificial Intelligence and Machine Learning in predicting occupational "
        "burnout among medical practitioners, reflecting a passion for innovation, evidence-based "
        "leadership and organisational excellence.",
    )

    add_section_heading(doc, "Key Career Highlights")
    plain_highlights = [
        "Nearly 30 years of progressive Human Resource and Executive Management experience.",
        "More than 18 years in Executive Human Resource Leadership.",
        "Strategic Human Resource leadership supporting 2,000+ employees.",
        "Oversight of Human Resource functions across 50+ operational departments and business units.",
        "Extensive experience in Human Capital Strategy, Organisational Development and Business Transformation.",
        "Successfully led organisation-wide workforce planning, recruitment, employee relations and performance management initiatives.",
        "Executive advisor to senior leadership on Human Capital Strategy, organisational effectiveness and change management.",
        "Founder and Managing Director of FairBanks Medical Centre Ltd.",
        "Successfully established the FairBanks Social Enterprise Initiative, reaching over 10,000 community beneficiaries.",
        "Member of the Federation of Uganda Employers (FUE).",
        "PhD Candidate in Management specialising in Human Resource Analytics, Artificial Intelligence and Organisational Behaviour.",
    ]
    for item in plain_highlights:
        add_bullet(doc, item)

    add_section_heading(doc, "Executive Core Competencies")
    # 3-column competency table for first three, then 2-col
    groups = list(COMPETENCIES.items())
    comp = doc.add_table(rows=1, cols=3)
    for i in range(3):
        cell = comp.rows[0].cells[i]
        set_cell_shading(cell, "F5F7FA")
        title, items = groups[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=8.5, bold=True, color=DOCX_NAVY)
        for item in items:
            ip = cell.add_paragraph()
            ip.paragraph_format.space_before = Pt(0)
            ip.paragraph_format.space_after = Pt(0)
            ir = ip.add_run(f"• {item}")
            set_run_font(ir, size=8, color=DOCX_INK)

    comp2 = doc.add_table(rows=1, cols=2)
    for i, idx in enumerate((3, 4)):
        cell = comp2.rows[0].cells[i]
        set_cell_shading(cell, "F5F7FA")
        title, items = groups[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=8.5, bold=True, color=DOCX_NAVY)
        for item in items:
            ip = cell.add_paragraph()
            ip.paragraph_format.space_before = Pt(0)
            ip.paragraph_format.space_after = Pt(0)
            ir = ip.add_run(f"• {item}")
            set_run_font(ir, size=8, color=DOCX_INK)

    add_section_heading(doc, "Professional Experience")
    for role in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(role["title"])
        set_run_font(r, size=10.5, bold=True, color=DOCX_NAVY)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{role['org']}  |  {role['dates']}")
        set_run_font(r, size=9.5, bold=True, color=DOCX_TEAL)

        add_body(doc, role["summary"])

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("Major Responsibilities")
        set_run_font(r, size=9.5, bold=True, color=DOCX_NAVY)
        for item in role["responsibilities"]:
            add_bullet(doc, item)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("Key Achievements")
        set_run_font(r, size=9.5, bold=True, color=DOCX_NAVY)
        for item in role["achievements"]:
            add_bullet(doc, item)

    add_section_heading(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Doctor of Philosophy (PhD) in Management (Ongoing)")
    set_run_font(r, size=10, bold=True, color=DOCX_NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Uganda Christian University")
    set_run_font(r, size=9.2, color=DOCX_MUTED)
    add_body(
        doc,
        "Research Topic: Machine Learning-Based Prediction of Occupational Burnout Among "
        "Medical Practitioners in Uganda",
    )

    p = doc.add_paragraph()
    r = p.add_run("Master of Social Sector Planning and Management")
    set_run_font(r, size=10, bold=True, color=DOCX_NAVY)
    p = doc.add_paragraph()
    r = p.add_run("Makerere University")
    set_run_font(r, size=9.2, color=DOCX_MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Bachelor of Arts in Social Sciences")
    set_run_font(r, size=10, bold=True, color=DOCX_NAVY)
    p = doc.add_paragraph()
    r = p.add_run("Makerere University")
    set_run_font(r, size=9.2, color=DOCX_MUTED)

    add_section_heading(doc, "Professional Membership")
    add_bullet(doc, "Member, Federation of Uganda Employers (FUE)")

    add_section_heading(doc, "Technical Competencies")
    for item in TECHNICAL:
        add_bullet(doc, item)

    add_section_heading(doc, "Academic Research")
    p = doc.add_paragraph()
    r = p.add_run("Doctoral Research")
    set_run_font(r, size=9.5, bold=True, color=DOCX_NAVY)
    add_body(
        doc,
        "Machine Learning-Based Prediction of Occupational Burnout Among Medical Practitioners in Uganda",
    )
    p = doc.add_paragraph()
    r = p.add_run("Research Interests")
    set_run_font(r, size=9.5, bold=True, color=DOCX_NAVY)
    for item in RESEARCH_INTERESTS:
        add_bullet(doc, item)

    add_section_heading(doc, "Languages")
    add_bullet(doc, "English - Fluent")
    add_bullet(doc, "Luganda - Fluent")

    # Footer with page numbers
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Racheal Nabukeera | Curriculum Vitae")
    set_run_font(run, size=8, color=DOCX_MUTED)

    doc.save(str(OUT_DOCX))


def render_preview() -> list[Path]:
    import fitz

    out_dir = ROOT / "tmp" / "preview"
    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("*.png"):
        old.unlink()
    pdf = fitz.open(str(OUT_PDF))
    paths = []
    for i, page in enumerate(pdf, start=1):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        path = out_dir / f"cv_page_{i}.png"
        pix.save(str(path))
        paths.append(path)
    pdf.close()
    return paths


def main():
    if not PHOTO_SRC.exists():
        raise SystemExit(f"Photo not found: {PHOTO_SRC}")
    build_pdf()
    build_docx()
    pages = render_preview()
    print(f"PDF:  {OUT_PDF}")
    print(f"DOCX: {OUT_DOCX}")
    print(f"Pages previewed: {len(pages)}")
    for p in pages:
        print(f"  {p}")


if __name__ == "__main__":
    main()
