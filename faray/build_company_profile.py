"""
Build Fravent Limited company profile PDF.
Brand colors: Deep Blue #071B36, Lemon Green #A3C324, White #FFFFFF
"""

from __future__ import annotations

from pathlib import Path

from PIL import Image as PILImage
from reportlab.lib.colors import Color, HexColor, white, black
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
PHOTOS = ROOT
OUT_PDF = ROOT / "profile-2026.pdf"
OUT_DOCX = ROOT / "profile-2026.docx"
OUT = OUT_PDF  # backward-compatible alias
TMP = ROOT.parent / "tmp" / "pdfs"
TMP.mkdir(parents=True, exist_ok=True)

PAGE_W, PAGE_H = A4
MARGIN = 12 * mm
HEADER_H = 16 * mm
FOOTER_H = 11 * mm
CAPTION_H = 7 * mm
CONTENT_TOP = PAGE_H - HEADER_H - 6 * mm
CONTENT_BOTTOM = FOOTER_H + 4 * mm

DEEP_BLUE = HexColor("#071B36")
DEEP_BLUE_MID = HexColor("#0C2E52")
LEMON = HexColor("#A3C324")
LEMON_DARK = HexColor("#8AAA1A")
LIGHT_GRAY = HexColor("#F4F6F8")
SOFT_BLUE = HexColor("#E8EEF4")
TEXT = HexColor("#1A2330")
MUTED = HexColor("#5A6570")

TAGLINE = "Your toolkit for real solutions"
SLOGAN = "Safe. Smart. Sustainable."


def p(name: str) -> Path:
    return PHOTOS / name


def a(name: str) -> Path:
    return ASSETS / name


def draw_logo_mark(c: canvas.Canvas, x: float, y: float, size: float = 12 * mm, light: bool = True):
    """Simple Fravent F + lemon arc mark."""
    f_color = white if light else DEEP_BLUE
    # green arc
    c.setStrokeColor(LEMON)
    c.setLineWidth(max(1.8, size * 0.08))
    c.arc(x, y, x + size, y + size, 40, 160)
    # stylized F
    c.setFillColor(f_color)
    stem_w = size * 0.18
    c.rect(x + size * 0.28, y + size * 0.18, stem_w, size * 0.62, fill=1, stroke=0)
    c.rect(x + size * 0.28, y + size * 0.68, size * 0.42, size * 0.12, fill=1, stroke=0)
    c.rect(x + size * 0.28, y + size * 0.42, size * 0.32, size * 0.10, fill=1, stroke=0)


def draw_header(c: canvas.Canvas, title: str, subtitle: str | None = None):
    c.setFillColor(DEEP_BLUE)
    c.rect(0, PAGE_H - HEADER_H, PAGE_W, HEADER_H, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.rect(0, PAGE_H - HEADER_H - 1.2 * mm, PAGE_W, 1.2 * mm, fill=1, stroke=0)

    draw_logo_mark(c, MARGIN, PAGE_H - HEADER_H + 2.5 * mm, size=10 * mm, light=True)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + 12 * mm, PAGE_H - 9 * mm, "FRAVENT LIMITED")
    c.setFont("Helvetica", 7)
    c.setFillColor(LEMON)
    c.drawString(MARGIN + 12 * mm, PAGE_H - 13 * mm, TAGLINE)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 12)
    c.drawRightString(PAGE_W - MARGIN, PAGE_H - 9 * mm, title)
    if subtitle:
        c.setFont("Helvetica", 7.5)
        c.setFillColor(LEMON)
        c.drawRightString(PAGE_W - MARGIN, PAGE_H - 13.5 * mm, subtitle)


def draw_footer(c: canvas.Canvas, page_no: int, total: int):
    y = 8 * mm
    c.setFillColor(DEEP_BLUE)
    c.rect(0, 0, PAGE_W, FOOTER_H, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.rect(0, FOOTER_H, PAGE_W, 1.0 * mm, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("Helvetica", 7)
    c.drawString(MARGIN, y - 0.5 * mm, "fraventlimited@gmail.com  |  +256 778 455 042  |  Kampala, Uganda")
    c.drawRightString(PAGE_W - MARGIN, y - 0.5 * mm, f"{page_no} / {total}")


def fit_image(path: Path, max_w: float, max_h: float) -> tuple[float, float]:
    with PILImage.open(path) as im:
        w, h = im.size
    scale = min(max_w / w, max_h / h)
    return w * scale, h * scale


def draw_image(c: canvas.Canvas, path: Path, x: float, y: float, w: float, h: float, cover: bool = True):
    if not path.exists():
        c.setFillColor(SOFT_BLUE)
        c.rect(x, y, w, h, fill=1, stroke=0)
        return
    if cover:
        with PILImage.open(path) as im:
            im = im.convert("RGB")
            src_w, src_h = im.size
            target_ratio = w / h
            src_ratio = src_w / src_h
            if src_ratio > target_ratio:
                new_w = int(src_h * target_ratio)
                left = (src_w - new_w) // 2
                im = im.crop((left, 0, left + new_w, src_h))
            else:
                new_h = int(src_w / target_ratio)
                top = (src_h - new_h) // 2
                im = im.crop((0, top, src_w, top + new_h))
            tmp = TMP / f"crop_{path.stem}_{int(w)}_{int(h)}.jpg"
            im.save(tmp, quality=90)
            c.drawImage(str(tmp), x, y, width=w, height=h, mask="auto")
    else:
        iw, ih = fit_image(path, w, h)
        c.drawImage(str(path), x + (w - iw) / 2, y + (h - ih) / 2, width=iw, height=ih, mask="auto", preserveAspectRatio=True)


def draw_caption(c: canvas.Canvas, text: str, x: float, y: float, w: float):
    """Caption bar under a photo. y is the bottom of the photo."""
    bar_h = CAPTION_H
    c.setFillColor(DEEP_BLUE)
    c.rect(x, y - bar_h, w, bar_h, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.rect(x, y - 1.0, w, 1.0, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("Helvetica", 6)
    lines = wrap_text(c, text, "Helvetica", 6, w - 3 * mm)
    ty = y - 3.6 * mm if len(lines) == 1 else y - 3.0 * mm
    for line in lines[:2]:
        c.drawString(x + 1.5 * mm, ty, line)
        ty -= 6.2


def draw_captioned_image(
    c: canvas.Canvas,
    path: Path,
    x: float,
    y: float,
    w: float,
    h: float,
    caption: str,
    cover: bool = True,
):
    """Draw image with caption bar. (x,y) is bottom-left of full block including caption."""
    bar_h = CAPTION_H
    img_h = max(h - bar_h, 8 * mm)
    draw_image(c, path, x, y + bar_h, w, img_h, cover=cover)
    draw_caption(c, caption, x, y + bar_h, w)


def draw_photo_grid(
    c: canvas.Canvas,
    items: list[tuple[Path, str]],
    x: float,
    y_top: float,
    y_bottom: float,
    cols: int = 3,
    gap: float | None = None,
):
    """Fill a vertical band with a captioned photo grid."""
    if not items or y_top <= y_bottom + 10 * mm:
        return
    gap = 2.5 * mm if gap is None else gap
    n = len(items)
    cols = min(cols, n)
    rows = (n + cols - 1) // cols
    usable_w = PAGE_W - 2 * MARGIN
    usable_h = y_top - y_bottom
    cw = (usable_w - (cols - 1) * gap) / cols
    ch = (usable_h - (rows - 1) * gap) / rows
    for i, (ph, cap) in enumerate(items):
        col = i % cols
        row = i // cols
        xx = x + col * (cw + gap)
        yy = y_top - (row + 1) * ch - row * gap
        draw_captioned_image(c, ph, xx, yy, cw, ch, cap)


def circle_clip_image(c: canvas.Canvas, path: Path, cx: float, cy: float, r: float):
    if not path.exists():
        c.setFillColor(SOFT_BLUE)
        c.circle(cx, cy, r, fill=1, stroke=0)
        return
    # Prefer a pre-squared cover crop saved for stable circular clipping
    with PILImage.open(path) as im:
        im = im.convert("RGB")
        src_w, src_h = im.size
        side = min(src_w, src_h)
        left = (src_w - side) // 2
        top = (src_h - side) // 2
        im = im.crop((left, top, left + side, top + side))
        tmp = TMP / f"circle_{path.stem}_{int(r)}.jpg"
        im.save(tmp, quality=92)
    c.saveState()
    pth = c.beginPath()
    pth.circle(cx, cy, r)
    c.clipPath(pth, stroke=0)
    c.drawImage(str(tmp), cx - r, cy - r, width=2 * r, height=2 * r, mask="auto")
    c.restoreState()
    c.setStrokeColor(LEMON)
    c.setLineWidth(2.2)
    c.circle(cx, cy, r, fill=0, stroke=1)


def wrap_text(c: canvas.Canvas, text: str, font: str, size: float, max_w: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    cur = ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if c.stringWidth(trial, font, size) <= max_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def draw_paragraph(c: canvas.Canvas, text: str, x: float, y: float, max_w: float, size: float = 9.5, leading: float = 13, color=TEXT, font="Helvetica") -> float:
    c.setFont(font, size)
    c.setFillColor(color)
    for line in wrap_text(c, text, font, size, max_w):
        c.drawString(x, y, line)
        y -= leading
    return y


def section_label(c: canvas.Canvas, text: str, x: float, y: float):
    c.setFillColor(LEMON)
    c.roundRect(x, y - 2, 3 * mm, 8, 1, fill=1, stroke=0)
    c.setFillColor(DEEP_BLUE)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(x + 5 * mm, y, text)


def meta_row(c: canvas.Canvas, label: str, value: str, x: float, y: float, label_w: float = 38 * mm, max_w: float = 120 * mm) -> float:
    c.setFont("Helvetica-Bold", 9)
    c.setFillColor(DEEP_BLUE)
    c.drawString(x, y, label)
    c.setFont("Helvetica", 9)
    c.setFillColor(TEXT)
    lines = wrap_text(c, value, "Helvetica", 9, max_w)
    for i, line in enumerate(lines):
        c.drawString(x + label_w, y - i * 12, line)
    return y - max(12, len(lines) * 12) - 4


# ---------------------------------------------------------------------------
# Pages
# ---------------------------------------------------------------------------

def page_cover(c: canvas.Canvas):
    # Full-bleed hero
    hero = p("nakivale-solar-hybrid-ground-array-installation.jpeg")
    draw_image(c, hero, 0, 0, PAGE_W, PAGE_H, cover=True)

    # Dark overlay panels
    c.setFillColor(Color(0.03, 0.11, 0.21, alpha=0.72))
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)

    # Left accent bar
    c.setFillColor(LEMON)
    c.rect(0, 0, 6 * mm, PAGE_H, fill=1, stroke=0)

    # Brand block
    draw_logo_mark(c, MARGIN + 4 * mm, PAGE_H - 42 * mm, size=18 * mm, light=True)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 28)
    c.drawString(MARGIN + 26 * mm, PAGE_H - 30 * mm, "FRAVENT LIMITED")
    c.setFillColor(LEMON)
    c.rect(MARGIN + 26 * mm, PAGE_H - 33 * mm, 42 * mm, 1.6, fill=1, stroke=0)
    c.setFont("Helvetica", 11)
    c.setFillColor(white)
    c.drawString(MARGIN + 26 * mm, PAGE_H - 40 * mm, "Electrical & Solar Solutions")

    c.setFont("Helvetica-Bold", 36)
    c.setFillColor(white)
    c.drawString(MARGIN + 4 * mm, PAGE_H / 2 + 20 * mm, "COMPANY")
    c.setFillColor(LEMON)
    c.drawString(MARGIN + 4 * mm, PAGE_H / 2 + 4 * mm, "PROFILE")

    c.setFillColor(white)
    c.setFont("Helvetica", 11)
    intro = "An introduction to who we are, what we do, and how we create lasting value for communities and clients across Uganda."
    y = PAGE_H / 2 - 10 * mm
    for line in wrap_text(c, intro, "Helvetica", 11, PAGE_W - 2 * MARGIN - 10 * mm):
        c.drawString(MARGIN + 4 * mm, y, line)
        y -= 15

    # Quote
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Oblique", 12)
    c.drawString(MARGIN + 4 * mm, 48 * mm, f'"{TAGLINE}"')

    c.setFillColor(white)
    c.setFont("Helvetica", 9)
    c.drawString(MARGIN + 4 * mm, 32 * mm, SLOGAN)
    c.drawString(MARGIN + 4 * mm, 24 * mm, "www.fraventlimited.com")
    c.drawString(MARGIN + 4 * mm, 16 * mm, "fraventlimited@gmail.com  ·  +256 778 455 042")


def page_about(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "About Us", "Our Story")
    y = CONTENT_TOP

    section_label(c, "Who We Are", MARGIN, y)
    y -= 7 * mm

    left_w = (PAGE_W - 2 * MARGIN) * 0.58
    right_w = (PAGE_W - 2 * MARGIN) * 0.40
    story = [
        "Fravent Limited is a fully registered company with the Registrar of Companies (Registration Number: 80020003614101). We help clients realise goals in electrical engineering, solar and renewable energy, property management, consultancy and advisory, import and export, logistics and transport, land brokerage, and business intelligence.",
        "We employ experienced, exposed, and professional staff for every project. Our clientele includes private individuals, schools, private and public companies, organisations, local and international investors, and government ministries.",
        "A steady flow of orders is an encouraging sign of trust in Fravent management's rich heritage and diversified experience, built on successful delivery year after year.",
    ]
    ty = y
    for para in story:
        ty = draw_paragraph(c, para, MARGIN, ty, left_w - 3 * mm, size=8.5, leading=11.5)
        ty -= 3.5

    photos_side = [
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop solar installation."),
        (a("cover_field_team_main.png"), "Field team delivering on site."),
        (a("cover_pole_climber.png"), "Power-line works at height."),
    ]
    side_h = y - ty + 4 * mm
    side_h = max(side_h, 72 * mm)
    draw_photo_grid(c, photos_side, MARGIN + left_w + 2 * mm, y + 2 * mm, y + 2 * mm - side_h, cols=1, gap=2 * mm)

    y = min(ty, y + 2 * mm - side_h) - 6 * mm

    boxes = [
        ("Registered", "Company No. 80020003614101"),
        ("Focus", "Electrical & Solar Solutions"),
        ("Base", "Kampala, Uganda"),
        ("Promise", "Quality, Affordable, On-time"),
        ("Clients", "Private, Public & NGO"),
        ("Delivery", "Design to Commissioning"),
    ]
    gap = 2.5 * mm
    bw = (PAGE_W - 2 * MARGIN - 5 * gap) / 6
    bh = 18 * mm
    for i, (t, s) in enumerate(boxes):
        x = MARGIN + i * (bw + gap)
        c.setFillColor(DEEP_BLUE if i % 2 == 0 else LEMON)
        c.roundRect(x, y - bh, bw, bh, 2, fill=1, stroke=0)
        c.setFillColor(white if i % 2 == 0 else DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7)
        c.drawCentredString(x + bw / 2, y - 6 * mm, t)
        c.setFont("Helvetica", 5.5)
        for j, line in enumerate(wrap_text(c, s, "Helvetica", 5.5, bw - 3 * mm)):
            c.drawCentredString(x + bw / 2, y - 11 * mm - j * 6.5, line)

    y = y - bh - 5 * mm
    more = [
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Ground array works — Nakivale."),
        (p("ugafode-nakivale-branch-site.jpeg"), "UGAFODE Nakivale branch site."),
        (p("kiwoko-hospital-adara-team-onsite.jpeg"), "Team on site — Kiwoko Hospital."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Commissioned hybrid inverter system."),
    ]
    draw_photo_grid(c, more, MARGIN, y, CONTENT_BOTTOM, cols=4)
    draw_footer(c, page_no, total)


def page_mission(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Purpose", "Mission · Vision · Values")
    y = CONTENT_TOP

    # Compact mission + vision side by side
    gap = 3 * mm
    cw = (PAGE_W - 2 * MARGIN - gap) / 2
    ch = 38 * mm
    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - ch, cw, ch, 3, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.rect(MARGIN, y - ch, 3 * mm, ch, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + 6 * mm, y - 8 * mm, "OUR MISSION")
    draw_paragraph(
        c,
        "Establish a strong and reliable working bond with clients, delivering quality, affordable and up-to-standard services for all stakeholders.",
        MARGIN + 6 * mm,
        y - 16 * mm,
        cw - 10 * mm,
        size=8.5,
        leading=11,
        color=white,
    )

    c.setFillColor(SOFT_BLUE)
    c.roundRect(MARGIN + cw + gap, y - ch, cw, ch, 3, fill=1, stroke=0)
    c.setFillColor(DEEP_BLUE)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + cw + gap + 6 * mm, y - 8 * mm, "OUR VISION")
    draw_paragraph(
        c,
        "Evolve into the leading company providing unsurpassed services to participants across the globe.",
        MARGIN + cw + gap + 6 * mm,
        y - 16 * mm,
        cw - 10 * mm,
        size=8.5,
        leading=11,
        color=TEXT,
    )

    y = y - ch - 6 * mm
    section_label(c, "What Guides Us", MARGIN, y)
    y -= 7 * mm
    values = [
        ("Reliable", "We deliver systems that last and support clients can trust."),
        ("Efficient", "We plan carefully, execute cleanly, and finish on time."),
        ("Sustainable", "We build solar and electrical solutions for a brighter tomorrow."),
        ("Client-first", "We do not just install systems — we deliver lasting solutions."),
        ("Safe", "Site discipline, PPE, and careful commissioning on every job."),
        ("Smart", "Practical engineering choices matched to real client needs."),
    ]
    gap = 2.5 * mm
    vw = (PAGE_W - 2 * MARGIN - 2 * gap) / 3
    vh = 24 * mm
    for i, (t, d) in enumerate(values):
        col = i % 3
        row = i // 3
        x = MARGIN + col * (vw + gap)
        yy = y - row * (vh + gap)
        c.setFillColor(white)
        c.setStrokeColor(LEMON)
        c.setLineWidth(1.1)
        c.roundRect(x, yy - vh, vw, vh, 2, fill=1, stroke=1)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 8.5)
        c.drawString(x + 3 * mm, yy - 7 * mm, t)
        draw_paragraph(c, d, x + 3 * mm, yy - 13 * mm, vw - 6 * mm, size=7, leading=9, color=MUTED)

    y = y - 2 * (vh + gap) - 2 * mm
    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - 14 * mm, PAGE_W - 2 * MARGIN, 14 * mm, 2, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Oblique", 9)
    c.drawCentredString(PAGE_W / 2, y - 5.5 * mm, f'"{TAGLINE}"')
    c.setFillColor(white)
    c.setFont("Helvetica", 6.5)
    c.drawCentredString(PAGE_W / 2, y - 10.5 * mm, "Message from the leadership · Fravent Limited")

    y = y - 17 * mm
    photos = [
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Branded PPE and panel commissioning."),
        (p("nakivale-solar-hybrid-technician-array-rear.jpeg"), "Array installation with Fravent toolkit."),
        (a("cover_field_team_bg.png"), "Teams working safely in the field."),
        (p("ugafode-solar-hybrid-transfer-switch-installation.jpeg"), "Indoor hybrid system integration."),
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Ground array mounting works."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop solar panel fastening."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def draw_people_grid(c: canvas.Canvas, people: list[tuple[str, str, Path]], start_y: float, cols: int = 3):
    gap = 3 * mm
    usable = PAGE_W - 2 * MARGIN
    card_w = (usable - (cols - 1) * gap) / cols
    card_h = 42 * mm
    for i, (name, role, photo) in enumerate(people):
        col = i % cols
        row = i // cols
        x = MARGIN + col * (card_w + gap)
        y = start_y - row * (card_h + 4 * mm)
        c.setFillColor(LIGHT_GRAY)
        c.roundRect(x, y - card_h, card_w, card_h, 3, fill=1, stroke=0)
        r = 11 * mm
        circle_clip_image(c, photo, x + card_w / 2, y - 14 * mm, r)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7.5)
        c.drawCentredString(x + card_w / 2, y - 30 * mm, name)
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 6.5)
        for j, line in enumerate(wrap_text(c, role, "Helvetica", 6.5, card_w - 4 * mm)):
            c.drawCentredString(x + card_w / 2, y - 35 * mm - j * 8, line)
        c.setStrokeColor(LEMON)
        c.setLineWidth(1)
        c.line(x + card_w / 2 - 8 * mm, y - card_h + 3.5 * mm, x + card_w / 2 + 8 * mm, y - card_h + 3.5 * mm)
    rows = (len(people) + cols - 1) // cols
    return start_y - rows * (card_h + 4 * mm)


def page_management(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Leadership", "Management Team")
    y = CONTENT_TOP
    section_label(c, "Management Team", MARGIN, y)
    people = [
        ("Faray Joseph", "Managing Director", a("portrait_faray_joseph.png")),
        ("Olowo Bwenge Norman", "Secretary", a("portrait_olowo_bwenge_norman.png")),
        ("Ekapuloni Francis", "Nominee Director", a("portrait_ekapuloni_francis.png")),
        ("Mutebi Alvin", "Partner", a("portrait_mutebi_alvin.png")),
        ("Tumwesigye Henry", "Alternate Director", a("portrait_tumwesigye_henry.png")),
    ]
    y = draw_people_grid(c, people, y - 5 * mm, cols=5) - 3 * mm

    photos = [
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Hands-on leadership on solar array works."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Technical oversight during commissioning."),
        (p("kiwoko-hospital-adara-team-onsite.jpeg"), "Project leadership with partners on site."),
        (a("cover_field_team_main.png"), "Field teams coordinated for delivery."),
        (p("ugafode-nakivale-branch-site.jpeg"), "Completed client sites under management care."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Commissioned systems ready for clients."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_advisors(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Governance", "Board of Advisors")
    y = CONTENT_TOP
    section_label(c, "Board of Advisors", MARGIN, y)
    people = [
        ("Mr. Mukaaya Eddie", "Advisor", a("portrait_mukaaya_eddie.png")),
        ("Mr. Mawanda Hamzah", "Advisor", a("portrait_mawanda_hamzah.png")),
        ("Mr. Kibirige David", "Advisor", a("portrait_kibirige_david.png")),
        ("Mr. Joshua S. Mugabi", "Advisor", a("portrait_joshua_s_mugabi.png")),
    ]
    y = draw_people_grid(c, people, y - 5 * mm, cols=4) - 3 * mm
    section_label(c, "In the Field", MARGIN, y)
    y -= 5 * mm
    photos = [
        (a("cover_field_team_main.png"), "Advisor-supported field delivery teams."),
        (a("cover_field_team_bg.png"), "Community and site engagement."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Technical works guided by experience."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Safe rooftop installation practice."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Reliable switchgear integration."),
        (p("kiwoko-jinko-hybrid-commissioning-work.jpeg"), "Commissioning discipline on larger systems."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_technical(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Our People", "Technical & Operations")
    y = CONTENT_TOP
    section_label(c, "Electrical Technical Leadership", MARGIN, y)
    y -= 5 * mm

    leaders = [
        ("Ekapuloni Francis", "Project & Technical Manager"),
        ("Tumwesigye Henry", "Project & Technical Advisor"),
        ("John Odutu", "Senior Technical Supervisor"),
        ("Osinde Mark", "Supervisor"),
        ("Richard Olec", "Supervisor"),
        ("James Kayeny", "Supervisor"),
        ("Melissa Johnson", "Supervisor"),
        ("Nicholas Wakatama", "Supervisor"),
        ("Egesa Deo", "Solar Sales Expert"),
    ]

    gap = 2.5 * mm
    cw = (PAGE_W - 2 * MARGIN - 2 * gap) / 3
    ch = 14 * mm
    for i, (name, role) in enumerate(leaders):
        col = i % 3
        row = i // 3
        x = MARGIN + col * (cw + gap)
        yy = y - row * (ch + 2 * mm)
        c.setFillColor(LIGHT_GRAY if row % 2 == 0 else SOFT_BLUE)
        c.roundRect(x, yy - ch, cw, ch, 2, fill=1, stroke=0)
        c.setFillColor(LEMON)
        c.circle(x + 3.5 * mm, yy - ch / 2, 1.8 * mm, fill=1, stroke=0)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7)
        c.drawString(x + 8 * mm, yy - 5 * mm, name)
        c.setFillColor(MUTED)
        c.setFont("Helvetica", 6)
        c.drawString(x + 8 * mm, yy - 10 * mm, role)

    y = y - 3 * (ch + 2 * mm) - 4 * mm
    section_label(c, "Electrical Technical & Operations Team", MARGIN, y)
    y -= 5 * mm
    ops = [
        "Sebufu Moses — Technician & Foreman",
        "Oburu David — Technician & Foreman",
        "Ongala Denish — Technician & Foreman",
        "Omella Thomas — Technician & Linesman",
        "Segawa Marvin — Technician",
        "Katushabe Bernard — Technician",
        "Mbazira Ronald — Technician",
        "Sendagire Derrick — Technician",
        "Mwesigwa Joseph — Linesman",
    ]
    for i, item in enumerate(ops):
        col = i % 3
        row = i // 3
        x = MARGIN + col * ((PAGE_W - 2 * MARGIN) / 3)
        yy = y - row * 8 * mm
        c.setFillColor(DEEP_BLUE)
        c.circle(x + 1.5 * mm, yy + 1.5, 1.3, fill=1, stroke=0)
        c.setFont("Helvetica", 7)
        c.setFillColor(TEXT)
        c.drawString(x + 5 * mm, yy, item)

    y = y - 3 * 8 * mm - 4 * mm
    section_label(c, "Property Team", MARGIN, y)
    y -= 5 * mm
    prop = [
        ("Mutesi Jamirah", "Property Administrator"),
        ("Odoch Derrick", "Sales & Marketing Executive"),
        ("Odida Asiimwe Gerald", "Sales & Marketing Executive"),
        ("Mugaya Edrine", "Marketing Executive"),
    ]
    cw = (PAGE_W - 2 * MARGIN - 3 * 2.5 * mm) / 4
    for i, (name, role) in enumerate(prop):
        x = MARGIN + i * (cw + 2.5 * mm)
        c.setFillColor(DEEP_BLUE)
        c.roundRect(x, y - 16 * mm, cw, 16 * mm, 2, fill=1, stroke=0)
        c.setFillColor(LEMON)
        c.setFont("Helvetica-Bold", 6.5)
        for j, line in enumerate(wrap_text(c, name, "Helvetica-Bold", 6.5, cw - 3 * mm)):
            c.drawCentredString(x + cw / 2, y - 5.5 * mm - j * 7, line)
        c.setFillColor(white)
        c.setFont("Helvetica", 5.8)
        for j, line in enumerate(wrap_text(c, role, "Helvetica", 5.8, cw - 3 * mm)):
            c.drawCentredString(x + cw / 2, y - 12 * mm - j * 6.5, line)

    y = y - 20 * mm
    note = "Our team is skilled and experienced, with nationals of Ugandan origin — dedicated to timely project delivery and seeing Fravent Limited excel."
    y = draw_paragraph(c, note, MARGIN, y, PAGE_W - 2 * MARGIN, size=7.5, leading=9.5, color=MUTED)
    y -= 3 * mm
    photos = [
        (p("nakivale-solar-hybrid-technician-array-rear.jpeg"), "Technicians on solar array works."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Panel commissioning in progress."),
        (p("kiwoko-hospital-adara-team-onsite.jpeg"), "Project team with partners."),
        (a("cover_field_team_main.png"), "Field crew coordination."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=4)
    draw_footer(c, page_no, total)


def page_services_electrical(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Services", "Electrical & Solar")
    y = CONTENT_TOP
    section_label(c, "Electrical Engineering & Renewable Energy", MARGIN, y)
    y -= 6 * mm

    left = [
        "General electrical works (domestic & industrial wiring)",
        "Power line construction",
        "Transformer installation",
        "Solar backup systems",
        "Mini-grid installation",
        "Grid-tied solar systems",
        "Solar water irrigation installation",
        "Renewable energy solutions",
        "Earthing works",
        "Solar hybrid systems (off-grid & hybrid)",
    ]
    right = [
        "Facilities maintenance & engineering",
        "Power generators",
        "Air conditioners",
        "Solar PV systems",
        "Security cameras",
        "Low voltage systems",
        "LV and MV installations",
        "UPS systems & capacity batteries",
        "Switchgear, protection & control relays",
        "Testing, commissioning & handover",
    ]

    col_w = (PAGE_W - 2 * MARGIN - 4 * mm) / 2
    for i, item in enumerate(left):
        yy = y - i * 8.5 * mm
        c.setFillColor(LEMON)
        c.circle(MARGIN + 1.8 * mm, yy + 1.5, 1.5, fill=1, stroke=0)
        c.setFillColor(TEXT)
        c.setFont("Helvetica", 7.5)
        c.drawString(MARGIN + 5.5 * mm, yy, item)
    for i, item in enumerate(right):
        yy = y - i * 8.5 * mm
        c.setFillColor(LEMON)
        c.circle(MARGIN + col_w + 5 * mm, yy + 1.5, 1.5, fill=1, stroke=0)
        c.setFillColor(TEXT)
        c.setFont("Helvetica", 7.5)
        c.drawString(MARGIN + col_w + 9 * mm, yy, item)

    y = y - 10 * 8.5 * mm - 4 * mm
    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - 22 * mm, PAGE_W - 2 * MARGIN, 22 * mm, 3, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(MARGIN + 4 * mm, y - 7 * mm, "Turnkey renewable energy delivery")
    draw_paragraph(
        c,
        "From load assessment and system design to construction, installation, testing, commissioning and monitoring — solar plants and mini-grids for grid-tied and off-grid needs.",
        MARGIN + 4 * mm,
        y - 14 * mm,
        PAGE_W - 2 * MARGIN - 8 * mm,
        size=7.5,
        leading=9.5,
        color=white,
    )

    y = y - 26 * mm
    photos = [
        (p("nakivale-solar-hybrid-electrical-panel-work.jpeg"), "Electrical panel integration."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop solar mounting."),
        (p("nakivale-solar-hybrid-completed-inverter-batteries-team.jpeg"), "Completed hybrid inverter bank."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Changeover & protection panels."),
        (p("kiwoko-jinko-hybrid-battery-inverter-room.jpeg"), "Larger hybrid storage rooms."),
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Ground-mounted solar arrays."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_services_property(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Services", "Property Management")
    y = CONTENT_TOP
    section_label(c, "Property Management Services", MARGIN, y)
    y -= 6 * mm

    services = [
        ("Property Management", "Day-to-day management that keeps properties productive and well maintained."),
        ("Property Sales Agency", "Support for buying and selling with clear market guidance."),
        ("Tenant Search & Placement", "Finding, screening, and placing the right tenants."),
        ("Preparing Property for Rental", "Getting units ready for the market, including repair and maintenance."),
        ("Retrofitting & Consulting", "Practical advice and upgrades that improve performance and value."),
        ("Site Inspections & Reporting", "Regular inspections, reporting, and follow-up action."),
        ("Sub-contractor Supervision", "Hiring and supervising trusted sub-contractors on site."),
        ("Compliance & Remittances", "Submission of governmental remittances and related paperwork."),
        ("Tenant Administration", "Up-to-date tenant lists, notices, leasing, and communication."),
    ]

    gap = 2.5 * mm
    cw = (PAGE_W - 2 * MARGIN - 2 * gap) / 3
    ch = 32 * mm
    for i, (t, d) in enumerate(services):
        col = i % 3
        row = i // 3
        x = MARGIN + col * (cw + gap)
        yy = y - row * (ch + gap)
        c.setFillColor(white)
        c.setStrokeColor(SOFT_BLUE)
        c.setLineWidth(1)
        c.roundRect(x, yy - ch, cw, ch, 2, fill=1, stroke=1)
        c.setFillColor(LEMON)
        c.rect(x, yy - ch, 2.2 * mm, ch, fill=1, stroke=0)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7.5)
        for j, line in enumerate(wrap_text(c, t, "Helvetica-Bold", 7.5, cw - 7 * mm)):
            c.drawString(x + 4 * mm, yy - 7 * mm - j * 9, line)
        draw_paragraph(c, d, x + 4 * mm, yy - 17 * mm, cw - 7 * mm, size=6.8, leading=8.5, color=MUTED)

    y = y - 3 * (ch + gap) - 3 * mm
    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - 14 * mm, PAGE_W - 2 * MARGIN, 14 * mm, 2, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("Helvetica", 8)
    c.drawCentredString(PAGE_W / 2, y - 8.5 * mm, "Electrical excellence · Property care · Lasting client partnerships")

    y = y - 18 * mm
    photos = [
        (a("cover_field_team_main.png"), "Site readiness and maintenance teams."),
        (p("ugafode-nakivale-branch-site.jpeg"), "Managed facility environments."),
        (p("ugafode-branch-completed-deye-customer-area.jpeg"), "Client spaces powered and maintained."),
        (a("cover_pole_climber.png"), "Technical support beyond the office."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=4)
    draw_footer(c, page_no, total)


def page_capabilities(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Capabilities", "How We Deliver")
    y = CONTENT_TOP

    blocks = [
        ("Power line construction", "We serve the power-line and energy industry with pride and integrity. Our teams deliver projects ranging from 11kV to 33kV, with continuous training across a wide scope of construction work."),
        ("Substation development", "For industrial, commercial, or residential power sub-stations, we assemble resources into a turnkey package — from ground-breaking to finished projects."),
        ("LV and MV installations", "Design, supply and construction of LV and MV installations, compact secondary sub-stations, switchgear, protection and control relays, transformers, batteries, filters, and UPS systems."),
        ("Renewable energy", "Load assessment, system design, construction, installation, testing, commissioning and monitoring of solar power plants and solar mini-grids — both grid-tied and off-grid."),
    ]
    for title, body in blocks:
        c.setFillColor(LIGHT_GRAY)
        c.roundRect(MARGIN, y - 26 * mm, PAGE_W - 2 * MARGIN, 26 * mm, 2, fill=1, stroke=0)
        c.setFillColor(LEMON)
        c.rect(MARGIN, y - 26 * mm, 2.5 * mm, 26 * mm, fill=1, stroke=0)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(MARGIN + 5 * mm, y - 7 * mm, title)
        draw_paragraph(c, body, MARGIN + 5 * mm, y - 14 * mm, PAGE_W - 2 * MARGIN - 10 * mm, size=7.5, leading=9.5, color=TEXT)
        y -= 28 * mm

    y -= 2 * mm
    section_label(c, "Delivery path we follow", MARGIN, y)
    y -= 5 * mm
    steps = ["Assess", "Design", "Procure", "Install", "Test", "Commission", "Support"]
    gap = 2 * mm
    sw = (PAGE_W - 2 * MARGIN - 6 * gap) / 7
    for i, s in enumerate(steps):
        x = MARGIN + i * (sw + gap)
        c.setFillColor(DEEP_BLUE if i % 2 == 0 else LEMON)
        c.roundRect(x, y - 12 * mm, sw, 12 * mm, 2, fill=1, stroke=0)
        c.setFillColor(white if i % 2 == 0 else DEEP_BLUE)
        c.setFont("Helvetica-Bold", 6.5)
        c.drawCentredString(x + sw / 2, y - 7.5 * mm, s)

    y = y - 16 * mm
    photos = [
        (a("project_prominigrid_1.png"), "Power-line and last-mile works."),
        (a("project_gomba_1.png"), "Community solar plant delivery."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Hybrid systems commissioned."),
        (p("kiwoko-jinko-hybrid-battery-inverter-room.jpeg"), "Larger storage installations."),
        (a("project_mazima_1.png"), "Community solar commissioning."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Safe changeover integration."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_featured_overview(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Featured Project", "Solar Hybrid · Off Grid")
    y = CONTENT_TOP

    left_w = (PAGE_W - 2 * MARGIN) * 0.56
    right_w = (PAGE_W - 2 * MARGIN) * 0.42
    draw_captioned_image(
        c,
        p("ugafode-nakivale-branch-site.jpeg"),
        MARGIN,
        y - 62 * mm,
        left_w,
        62 * mm,
        "UGAFODE Microfinance Limited (MDI) Nakivale Branch — project site.",
    )

    rx = MARGIN + left_w + 3 * mm
    ry = y
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 11)
    for line in wrap_text(c, "Solar Hybrid System — OFF Grid", "Helvetica-Bold", 11, right_w):
        c.drawString(rx, ry, line)
        ry -= 12
    c.setFillColor(MUTED)
    c.setFont("Helvetica", 7.5)
    c.drawString(rx, ry - 2 * mm, "Completed and Commissioned")

    meta = [
        ("Description:", "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System."),
        ("Client:", "UGAFODE Microfinance Limited - MDI"),
        ("Location:", "Nakivale Refugee Settlement - Isingiro District, Western Uganda"),
        ("System Size:", "5.58 kWh"),
        ("Status:", "Completed and Commissioned"),
        ("Scope:", "Solar array, hybrid inverter, battery storage, electrical panels, testing and handover for reliable branch power."),
    ]
    my = ry - 12 * mm
    for label, value in meta:
        my = meta_row(c, label, value, rx, my, label_w=28 * mm, max_w=right_w - 28 * mm)

    y = CONTENT_TOP - 66 * mm
    photos = [
        (p("ugafode-nakivale-branch-directional-sign.jpeg"), "UGAFODE Nakivale Branch directional signage."),
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Ground-mounted solar array installation."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop solar panel mounting."),
        (p("nakivale-solar-hybrid-technician-array-rear.jpeg"), "Fravent technician on array works."),
        (p("nakivale-solar-hybrid-electrical-panel-work.jpeg"), "Indoor electrical panel integration."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Control panel commissioning."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_featured_install(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Featured Project", "Installation in Progress")
    y = CONTENT_TOP
    section_label(c, "Field Installation — Nakivale (5.58 kWh)", MARGIN, y)
    y -= 5 * mm

    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - 18 * mm, PAGE_W - 2 * MARGIN, 18 * mm, 2, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(MARGIN + 4 * mm, y - 6 * mm, "Quality workmanship on site")
    draw_paragraph(
        c,
        "Solar array mounting, rooftop installation, electrical panel integration, and commissioning for reliable UGAFODE Nakivale branch operations.",
        MARGIN + 4 * mm,
        y - 12 * mm,
        PAGE_W - 2 * MARGIN - 8 * mm,
        size=7,
        leading=9,
        color=white,
    )
    y = y - 22 * mm

    photos = [
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Technician securing ground-array mounts."),
        (p("nakivale-solar-hybrid-technician-array-rear.jpeg"), "Branded PPE during array works."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop panel fastening."),
        (p("nakivale-solar-hybrid-electrical-panel-work.jpeg"), "Electrical panel integration."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Control panel wiring."),
        (p("ugafode-solar-hybrid-transfer-switch-installation.jpeg"), "Transfer switch installation."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_featured_completed(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Featured Project", "Completed System")
    y = CONTENT_TOP
    section_label(c, "Commissioned Off-Grid Hybrid System — 5.58 kWh", MARGIN, y)
    y -= 5 * mm

    cards = [
        ("Hybrid Inverter", "Deye hybrid inverter commissioned for off-grid branch power."),
        ("Battery Storage", "5.58 kWh lithium storage for overnight and backup supply."),
        ("Clean Integration", "Neat trunking, distribution boards, and isolators."),
        ("Client Outcome", "Stable power for UGAFODE MDI at Nakivale."),
    ]
    gap = 2.5 * mm
    cw = (PAGE_W - 2 * MARGIN - 3 * gap) / 4
    ch = 28 * mm
    for i, (t, d) in enumerate(cards):
        x = MARGIN + i * (cw + gap)
        c.setFillColor(DEEP_BLUE if i % 2 == 0 else LEMON)
        c.roundRect(x, y - ch, cw, ch, 2, fill=1, stroke=0)
        c.setFillColor(white if i % 2 == 0 else DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7)
        c.drawString(x + 2 * mm, y - 6 * mm, t)
        draw_paragraph(c, d, x + 2 * mm, y - 12 * mm, cw - 4 * mm, size=6.2, leading=8, color=white if i % 2 == 0 else DEEP_BLUE)

    y = y - ch - 4 * mm
    photos = [
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Completed Deye hybrid inverter and batteries."),
        (p("nakivale-solar-hybrid-completed-inverter-batteries-team.jpeg"), "Commissioned wall-mounted system."),
        (p("ugafode-branch-completed-deye-customer-area.jpeg"), "System view from customer area."),
        (p("ugafode-branch-completed-deye-system-bw.jpeg"), "Clean completed wall installation."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Mains / generator / inverter changeover."),
        (p("ugafode-nakivale-branch-site.jpeg"), "Powered UGAFODE Nakivale Branch."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_ugafode_portfolio(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "More Projects", "UGAFODE Solar Hybrid Sites")
    y = CONTENT_TOP
    section_label(c, "Completed Solar Hybrid Systems — UGAFODE Microfinance Limited (MDI)", MARGIN, y)
    y -= 6 * mm

    projects = [
        {
            "title": "Rushere — Western Uganda",
            "size": "6.0 kWh",
            "status": "Completed and Commissioned",
            "note": "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System.",
        },
        {
            "title": "Ntungamo — Western Uganda",
            "size": "6.0 kWh",
            "status": "Completed and Commissioned",
            "note": "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System.",
        },
        {
            "title": "Kyaka Sales Branch",
            "size": "5.0 kWp DC",
            "status": "Completed",
            "note": "Procurement, Delivery, Installation, Testing and Commissioning of a Hybrid Solar System.",
        },
        {
            "title": "Rwamwanja Sales Branch",
            "size": "6.25 kWp DC",
            "status": "Completed",
            "note": "Procurement, Delivery, Installation, Testing and Commissioning of a Hybrid Solar System.",
        },
        {
            "title": "Nakivale — Isingiro District",
            "size": "5.58 kWh",
            "status": "Completed and Commissioned",
            "note": "Off-grid solar hybrid system for the Nakivale Refugee Settlement branch.",
        },
        {
            "title": "Programme focus",
            "size": "Multi-site",
            "status": "Active client account",
            "note": "Repeat UGAFODE branch deployments across Western Uganda and settlement markets.",
        },
    ]

    gap = 2.5 * mm
    cw = (PAGE_W - 2 * MARGIN - 2 * gap) / 3
    ch = 34 * mm
    for i, proj in enumerate(projects):
        col = i % 3
        row = i // 3
        x = MARGIN + col * (cw + gap)
        yy = y - row * (ch + gap)
        c.setFillColor(LIGHT_GRAY)
        c.roundRect(x, yy - ch, cw, ch, 2, fill=1, stroke=0)
        c.setFillColor(LEMON)
        c.rect(x, yy - ch, 2.5 * mm, ch, fill=1, stroke=0)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 7.5)
        c.drawString(x + 4.5 * mm, yy - 7 * mm, proj["title"])
        c.setFont("Helvetica", 6.5)
        c.setFillColor(TEXT)
        c.drawString(x + 4.5 * mm, yy - 13 * mm, f"Size: {proj['size']}")
        c.drawString(x + 4.5 * mm, yy - 18.5 * mm, f"Status: {proj['status']}")
        draw_paragraph(c, proj["note"], x + 4.5 * mm, yy - 24 * mm, cw - 7 * mm, size=6.2, leading=8, color=MUTED)

    y = y - 2 * (ch + gap) - 3 * mm
    section_label(c, "Branch installation gallery", MARGIN, y)
    y -= 4 * mm
    photos = [
        (p("ugafode-solar-hybrid-transfer-switch-installation.jpeg"), "Transfer switch and Deye battery works."),
        (p("ugafode-branch-completed-deye-system-bw.jpeg"), "Completed Deye hybrid wall mount."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Manual changeover controls."),
        (p("ugafode-branch-completed-deye-customer-area.jpeg"), "Completed system in branch interior."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Commissioned inverter and batteries."),
        (p("ugafode-nakivale-branch-directional-sign.jpeg"), "UGAFODE Nakivale branch identity."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_kiwoko_project(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "On-going Project", "Kiwoko Hospital")
    y = CONTENT_TOP
    section_label(c, "Solar Hybrid System — Kiwoko Hospital", MARGIN, y)
    y -= 5 * mm

    # Compact meta strip
    c.setFillColor(SOFT_BLUE)
    c.roundRect(MARGIN, y - 28 * mm, PAGE_W - 2 * MARGIN, 28 * mm, 2, fill=1, stroke=0)
    meta = [
        ("Description:", "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System."),
        ("Location:", "Kiwoko Hospital - Nakaseke District"),
        ("System Size:", "10 kWh"),
        ("Client:", "Adara Group"),
        ("Status:", "On-going"),
    ]
    my = y - 6 * mm
    for label, value in meta:
        my = meta_row(c, label, value, MARGIN + 3 * mm, my, label_w=28 * mm, max_w=PAGE_W - 2 * MARGIN - 34 * mm)

    y = y - 32 * mm
    photos = [
        (p("kiwoko-hospital-adara-team-onsite.jpeg"), "Fravent team on site at Adara Development Uganda — Kiwoko."),
        (p("kiwoko-jinko-hybrid-battery-inverter-room.jpeg"), "Jinko hybrid inverters and battery racks."),
        (p("kiwoko-jinko-hybrid-commissioning-work.jpeg"), "Commissioning checks on the 10 kWh system."),
        (p("nakivale-solar-hybrid-electrical-panel-work.jpeg"), "Comparable panel integration standards."),
        (p("ugafode-solar-hybrid-transfer-switch-installation.jpeg"), "Transfer and protection practices."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Detailed commissioning discipline."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=3)
    draw_footer(c, page_no, total)


def page_project_block(
    c: canvas.Canvas,
    page_no: int,
    total: int,
    title: str,
    meta: list[tuple[str, str]],
    photos: list[Path] | list[tuple[Path, str]],
    subtitle: str = "Previous Projects",
):
    draw_header(c, "Portfolio", subtitle)
    y = CONTENT_TOP
    c.setFillColor(DEEP_BLUE)
    c.setFont("Helvetica-Bold", 11)
    for line in wrap_text(c, title, "Helvetica-Bold", 11, PAGE_W - 2 * MARGIN):
        c.drawString(MARGIN, y, line)
        y -= 12
    y -= 2
    for label, value in meta:
        y = meta_row(c, label, value, MARGIN, y, max_w=PAGE_W - 2 * MARGIN - 38 * mm)

    y -= 3 * mm
    items: list[tuple[Path, str]] = []
    for item in photos:
        if isinstance(item, tuple):
            items.append(item)
        else:
            items.append((item, item.stem.replace("-", " ").replace("_", " ").title()))

    if items:
        draw_photo_grid(c, items, MARGIN, y, CONTENT_BOTTOM, cols=2 if len(items) <= 4 else 3)
    draw_footer(c, page_no, total)


def page_other_projects(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Portfolio", "Other Projects")
    y = CONTENT_TOP
    projects = [
        {
            "name": "Hybrid Solar System — Ernest Cook University",
            "location": "Balintuma Road, Mengo",
            "client": "Ernest Cook University",
            "status": "Completed",
            "desc": "Procurement, Delivery, Installation, Testing and Commissioning of a Hybrid Solar System.",
        },
        {
            "name": "Electrical Upgrade and Installation of Automatic Change Over System",
            "location": "Uganda Museum",
            "client": "Equatorial Power",
            "status": "Complete",
            "desc": "Upgrade works and automatic change-over installation for reliable facility power.",
        },
        {
            "name": "Residential Apartment Wiring",
            "location": "Sentema Road, Mengo",
            "purpose": "Domestic Wiring",
            "status": "Complete",
            "desc": "Full domestic wiring package for a residential apartment development.",
        },
        {
            "name": "Maintenance of Lighting System",
            "location": "Luthuli Avenue, Bugolobi",
            "client": "Madam Ivy",
            "status": "Complete",
            "desc": "Lighting system maintenance and service support for a private client.",
        },
    ]
    for proj in projects:
        c.setFillColor(LIGHT_GRAY)
        c.roundRect(MARGIN, y - 28 * mm, PAGE_W - 2 * MARGIN, 28 * mm, 2, fill=1, stroke=0)
        c.setFillColor(LEMON)
        c.rect(MARGIN, y - 28 * mm, 2.5 * mm, 28 * mm, fill=1, stroke=0)
        c.setFillColor(DEEP_BLUE)
        c.setFont("Helvetica-Bold", 8)
        ty = y - 7 * mm
        for line in wrap_text(c, proj["name"], "Helvetica-Bold", 8, PAGE_W - 2 * MARGIN - 10 * mm):
            c.drawString(MARGIN + 5 * mm, ty, line)
            ty -= 9
        c.setFont("Helvetica", 6.8)
        c.setFillColor(TEXT)
        details = []
        for key in ("location", "client", "purpose", "status"):
            if key in proj:
                details.append(f"{key.title()}: {proj[key]}")
        c.drawString(MARGIN + 5 * mm, y - 18 * mm, "  ·  ".join(details))
        if "desc" in proj:
            c.setFillColor(MUTED)
            c.setFont("Helvetica", 6.5)
            c.drawString(MARGIN + 5 * mm, y - 24 * mm, proj["desc"][:120])
        y -= 30 * mm

    c.setFillColor(DEEP_BLUE)
    c.roundRect(MARGIN, y - 14 * mm, PAGE_W - 2 * MARGIN, 14 * mm, 2, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(MARGIN + 4 * mm, y - 5.5 * mm, "Ready for the next assignment")
    c.setFillColor(white)
    c.setFont("Helvetica", 7)
    c.drawString(MARGIN + 4 * mm, y - 11 * mm, "From UGAFODE hybrids and hospital systems to mini-grids and facility upgrades across Uganda.")

    y = y - 18 * mm
    photos = [
        (a("cover_pole_climber.png"), "Power-line and upgrade capability."),
        (a("project_prominigrid_2.png"), "Last-mile connection experience."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Hybrid systems ready for clients."),
        (p("kiwoko-jinko-hybrid-battery-inverter-room.jpeg"), "Larger facility-scale installs."),
    ]
    draw_photo_grid(c, photos, MARGIN, y, CONTENT_BOTTOM, cols=4)
    draw_footer(c, page_no, total)


def page_partners(c: canvas.Canvas, page_no: int, total: int):
    draw_header(c, "Network", "Business Partners")
    y = PAGE_H - 36 * mm
    section_label(c, "Organisations We Work With", MARGIN, y)
    y -= 8 * mm

    partner_files = sorted(ASSETS.glob("partner_*.png"))
    # Prefer larger logos
    usable = [p for p in partner_files if PILImage.open(p).size[0] >= 100][:8]
    if not usable:
        usable = partner_files[:8]

    gap = 5 * mm
    cols = 4
    cw = (PAGE_W - 2 * MARGIN - (cols - 1) * gap) / cols
    ch = 28 * mm
    for i, ph in enumerate(usable):
        col = i % cols
        row = i // cols
        x = MARGIN + col * (cw + gap)
        yy = y - row * (ch + gap) - ch
        c.setFillColor(white)
        c.setStrokeColor(SOFT_BLUE)
        c.roundRect(x, yy, cw, ch, 3, fill=1, stroke=1)
        draw_image(c, ph, x + 3 * mm, yy + 3 * mm, cw - 6 * mm, ch - 6 * mm, cover=False)

    y = y - 2 * (ch + gap) - 10 * mm
    names = [
        "Rescue a Child Uganda (RACU)",
        "Hear His Voice Uganda",
        "Forbes Coffee Training Center",
        "Nirvana Wellness Spa",
        "Hope & Glory Foundation",
        "ESO J Technology Solutions Point",
        "RIA ROUND Creations",
        "MEGA Property Services",
    ]
    for i, name in enumerate(names):
        col = i % 2
        row = i // 2
        x = MARGIN + col * ((PAGE_W - 2 * MARGIN) / 2)
        yy = y - row * 10 * mm
        c.setFillColor(LEMON)
        c.circle(x + 2 * mm, yy + 2, 1.6, fill=1, stroke=0)
        c.setFillColor(TEXT)
        c.setFont("Helvetica", 8.5)
        c.drawString(x + 6 * mm, yy, name)

    y = y - 5 * 10 * mm - 6 * mm
    draw_image(c, a("cover_field_team_bg.png"), MARGIN, 18 * mm, PAGE_W - 2 * MARGIN, max(30 * mm, y - 18 * mm))

    draw_footer(c, page_no, total)


def page_contact(c: canvas.Canvas, page_no: int, total: int):
    # Full branded closing
    c.setFillColor(DEEP_BLUE)
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    c.setFillColor(LEMON)
    c.rect(0, 0, 8 * mm, PAGE_H, fill=1, stroke=0)

    draw_logo_mark(c, MARGIN + 4 * mm, PAGE_H - 48 * mm, size=20 * mm, light=True)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 26)
    c.drawString(MARGIN + 28 * mm, PAGE_H - 32 * mm, "FRAVENT LIMITED")
    c.setFillColor(LEMON)
    c.setFont("Helvetica", 11)
    c.drawString(MARGIN + 28 * mm, PAGE_H - 40 * mm, TAGLINE)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 22)
    c.drawString(MARGIN + 4 * mm, PAGE_H - 70 * mm, "Let's power")
    c.setFillColor(LEMON)
    c.drawString(MARGIN + 4 * mm, PAGE_H - 82 * mm, "your next project.")

    y = PAGE_H - 105 * mm
    details = [
        ("Address", "P. O. Box 108053, Kampala (U)\nRofra House, 04th Floor,\nKansanga – Ggaba Road"),
        ("Phone", "+256 778 455 042\n+256 703 305 262"),
        ("Email", "fraventlimited@gmail.com"),
        ("Website", "www.fraventlimited.com"),
    ]
    for label, value in details:
        c.setFillColor(LEMON)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(MARGIN + 4 * mm, y, label.upper())
        c.setFillColor(white)
        c.setFont("Helvetica", 10)
        y -= 12
        for line in value.split("\n"):
            c.drawString(MARGIN + 4 * mm, y, line)
            y -= 12
        y -= 6

    # Photo strip
    photos = [
        p("nakivale-solar-hybrid-ground-array-installation.jpeg"),
        p("nakivale-solar-hybrid-completed-deye-system.jpeg"),
        p("fravent-brand-colors-flyer.jpeg"),
    ]
    gap = 3 * mm
    cw = (PAGE_W - MARGIN - 12 * mm - 2 * gap) / 3
    ch = 40 * mm
    for i, ph in enumerate(photos):
        draw_image(c, ph, MARGIN + 4 * mm + i * (cw + gap), 28 * mm, cw, ch)

    c.setFillColor(LEMON)
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(PAGE_W / 2, 16 * mm, SLOGAN)
    c.setFillColor(white)
    c.setFont("Helvetica", 8)
    c.drawRightString(PAGE_W - MARGIN, 16 * mm, f"{page_no} / {total}")


def build_pdf():
    c = canvas.Canvas(str(OUT_PDF), pagesize=A4)
    pages = [
        page_cover,
        page_about,
        page_mission,
        page_management,
        page_advisors,
        page_technical,
        page_services_electrical,
        page_services_property,
        page_capabilities,
        page_featured_overview,
        page_featured_install,
        page_featured_completed,
        page_ugafode_portfolio,
        page_kiwoko_project,
        lambda cv, n, t: page_project_block(
            cv,
            n,
            t,
            "PROMOTION OF MINI-GRID FOR RURAL ELECTRIFICATION (Pro Mini Grids)",
            [
                ("Location:", "Lamwo District - Uganda (Paloga, Potika, Pangira, Labayango, Lapidyeni, Pawena, Apwoyo, Apyetta Town Council & Apyeta West, Kapeta, Muddu, Agoro & Lollimebeng)"),
                ("Duration:", "Three (3) Months"),
                ("Purpose:", "Last Mile Connection / Home Connection / House wiring"),
                ("Funders:", "GIZ, REA, Government of Uganda"),
                ("Contractor:", "Sagemcom Limited"),
                ("Client:", "WINCH Energy"),
            ],
            [
                (a(f"project_prominigrid_{i}.png"), f"Pro Mini Grids field works — Lamwo District ({i}/4).")
                for i in range(1, 5)
            ],
        ),
        lambda cv, n, t: page_project_block(
            cv,
            n,
            t,
            "SHARED COMMUNITY SOLAR PROJECT - Pro Mini Grid",
            [
                ("Location:", "Kalyamawolu Village, Kyabagamba Sub-county, Maddu Parish, Gomba District"),
                ("Client:", "Ministry of Local Government"),
                ("Status:", "Completed & Commissioned"),
            ],
            [
                (a(f"project_gomba_{i}.png"), f"Shared community solar project — Gomba District ({i}/4).")
                for i in range(1, 5)
            ],
        ),
        lambda cv, n, t: page_project_block(
            cv,
            n,
            t,
            "MAZIMA COMMUNITY SOLAR SYSTEM",
            [
                ("Location:", "Buwenge, Jinja District - Uganda"),
                ("Client:", "Mazima Community Development Center"),
                ("Status:", "Completed and Commissioned"),
            ],
            [
                (a(f"project_mazima_{i}.png"), f"Mazima Community Solar System — Jinja District ({i}/4).")
                for i in range(1, 5)
            ],
        ),
        page_other_projects,
        page_partners,
        page_contact,
    ]

    total = len(pages)
    for i, fn in enumerate(pages, start=1):
        if fn is page_cover:
            fn(c)
        else:
            fn(c, i, total)
        c.showPage()

    c.save()
    print(f"Wrote {OUT_PDF} ({total} pages)")
    return total


def build_docx():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    def set_run_color(run, hex_color: str):
        run.font.color.rgb = RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

    def add_heading_bar(text: str, sub: str = ""):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        set_run_color(run, "071B36")
        if sub:
            p2 = doc.add_paragraph()
            r2 = p2.add_run(sub)
            r2.font.size = Pt(10)
            set_run_color(r2, "A3C324")

    def add_caption(text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        run.font.size = Pt(8)
        set_run_color(run, "5A6570")

    def add_picture(path: Path, width: float = 5.8, caption: str = ""):
        if path.exists():
            doc.add_picture(str(path), width=Inches(width))
            if caption:
                add_caption(caption)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("FRAVENT LIMITED")
    r.bold = True
    r.font.size = Pt(28)
    set_run_color(r, "071B36")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("COMPANY PROFILE 2026")
    r.bold = True
    r.font.size = Pt(18)
    set_run_color(r, "A3C324")

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run(f'"{TAGLINE}"')
    r.italic = True
    r.font.size = Pt(12)
    set_run_color(r, "071B36")

    add_picture(
        p("nakivale-solar-hybrid-ground-array-installation.jpeg"),
        5.8,
        "Cover photo: Solar Hybrid System installation for UGAFODE at Nakivale Refugee Settlement.",
    )
    doc.add_paragraph("Electrical & Solar Solutions  ·  Safe. Smart. Sustainable.")
    doc.add_paragraph("www.fraventlimited.com  ·  fraventlimited@gmail.com  ·  +256 778 455 042")
    doc.add_page_break()

    # About
    add_heading_bar("About Us", "Our Story")
    for para in [
        "Fravent Limited is a fully registered company with the Registrar of Companies (Registration Number: 80020003614101). We help clients realise their goals in electrical engineering, solar and renewable energy, property management, consultancy and advisory services, import and export, logistics and transport, land brokerage, and business intelligence.",
        "We employ experienced professional staff for every project. Our clients include private individuals, schools, private and public companies, organisations, local and international investors, and government ministries.",
        "A steady flow of orders reflects trust in Fravent management's heritage and diversified experience, built on successful project delivery year after year.",
    ]:
        doc.add_paragraph(para)
    doc.add_page_break()

    # Mission
    add_heading_bar("Purpose", "Mission · Vision · Values")
    para = doc.add_paragraph()
    r = para.add_run("Mission: ")
    r.bold = True
    para.add_run("Establish a strong and reliable working bond with clients, delivering quality, affordable and up-to-standard services for all stakeholders.")
    para = doc.add_paragraph()
    r = para.add_run("Vision: ")
    r.bold = True
    para.add_run("Evolve into the leading company providing unsurpassed services to participants across the globe.")
    for v in ["Reliable", "Efficient", "Sustainable", "Client-first"]:
        doc.add_paragraph(v, style="List Bullet")
    doc.add_page_break()

    # Featured Nakivale
    add_heading_bar("Featured Project", "Solar Hybrid System — OFF Grid")
    meta = [
        ("Description", "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System."),
        ("Client", "UGAFODE Microfinance Limited - MDI"),
        ("Location", "Nakivale Refugee Settlement - Isingiro District, Western Uganda"),
        ("System Size", "5.58 kWh"),
        ("Status", "Completed and Commissioned"),
    ]
    for k, v in meta:
        para = doc.add_paragraph()
        r = para.add_run(f"{k}: ")
        r.bold = True
        para.add_run(v)

    nakivale_photos = [
        (p("ugafode-nakivale-branch-site.jpeg"), "UGAFODE Microfinance Limited (MDI) Nakivale Branch — project site."),
        (p("ugafode-nakivale-branch-directional-sign.jpeg"), "UGAFODE Nakivale Branch directional signage."),
        (p("nakivale-solar-hybrid-ground-array-installation.jpeg"), "Ground-mounted solar array installation at Nakivale."),
        (p("nakivale-solar-hybrid-rooftop-panel-installation.jpeg"), "Rooftop solar panel mounting by Fravent technicians."),
        (p("nakivale-solar-hybrid-technician-array-rear.jpeg"), "Fravent branded PPE during array works."),
        (p("nakivale-solar-hybrid-electrical-panel-work.jpeg"), "Indoor electrical panel integration."),
        (p("nakivale-solar-hybrid-control-panel-installation.jpeg"), "Control panel wiring and commissioning."),
        (p("nakivale-solar-hybrid-completed-deye-system.jpeg"), "Completed Deye hybrid inverter and battery installation."),
        (p("nakivale-solar-hybrid-completed-inverter-batteries-team.jpeg"), "Commissioned wall-mounted inverter and battery bank."),
        (p("ugafode-branch-completed-deye-customer-area.jpeg"), "Completed system viewed from the UGAFODE customer area."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Changeover panel — mains, generator, and inverter paths."),
    ]
    for path, cap in nakivale_photos:
        add_picture(path, 5.5, cap)
    doc.add_page_break()

    # UGAFODE portfolio
    add_heading_bar("More Projects", "UGAFODE Solar Hybrid Sites")
    ugafode_sites = [
        ("Rushere — Western Uganda", "6.0 kWh", "Completed and Commissioned"),
        ("Ntungamo — Western Uganda", "6.0 kWh", "Completed and Commissioned"),
        ("Nakivale Refugee Settlement — Isingiro District", "5.58 kWh", "Completed and Commissioned"),
        ("Kyaka Sales Branch", "5.0 kWp DC", "Completed"),
        ("Rwamwanja Sales Branch", "6.25 kWp DC", "Completed"),
    ]
    for title, size, status in ugafode_sites:
        para = doc.add_paragraph()
        r = para.add_run(title)
        r.bold = True
        doc.add_paragraph("Description: Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System.")
        doc.add_paragraph(f"Client: UGAFODE Microfinance Limited - MDI")
        doc.add_paragraph(f"System Size: {size}")
        doc.add_paragraph(f"Status: {status}")
        doc.add_paragraph("")

    for path, cap in [
        (p("ugafode-solar-hybrid-transfer-switch-installation.jpeg"), "Transfer switch and Deye battery works at a UGAFODE branch."),
        (p("ugafode-branch-completed-deye-system-bw.jpeg"), "Completed Deye hybrid inverter and battery wall mount."),
        (p("ugafode-solar-hybrid-changeover-control-panel.jpeg"), "Manual changeover controls for mains, generator, and inverter."),
    ]:
        add_picture(path, 5.2, cap)
    doc.add_page_break()

    # Kiwoko
    add_heading_bar("On-going Project", "Kiwoko Hospital — Adara Group")
    for k, v in [
        ("Description", "Procurement, Delivery, Installation, Testing and Commissioning of a Solar Hybrid System."),
        ("Location", "Kiwoko Hospital - Nakaseke District"),
        ("System Size", "10 kWh"),
        ("Client", "Adara Group"),
        ("Status", "On-going"),
    ]:
        para = doc.add_paragraph()
        r = para.add_run(f"{k}: ")
        r.bold = True
        para.add_run(v)
    for path, cap in [
        (p("kiwoko-hospital-adara-team-onsite.jpeg"), "Fravent team on site at Adara Development Uganda — Kiwoko, Nakaseke District."),
        (p("kiwoko-jinko-hybrid-battery-inverter-room.jpeg"), "Jinko hybrid inverters and battery racks during installation."),
        (p("kiwoko-jinko-hybrid-commissioning-work.jpeg"), "Commissioning checks on the 10 kWh hybrid storage system."),
    ]:
        add_picture(path, 5.5, cap)
    doc.add_page_break()

    # Earlier portfolio
    add_heading_bar("Earlier Portfolio", "Community Solar & Mini-Grids")
    doc.add_paragraph("Pro Mini Grids — Lamwo District (WINCH Energy / Sagemcom / GIZ, REA, Government of Uganda).")
    for i in range(1, 5):
        add_picture(a(f"project_prominigrid_{i}.png"), 4.8, f"Pro Mini Grids field works — Lamwo District ({i}/4).")
    doc.add_paragraph("Shared Community Solar — Gomba District (Ministry of Local Government).")
    for i in range(1, 5):
        add_picture(a(f"project_gomba_{i}.png"), 4.8, f"Shared community solar project — Gomba District ({i}/4).")
    doc.add_paragraph("Mazima Community Solar System — Buwenge, Jinja District.")
    for i in range(1, 5):
        add_picture(a(f"project_mazima_{i}.png"), 4.8, f"Mazima Community Solar System — Jinja District ({i}/4).")
    doc.add_page_break()

    # Other
    add_heading_bar("Other Projects")
    for name, detail in [
        ("Ernest Cook University — Balintuma Road, Mengo", "Hybrid Solar System. Status: Completed."),
        ("Uganda Museum", "Electrical Upgrade and Automatic Change Over — Equatorial Power. Status: Complete."),
        ("Sentema Road, Mengo", "Residential Apartment Wiring. Status: Complete."),
        ("Luthuli Avenue, Bugolobi", "Maintenance of Lighting System — Madam Ivy. Status: Complete."),
    ]:
        para = doc.add_paragraph()
        r = para.add_run(name)
        r.bold = True
        doc.add_paragraph(detail)

    doc.add_page_break()
    add_heading_bar("Contact")
    doc.add_paragraph("Address: P. O. Box 108053, Kampala (U), Rofra House, 04th Floor, Kansanga - Ggaba Road")
    doc.add_paragraph("Phone: +256 778 455 042  ·  +256 703 305 262")
    doc.add_paragraph("Email: fraventlimited@gmail.com")
    doc.add_paragraph("Website: www.fraventlimited.com")
    doc.add_paragraph(SLOGAN)

    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


def build():
    build_pdf()
    build_docx()


if __name__ == "__main__":
    build()
