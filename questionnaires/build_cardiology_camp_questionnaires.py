"""
Build FairBanks Community Cardiology Camp evaluation questionnaires.

Outputs (in this folder):
  - consultant_sho_evaluation.docx
  - consultant_sho_evaluation_fillable.pdf
  - internal_staff_evaluation.docx
  - internal_staff_evaluation_fillable.pdf
"""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

OUT_DIR = Path(__file__).resolve().parent
REPO = OUT_DIR.parent
LOGO = REPO / "assets" / "fairbanks_logo.jpeg"

# Brand aligned to FairBanks logo (green + orange)
BRAND = HexColor("#008C45")
BRAND_DARK = HexColor("#006B35")
ACCENT = HexColor("#F58220")
LIGHT = HexColor("#E8F6EE")
MUTED = HexColor("#5A6A6A")
LINE = HexColor("#C5D5D4")
INK = HexColor("#1C2B2A")
WARN = HexColor("#FFF8E7")

PAGE_W, PAGE_H = A4
MARGIN_L = 16 * mm
MARGIN_R = 16 * mm
MARGIN_T = 12 * mm
MARGIN_B = 20 * mm
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R

SUBMIT_EMAIL = "info@fairbanksmedicalcentre.org"
WEBSITE = "www.fairbanksmedicalcentre.ug"
PHONE = "0701 849258"
WHATSAPP = "0777 462398"
ADDRESS = "Tirupati Road, Kyebando-Kisalosalo (Opp. Northern Bypass Roundabout), Kampala"
SLOGAN = "Your health, our mission."
CONTACT_LINE_1 = ADDRESS
CONTACT_LINE_2 = f"Tel: {PHONE}  |  WhatsApp: {WHATSAPP}"
CONTACT_LINE_3 = f"Email: {SUBMIT_EMAIL}  |  Web: {WEBSITE}"


# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

CONSULTANT_RATINGS = [
    "Overall organisation",
    "Communication before the camp",
    "Coordination during the camp",
    "Patient registration process",
    "Patient flow / triage",
    "Availability of investigations",
    "Availability of medicines",
    "Support from organising team",
    "Teamwork among health workers",
    "Quality of clinical workspace",
    "Overall experience",
]

STAFF_RATINGS = [
    "Pre-camp briefing and role clarity",
    "Communication during the camp",
    "Registration and crowd control",
    "Patient flow and triage",
    "Station setup and workspace",
    "Availability of supplies / consumables",
    "Lab / investigation support",
    "Pharmacy / medicine support",
    "Staffing levels for your station",
    "Breaks, meals and welfare support",
    "Safety and infection prevention",
    "Overall organisation",
]

CONSULTANT_DESIGNATIONS = [
    ("Consultant Physician", "des_consult"),
    ("Cardiologist", "des_cardio"),
    ("Senior House Officer (SHO)", "des_sho"),
    ("Medical Officer", "des_mo"),
    ("Other", "des_other"),
]

STAFF_ROLES = [
    ("Nursing", "role_nursing"),
    ("Clinical / Medical Officer", "role_clinical"),
    ("Pharmacy", "role_pharmacy"),
    ("Laboratory", "role_lab"),
    ("Registration / Records", "role_reg"),
    ("CHW / VHT", "role_chw"),
    ("Admin / Logistics", "role_admin"),
    ("Volunteer", "role_vol"),
    ("Other", "role_other"),
]

PRIORITY_SERVICES = [
    ("Cardiology", "svc_cardio"),
    ("Hypertension", "svc_htn"),
    ("Diabetes", "svc_dm"),
    ("Women's Health", "svc_wh"),
    ("Men's Health", "svc_mh"),
    ("Children's Health", "svc_ch"),
    ("Eye Care", "svc_eye"),
    ("Mental Health", "svc_mental"),
    ("Cancer Screening", "svc_cancer"),
    ("Elderly Care / Gericare", "svc_elderly"),
]

COLLAB_CAPACITY = [
    ("Clinical services", "cap_clinical"),
    ("Training and mentorship", "cap_train"),
    ("Research", "cap_research"),
    ("Community outreach", "cap_outreach"),
    ("Health education", "cap_educ"),
    ("Programme planning", "cap_plan"),
]


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_run_font(run, size=11, bold=False, color=None, name="Calibri") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading_block(doc: Document, title: str, subtitle: str) -> None:
    """Logo left, camp metadata right — vertically centred as one branding row."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left.text = ""
    if LOGO.exists():
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(2.15))

    right.text = ""
    lines = [
        ("Community Cardiology Camp", 12, True, (28, 43, 42)),
        (SLOGAN, 9, False, (0, 140, 69)),
        (CONTACT_LINE_2, 8.5, False, (90, 106, 106)),
        (CONTACT_LINE_3, 8.5, False, (90, 106, 106)),
        (CONTACT_LINE_1, 8, False, (90, 106, 106)),
    ]
    for i, (text, size, bold, color) in enumerate(lines):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        _set_run_font(r, size=size, bold=bold, color=color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(title)
    _set_run_font(r, size=12, bold=True, color=(0, 107, 53))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    _set_run_font(r, size=10, color=(90, 106, 106))
    doc.add_paragraph()


def _add_para(doc: Document, text: str, *, size=10.5, bold=False, color=(28, 43, 42),
              space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, color=color)


def _add_section(doc: Document, code: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{code}: {title}")
    _set_run_font(r, size=11.5, bold=True, color=(0, 140, 69))


def _add_field_line(doc: Document, label: str, width_hint: str = "_" * 48) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label} ")
    _set_run_font(r, size=10.5, bold=True)
    r2 = p.add_run(width_hint)
    _set_run_font(r2, size=10.5, color=(90, 106, 106))


def _add_date_field(doc: Document, label: str, *, optional_note: str = "") -> None:
    """Printable DD / MM / YYYY date boxes."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    _set_run_font(r, size=10.5, bold=True)
    if optional_note:
        n = p.add_run(f"  {optional_note}")
        _set_run_font(n, size=9, color=(90, 106, 106))

    table = doc.add_table(rows=1, cols=5)
    table.autofit = True
    parts = [("DD", 3), ("/", 1), ("MM", 3), ("/", 1), ("YYYY", 5)]
    for i, (text, _) in enumerate(parts):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text in ("/",):
            run = p.add_run("/")
            _set_run_font(run, size=14, bold=True, color=(0, 140, 69))
        else:
            run = p.add_run(f"  {text}  ")
            _set_run_font(run, size=11, color=(90, 106, 106))
            _set_cell_shading(cell, "F7FBFF")
    hint = doc.add_paragraph()
    hint.paragraph_format.space_after = Pt(8)
    hr = hint.add_run("Format: day / month / year  (example: 18 / 07 / 2026)")
    _set_run_font(hr, size=8.5, color=(90, 106, 106))


def _add_check_options(doc: Document, options: list[str], per_row: int = 2) -> None:
    """Clear printable checkbox grid (not cramped inline text)."""
    # Pad to full rows for a clean table
    cells = list(options)
    while len(cells) % per_row:
        cells.append("")
    rows = len(cells) // per_row
    table = doc.add_table(rows=rows, cols=per_row)
    table.autofit = True
    for r in range(rows):
        for c in range(per_row):
            opt = cells[r * per_row + c]
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if not opt:
                continue
            box = p.add_run("☐")
            _set_run_font(box, size=14, color=(0, 140, 69))
            gap = p.add_run("  ")
            _set_run_font(gap, size=11)
            label = p.add_run(opt)
            _set_run_font(label, size=10.5, color=(28, 43, 42))
            _set_cell_shading(cell, "F3FBF6")
    doc.add_paragraph()


def _add_radio_options(doc: Document, options: list[str], per_row: int = 3) -> None:
    """Printable radio-style choices (circles, not squares)."""
    cells = list(options)
    while len(cells) % per_row:
        cells.append("")
    rows = max(1, len(cells) // per_row)
    table = doc.add_table(rows=rows, cols=per_row)
    table.autofit = True
    for r in range(rows):
        for c in range(per_row):
            opt = cells[r * per_row + c]
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if not opt:
                continue
            circle = p.add_run("☐")
            _set_run_font(circle, size=14, color=(0, 140, 69))
            gap = p.add_run("  ")
            _set_run_font(gap, size=11)
            label = p.add_run(opt)
            _set_run_font(label, size=10.5, bold=True, color=(28, 43, 42))
            _set_cell_shading(cell, "F3FBF6")
    doc.add_paragraph()


def _add_open_prompt(doc: Document, number: str, prompt: str, lines: int = 3) -> None:
    _add_para(doc, f"{number}  {prompt}", size=10.5, bold=True, space_after=3)
    for _ in range(lines):
        _add_para(doc, "_" * 92, size=9, color=(165, 180, 179), space_after=2)


def _add_rating_table(doc: Document, items: list[str]) -> None:
    headers = ["Area", "Excellent\n(5)", "Very Good\n(4)", "Good\n(3)", "Fair\n(2)", "Poor\n(1)"]
    table = doc.add_table(rows=1 + len(items), cols=6)
    table.style = "Table Grid"
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        _set_run_font(r, size=8.5, bold=True, color=(255, 255, 255))
        _set_cell_shading(cell, "008C45")

    for r_idx, item in enumerate(items, 1):
        row = table.rows[r_idx]
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        r = p.add_run(item)
        _set_run_font(r, size=9)
        if r_idx % 2 == 0:
            _set_cell_shading(row.cells[0], "E8F6EE")
        for c in range(1, 6):
            row.cells[c].text = ""
            p = row.cells[c].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run("☐")
            _set_run_font(rr, size=14, color=(0, 140, 69))
            if r_idx % 2 == 0:
                _set_cell_shading(row.cells[c], "E8F6EE")

    doc.add_paragraph()


def _add_footer_block(doc: Document) -> None:
    _add_section(doc, "SUBMISSION", "How to return this form")
    _add_para(
        doc,
        f"Please email your completed questionnaire to {SUBMIT_EMAIL}. "
        "You may type responses in this Word file, or print, fill by hand, and scan. "
        "A fillable PDF is also available: type in the fields, save a filled copy, and email it back. "
        "Your feedback is confidential and used only for programme improvement, "
        "reporting, and partnership development.",
        size=10,
    )
    _add_para(doc, "Contact FairBanks Medical Centre", size=10.5, bold=True, space_after=2)
    _add_para(doc, CONTACT_LINE_1, size=9.5, color=(90, 106, 106), space_after=1)
    _add_para(doc, CONTACT_LINE_2, size=9.5, color=(90, 106, 106), space_after=1)
    _add_para(doc, CONTACT_LINE_3, size=9.5, color=(90, 106, 106), space_after=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Thank you for strengthening community cardiovascular care with FairBanks.")
    _set_run_font(r, size=10, bold=True, color=(0, 140, 69))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"FairBanks Medical Centre  ·  {SLOGAN}")
    _set_run_font(r, size=9, color=(90, 106, 106))


def build_consultant_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _add_heading_block(
        doc,
        "Consultant & Senior House Officer (SHO) Post-Camp Evaluation",
        "FairBanks Reach Programme  ·  Confidential feedback  ·  ~10–12 minutes",
    )

    _add_para(
        doc,
        "Dear Consultant / SHO,",
        size=10.5,
        bold=True,
        space_after=4,
    )
    _add_para(
        doc,
        "On behalf of the Board, Management and FairBanks Medical Centre team, thank you for "
        "your contribution to our Community Cardiology Camp. Your expertise helped community "
        "members access specialist cardiovascular screening, diagnosis and advice.",
        size=10.5,
    )
    _add_para(
        doc,
        "As we plan future outreaches under the FairBanks Reach Programme, please share honest "
        "feedback. Your answers will help improve camp quality and strengthen partnership and "
        "sponsorship proposals for long-term community cardiovascular care.",
        size=10.5,
    )
    _add_para(
        doc,
        "Tip: You can edit this Word file on a computer, or print and fill by hand. "
        "A fillable PDF version is also available for soft-copy completion.",
        size=9.5,
        color=(8, 78, 76),
    )

    _add_section(doc, "SECTION A", "General information")
    _add_field_line(doc, "Name (optional):")
    _add_field_line(doc, "Institution / hospital:")
    _add_date_field(doc, "Camp date:")
    _add_field_line(doc, "Camp venue / community:", "_" * 36)
    _add_para(doc, "Designation (tick one):", size=10.5, bold=True, space_after=3)
    _add_check_options(doc, [d[0] for d in CONSULTANT_DESIGNATIONS], per_row=2)
    _add_field_line(doc, "If Other, please specify:")
    _add_field_line(doc, "Years of practice (optional):", "_" * 20)
    _add_date_field(doc, "Date you completed this form:")

    _add_section(doc, "SECTION B", "Overall rating")
    _add_para(
        doc,
        "Please rate each area. Scale: 5 = Excellent · 4 = Very Good · 3 = Good · 2 = Fair · 1 = Poor. "
        "Tick one box per row.",
        size=9.5,
        color=(90, 106, 106),
    )
    _add_rating_table(doc, CONSULTANT_RATINGS)

    _add_section(doc, "SECTION C", "Professional feedback")
    _add_open_prompt(doc, "1.", "How would you describe your overall experience at this camp?")
    _add_open_prompt(doc, "2.", "What impressed you most (clinical quality, organisation, teamwork, community response)?")
    _add_open_prompt(doc, "3.", "What challenges did you face (patient volume, investigations, medicines, space, flow)?")
    _add_open_prompt(doc, "4.", "What should we change to improve the next cardiology camp?")
    _add_open_prompt(
        doc,
        "5.",
        "Were there any clinical patterns or high-need findings we should plan for next time?",
        lines=2,
    )

    _add_section(doc, "SECTION D", "Patient follow-up and referral")
    _add_para(
        doc,
        "Many patients need ongoing cardiovascular care after the camp. Please advise how "
        "FairBanks Medical Centre can:",
        size=10,
    )
    _add_open_prompt(doc, "a)", "Improve post-camp follow-up of patients identified with CVD or risk factors.")
    _add_open_prompt(doc, "b)", "Strengthen referral pathways for complex cardiac cases.")
    _add_open_prompt(doc, "c)", "Improve medication adherence and continuity of care.")
    _add_open_prompt(
        doc,
        "d)",
        "Use CHWs, VHTs and digital tools to monitor patients in the community after the camp.",
    )

    _add_section(doc, "SECTION E", "Programme sustainability")
    _add_open_prompt(
        doc,
        "1.",
        "Which strengths of this programme would appeal most to partners and sponsors?",
    )
    _add_para(doc, "2. Which services should FairBanks prioritise in future community outreaches? (tick all that apply)", size=10.5, bold=True, space_after=3)
    _add_check_options(doc, [s[0] for s in PRIORITY_SERVICES] + ["Other"], per_row=3)
    _add_field_line(doc, "If Other, please specify:")
    _add_open_prompt(
        doc,
        "3.",
        "What would make a community cardiovascular programme sustainable over 2–3 years?",
    )

    _add_section(doc, "SECTION F", "Future collaboration")
    _add_para(doc, "Would you be willing to join future FairBanks Reach Programmes?", size=10.5, bold=True, space_after=3)
    _add_radio_options(doc, ["Yes", "No", "Maybe"], per_row=3)
    _add_para(doc, "If yes or maybe, in what capacity? (tick all that apply)", size=10.5, bold=True, space_after=3)
    _add_check_options(doc, [c[0] for c in COLLAB_CAPACITY] + ["Other"], per_row=2)
    _add_field_line(doc, "If Other, please specify:")
    _add_field_line(doc, "Preferred contact (phone or email, optional):")

    _add_section(doc, "SECTION G", "Closing remarks")
    _add_open_prompt(
        doc,
        "",
        "Any additional comments, lessons learned or suggestions for the Board and Management:",
        lines=4,
    )

    _add_footer_block(doc)
    doc.save(path)


def build_staff_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _add_heading_block(
        doc,
        "Internal Staff Post-Camp Evaluation",
        "FairBanks Reach Programme  ·  Confidential internal feedback  ·  ~8–10 minutes",
    )

    _add_para(doc, "Dear Colleague,", size=10.5, bold=True, space_after=4)
    _add_para(
        doc,
        "Thank you for serving at the Community Cardiology Camp. Your work at registration, "
        "clinical stations, pharmacy, laboratory, logistics and community support made the "
        "outreach possible.",
        size=10.5,
    )
    _add_para(
        doc,
        "This internal form collects practical feedback so we can improve the next FairBanks "
        "Reach Programme camp. Please be frank — responses are for internal learning and "
        "planning, not individual performance appraisal.",
        size=10.5,
    )
    _add_para(
        doc,
        "Tip: Edit this Word file on a computer, or print and fill by hand. "
        "A fillable PDF version is also available for soft-copy completion.",
        size=9.5,
        color=(8, 78, 76),
    )

    _add_section(doc, "SECTION A", "About you and your role")
    _add_field_line(doc, "Name (optional):")
    _add_field_line(doc, "Department / unit:")
    _add_date_field(doc, "Camp date:")
    _add_field_line(doc, "Camp venue / community:", "_" * 36)
    _add_para(doc, "Main role during the camp (tick one):", size=10.5, bold=True, space_after=3)
    _add_check_options(doc, [r[0] for r in STAFF_ROLES], per_row=2)
    _add_field_line(doc, "If Other, please specify:")
    _add_para(doc, "Was this your first FairBanks community camp?", size=10.5, bold=True, space_after=3)
    _add_radio_options(doc, ["Yes", "No"], per_row=2)
    _add_date_field(doc, "Date you completed this form:")

    _add_section(doc, "SECTION B", "Operations rating")
    _add_para(
        doc,
        "Rate each area from your station's point of view. "
        "Scale: 5 = Excellent · 4 = Very Good · 3 = Good · 2 = Fair · 1 = Poor. "
        "Tick one box per row.",
        size=9.5,
        color=(90, 106, 106),
    )
    _add_rating_table(doc, STAFF_RATINGS)

    _add_section(doc, "SECTION C", "What worked and what did not")
    _add_open_prompt(doc, "1.", "What worked well at your station or team?")
    _add_open_prompt(doc, "2.", "What caused delays, confusion or bottlenecks?")
    _add_open_prompt(doc, "3.", "Were supplies, forms, devices or medicines missing or short? Please list.")
    _add_open_prompt(doc, "4.", "How clear was your briefing and task assignment before and during the camp?")
    _add_open_prompt(
        doc,
        "5.",
        "How would you improve patient registration, triage and flow next time?",
    )

    _add_section(doc, "SECTION D", "Teamwork, welfare and safety")
    _add_open_prompt(doc, "1.", "How was teamwork between clinical, support and community teams?")
    _add_open_prompt(doc, "2.", "Did you have enough staff at your station? If not, what support was missing?")
    _add_open_prompt(doc, "3.", "Any safety, infection-prevention or patient-privacy concerns?")
    _add_open_prompt(doc, "4.", "Comments on breaks, meals, transport or staff welfare:")

    _add_section(doc, "SECTION E", "Community and continuity")
    _add_open_prompt(
        doc,
        "1.",
        "What community needs or patient patterns stood out (e.g. hypertension, diabetes, late presentation)?",
    )
    _add_open_prompt(
        doc,
        "2.",
        "How can CHWs / VHTs and FairBanks better support follow-up after the camp?",
    )
    _add_para(doc, "3. Which services should future camps prioritise? (tick all that apply)", size=10.5, bold=True, space_after=3)
    _add_check_options(doc, [s[0] for s in PRIORITY_SERVICES] + ["Other"], per_row=3)
    _add_field_line(doc, "If Other, please specify:")

    _add_section(doc, "SECTION F", "Future camps")
    _add_para(doc, "Would you like to serve again at future FairBanks Reach camps?", size=10.5, bold=True, space_after=3)
    _add_radio_options(doc, ["Yes", "No", "Maybe"], per_row=3)
    _add_open_prompt(doc, "1.", "What training or tools would help you perform better at the next camp?")
    _add_open_prompt(doc, "2.", "One change that would most improve the next camp:")

    _add_section(doc, "SECTION G", "Closing remarks")
    _add_open_prompt(
        doc,
        "",
        "Any other comments for Management (logistics, partnerships, recognition, ideas):",
        lines=4,
    )

    _add_footer_block(doc)
    doc.save(path)


# ---------------------------------------------------------------------------
# PDF helpers (fillable)
# ---------------------------------------------------------------------------

def finalize_fillable_pdf(path: Path) -> None:
    """Ensure AcroForm fields stay typeable, visible, and saveable across viewers."""
    doc = fitz.open(path)
    cat = doc.pdf_catalog()
    acro = doc.xref_get_key(cat, "AcroForm")
    if acro[0] == "xref":
        af_xref = int(acro[1].split()[0])
        doc.xref_set_key(af_xref, "NeedAppearances", "true")

    # Clear accidental default selections and ReadOnly flags without rewriting
    # widget appearances (calling Widget.update() on radios can mark them On).
    for page in doc:
        for widget in page.widgets() or []:
            flags = widget.field_flags or 0
            if flags & 1:  # ReadOnly
                widget.field_flags = flags & ~1
            # Radios/checkboxes must start unchecked for a blank questionnaire
            if widget.field_type_string in ("RadioButton", "CheckBox"):
                try:
                    if widget.field_value not in (None, "", "Off"):
                        widget.field_value = "Off"
                except Exception:
                    pass

    tmp = path.with_suffix(".tmp.pdf")
    doc.save(tmp, garbage=3, deflate=True, clean=True)
    doc.close()
    tmp.replace(path)


class FormPDF:
    def __init__(self, path: Path, doc_title: str):
        self.path = path
        self.c = canvas.Canvas(str(path), pagesize=A4)
        self.c.setTitle(doc_title)
        self.c.setAuthor("FairBanks Medical Centre")
        self.c.setSubject("Fillable evaluation questionnaire - type, save, and email back")
        self.y = PAGE_H - MARGIN_T
        self.page = 1
        self._field_n = 0
        # Help viewers regenerate field appearances when typing
        try:
            self.c.acroForm.needAppearances = True
        except Exception:
            pass

    def new_name(self, prefix: str) -> str:
        self._field_n += 1
        return f"{prefix}_{self._field_n}"

    def ensure(self, need: float) -> None:
        # Keep clear of footer band so fields never collide with contact lines
        if self.y - need < MARGIN_B + 8 * mm:
            self.footer()
            self.c.showPage()
            self.page += 1
            self.y = PAGE_H - MARGIN_T
            self.header_mini()

    def footer(self) -> None:
        self.c.setStrokeColor(LINE)
        self.c.setLineWidth(0.4)
        self.c.line(MARGIN_L, 16 * mm, PAGE_W - MARGIN_R, 16 * mm)
        self.c.setFillColor(MUTED)
        self.c.setFont("Helvetica", 6.5)
        self.c.drawString(MARGIN_L, 11.5 * mm, f"{CONTACT_LINE_2}  |  {SUBMIT_EMAIL}")
        self.c.drawString(MARGIN_L, 7.5 * mm, f"{WEBSITE}  ·  Confidential  ·  Type, Save As, then email filled copy")
        self.c.drawRightString(PAGE_W - MARGIN_R, 7.5 * mm, f"Page {self.page}")

    def header_mini(self) -> None:
        self.c.setFillColor(BRAND)
        self.c.rect(0, PAGE_H - 8 * mm, PAGE_W, 8 * mm, fill=1, stroke=0)
        self.c.setFillColor(white)
        self.c.setFont("Helvetica-Bold", 8)
        self.c.drawString(MARGIN_L, PAGE_H - 5.5 * mm, "FairBanks  ·  Community Cardiology Camp Evaluation")
        self.c.setFont("Helvetica", 7)
        self.c.drawRightString(PAGE_W - MARGIN_R, PAGE_H - 5.5 * mm, SUBMIT_EMAIL)
        self.y = PAGE_H - 14 * mm

    def header_full(self, title: str, audience: str) -> None:
        """Brand band: logo and metadata vertically centred and left-aligned as one unit."""
        band_h = 38 * mm
        band_bottom = PAGE_H - band_h

        self.c.setFillColor(BRAND)
        self.c.rect(0, band_bottom, PAGE_W, band_h, fill=1, stroke=0)
        self.c.setFillColor(ACCENT)
        self.c.rect(0, band_bottom - 1.6 * mm, PAGE_W, 1.6 * mm, fill=1, stroke=0)

        # Logo from true aspect ratio (269 x 101) — plate optically centred in band
        logo_h = 17.5 * mm
        logo_w = logo_h * (269 / 101)
        plate_pad = 2.0 * mm
        plate_w = logo_w + 2 * plate_pad
        plate_h = logo_h + 2 * plate_pad
        gap = 4.5 * mm

        plate_x = MARGIN_L
        # Slight downward nudge so the plate reads centred against the orange strip
        plate_y = band_bottom + (band_h - plate_h) / 2 - 0.8 * mm
        logo_x = plate_x + plate_pad
        logo_y = plate_y + plate_pad

        if LOGO.exists():
            self.c.setFillColor(white)
            self.c.roundRect(plate_x, plate_y, plate_w, plate_h, 3, fill=1, stroke=0)
            self.c.drawImage(
                str(LOGO),
                logo_x,
                logo_y,
                width=logo_w,
                height=logo_h,
                preserveAspectRatio=False,
                mask="auto",
            )
            text_x = plate_x + plate_w + gap
        else:
            text_x = MARGIN_L

        # Metadata shares the plate's top/bottom — lines evenly spaced between
        meta_texts = [
            ("Helvetica-Bold", 10.5, "Community Cardiology Camp"),
            ("Helvetica-Oblique", 7.2, SLOGAN),
            ("Helvetica", 6.5, CONTACT_LINE_2),
            ("Helvetica", 6.5, CONTACT_LINE_3),
            ("Helvetica", 6.1, CONTACT_LINE_1),
        ]
        first_y = plate_y + plate_h - 3.4 * mm
        last_y = plate_y + 2.2 * mm
        step = (first_y - last_y) / (len(meta_texts) - 1)

        self.c.setFillColor(white)
        for i, (font, size, text) in enumerate(meta_texts):
            self.c.setFont(font, size)
            self.c.drawString(text_x, first_y - i * step, text)

        self.y = band_bottom - 1.6 * mm - 6 * mm
        self.c.setFillColor(BRAND_DARK)
        self.c.setFont("Helvetica-Bold", 11)
        self.c.drawCentredString(PAGE_W / 2, self.y, title)
        self.y -= 5 * mm
        self.c.setFillColor(MUTED)
        self.c.setFont("Helvetica", 8.5)
        self.c.drawCentredString(PAGE_W / 2, self.y, audience)
        self.y -= 6 * mm

        # Soft-copy tip
        box_h = 13 * mm
        self.ensure(box_h + 4 * mm)
        self.c.setFillColor(WARN)
        self.c.setStrokeColor(ACCENT)
        self.c.roundRect(MARGIN_L, self.y - box_h + 2 * mm, CONTENT_W, box_h, 3, fill=1, stroke=1)
        self.c.setFillColor(INK)
        self.c.setFont("Helvetica-Bold", 8)
        self.c.drawString(MARGIN_L + 3 * mm, self.y - 1.5 * mm, "Fillable soft copy - type, save, and send back")
        self.c.setFont("Helvetica", 7.5)
        self.c.drawString(
            MARGIN_L + 3 * mm,
            self.y - 5.5 * mm,
            "1) Click any blue-bordered field to type   2) Tick boxes (a check appears)   3) File > Save As (keep as PDF)",
        )
        self.c.drawString(
            MARGIN_L + 3 * mm,
            self.y - 9.5 * mm,
            f"4) Email the filled PDF to {SUBMIT_EMAIL}   ·   Works in Adobe Reader, Edge, Chrome, and Foxit",
        )
        self.y -= box_h + 3 * mm

    def intro(self, paragraphs: list[str]) -> None:
        for text in paragraphs:
            self.ensure(12 * mm)
            self.c.setFillColor(INK)
            self.c.setFont("Helvetica", 8.5)
            self._wrap(text, MARGIN_L, CONTENT_W, leading=10)
            self.y -= 2 * mm

    def _wrap(self, text: str, x: float, width: float, leading: float = 10, font="Helvetica", size=8.5) -> None:
        self.c.setFont(font, size)
        words = text.split()
        line = ""
        for w in words:
            trial = f"{line} {w}".strip()
            if self.c.stringWidth(trial, font, size) <= width:
                line = trial
            else:
                self.ensure(leading + 1)
                self.c.drawString(x, self.y, line)
                self.y -= leading
                line = w
        if line:
            self.ensure(leading + 1)
            self.c.drawString(x, self.y, line)
            self.y -= leading

    def section(self, code: str, title: str) -> None:
        self.ensure(14 * mm)
        self.y -= 3 * mm
        bar_h = 7.5 * mm
        self.c.setFillColor(LIGHT)
        self.c.roundRect(MARGIN_L, self.y - bar_h + 2 * mm, CONTENT_W, bar_h, 2, fill=1, stroke=0)
        self.c.setFillColor(BRAND)
        self.c.setFont("Helvetica-Bold", 9.5)
        self.c.drawString(MARGIN_L + 2.5 * mm, self.y - 2.2 * mm, f"{code}  ·  {title}")
        self.y -= bar_h + 3 * mm

    def label(self, text: str, *, bold: bool = True, size: float = 8.5) -> None:
        self.ensure(7 * mm)
        self.c.setFillColor(INK)
        self.c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        self.c.drawString(MARGIN_L, self.y, text)
        self.y -= 5.5 * mm

    def text_line(self, label: str, name: str, field_w: float | None = None) -> None:
        """Label on its own line, field below - avoids label/field collision."""
        self.ensure(18 * mm)
        self.c.setFillColor(INK)
        self.c.setFont("Helvetica-Bold", 8.5)
        self.c.drawString(MARGIN_L, self.y, label)
        self.y -= 5.5 * mm
        fw = field_w or CONTENT_W
        field_h = 5.5 * mm
        field_bottom = self.y - field_h
        self.c.acroForm.textfield(
            name=name,
            tooltip=label,
            x=MARGIN_L,
            y=field_bottom,
            width=fw,
            height=field_h,
            borderWidth=0.8,
            borderColor=HexColor("#4AA3FF"),
            fillColor=HexColor("#F7FBFF"),
            textColor=INK,
            forceBorder=True,
            fontSize=9,
            fontName="Helvetica",
            fieldFlags="",
            maxlen=500,
        )
        self.y = field_bottom - 8.5 * mm

    def date_field(self, label: str, prefix: str, *, hint: str = "") -> None:
        """DD / MM / YYYY with labels under boxes so nothing overlaps."""
        self.ensure(26 * mm)
        self.c.setFillColor(INK)
        self.c.setFont("Helvetica-Bold", 8.5)
        self.c.drawString(MARGIN_L, self.y, label)
        if hint:
            self.c.setFont("Helvetica", 7.5)
            self.c.setFillColor(MUTED)
            self.c.drawString(
                MARGIN_L + self.c.stringWidth(label, "Helvetica-Bold", 8.5) + 2.5 * mm,
                self.y,
                hint,
            )
        self.y -= 5.5 * mm

        day_w = 16 * mm
        mon_w = 16 * mm
        year_w = 26 * mm
        slash_w = 6 * mm
        box_h = 7 * mm
        parts = [
            ("day", day_w, "Day", 2),
            ("month", mon_w, "Month", 2),
            ("year", year_w, "Year", 4),
        ]

        x = MARGIN_L
        box_y = self.y - box_h
        for i, (part, width, placeholder, maxlen) in enumerate(parts):
            self.c.acroForm.textfield(
                name=f"{prefix}_{part}",
                tooltip=f"{label} ({placeholder})",
                value="",
                x=x,
                y=box_y,
                width=width,
                height=box_h,
                borderWidth=0.9,
                borderColor=BRAND_DARK,
                fillColor=HexColor("#F7FBFF"),
                textColor=INK,
                forceBorder=True,
                fontSize=11,
                fontName="Helvetica-Bold",
                fieldFlags="",
                maxlen=maxlen,
            )
            x += width
            if i < 2:
                self.c.setFillColor(BRAND)
                self.c.setFont("Helvetica-Bold", 12)
                self.c.drawCentredString(x + slash_w / 2, box_y + box_h / 2 - 3.5, "/")
                x += slash_w

        self.c.setFillColor(MUTED)
        self.c.setFont("Helvetica", 7)
        self.c.drawString(x + 5 * mm, box_y + box_h / 2 - 2.5, "Example: 18 / 07 / 2026")

        # Captions under boxes (never overlap the inputs or the next field)
        x = MARGIN_L
        cap_y = box_y - 3.5 * mm
        self.c.setFont("Helvetica", 6.5)
        for i, (part, width, placeholder, maxlen) in enumerate(parts):
            self.c.drawCentredString(x + width / 2, cap_y, placeholder)
            x += width
            if i < 2:
                x += slash_w
        self.y = cap_y - 7 * mm

    def multiline(self, prompt: str, name: str, height: float = 18 * mm) -> None:
        self.ensure(height + 14 * mm)
        self.c.setFillColor(INK)
        self.c.setFont("Helvetica-Bold", 8.5)
        self._wrap(prompt, MARGIN_L, CONTENT_W, leading=10, font="Helvetica-Bold", size=8.5)
        self.y -= 3 * mm
        self.ensure(height + 6 * mm)
        field_bottom = self.y - height
        self.c.acroForm.textfield(
            name=name,
            tooltip=prompt[:120],
            x=MARGIN_L,
            y=field_bottom,
            width=CONTENT_W,
            height=height,
            borderWidth=0.8,
            borderColor=HexColor("#4AA3FF"),
            fillColor=HexColor("#F7FBFF"),
            textColor=INK,
            forceBorder=True,
            fontSize=9,
            fontName="Helvetica",
            fieldFlags="multiline",
            maxlen=4000,
        )
        self.y = field_bottom - 7.5 * mm

    def check_row(self, options: list[tuple[str, str]], cols: int = 2) -> None:
        """Checkboxes fully inside chips with clear padding."""
        col_w = CONTENT_W / cols
        box = 12
        pad = 2.2 * mm
        for i, (label, name) in enumerate(options):
            if i % cols == 0:
                self.ensure(12 * mm)
                row_y = self.y
            col = i % cols
            x = MARGIN_L + col * col_w
            chip_w = col_w - 2.5 * mm
            chip_h = 7.6 * mm
            chip_bottom = row_y - chip_h
            self.c.setFillColor(HexColor("#F3FBF6"))
            self.c.setStrokeColor(HexColor("#B7DFC6"))
            self.c.setLineWidth(0.7)
            self.c.roundRect(x, chip_bottom, chip_w, chip_h, 2.2, fill=1, stroke=1)
            cb_y = chip_bottom + (chip_h - box) / 2
            self.c.acroForm.checkbox(
                name=name,
                tooltip=f"Select: {label}",
                x=x + pad,
                y=cb_y,
                size=box,
                buttonStyle="check",
                shape="square",
                borderWidth=1.4,
                borderColor=BRAND_DARK,
                fillColor=HexColor("#FFFFFF"),
                textColor=BRAND_DARK,
                forceBorder=True,
                checked=False,
                fieldFlags="",
            )
            self.c.setFillColor(INK)
            self.c.setFont("Helvetica", 8.5)
            self.c.drawString(x + pad + box + 2.8 * mm, cb_y + 2.5, label)
            if col == cols - 1 or i == len(options) - 1:
                self.y = chip_bottom - 4.5 * mm

    def radio_choice_row(self, group: str, opts: list[tuple[str, str]]) -> None:
        """Yes/No/Maybe as padded checkbox chips (not stretched full width)."""
        box = 12
        self.ensure(12 * mm)
        chip_w = 38 * mm
        chip_h = 8 * mm
        gap = 4 * mm
        x = MARGIN_L
        chip_bottom = self.y - chip_h
        for label, val in opts:
            self.c.setFillColor(HexColor("#F3FBF6"))
            self.c.setStrokeColor(HexColor("#8FCFAB"))
            self.c.setLineWidth(0.8)
            self.c.roundRect(x, chip_bottom, chip_w, chip_h, 2.5, fill=1, stroke=1)
            cb_y = chip_bottom + (chip_h - box) / 2
            self.c.acroForm.checkbox(
                name=f"{group}_{val}",
                tooltip=f"Select {label} (tick one)",
                x=x + 2.5 * mm,
                y=cb_y,
                size=box,
                buttonStyle="check",
                shape="square",
                borderWidth=1.6,
                borderColor=BRAND_DARK,
                fillColor=HexColor("#FFFFFF"),
                textColor=BRAND_DARK,
                forceBorder=True,
                checked=False,
                fieldFlags="",
            )
            self.c.setFillColor(INK)
            self.c.setFont("Helvetica-Bold", 9)
            self.c.drawString(x + 2.5 * mm + box + 2.5 * mm, cb_y + 2.5, label)
            x += chip_w + gap
        self.y = chip_bottom - 4.5 * mm

    def radio_yes_no_maybe(self, group: str) -> None:
        self.radio_choice_row(group, [("Yes", "Yes"), ("No", "No"), ("Maybe", "Maybe")])

    def rating_table(self, items: list[str], prefix: str) -> None:
        headers = ["Area", "5", "4", "3", "2", "1"]
        col_w = 12.5 * mm
        col_area = CONTENT_W - 5 * col_w
        row_h = 7.6 * mm
        head_h = 7.2 * mm
        box = 11

        self.ensure(head_h + row_h * min(3, len(items)) + 6 * mm)
        self.c.setFillColor(MUTED)
        self.c.setFont("Helvetica", 7)
        self.c.drawString(
            MARGIN_L,
            self.y,
            "Scale: 5 Excellent · 4 Very Good · 3 Good · 2 Fair · 1 Poor  "
            "(tick one checkbox per row)",
        )
        self.y -= 4.5 * mm

        self.ensure(head_h + 2 * mm)
        head_bottom = self.y - head_h
        self.c.setFillColor(BRAND)
        self.c.rect(MARGIN_L, head_bottom, CONTENT_W, head_h, fill=1, stroke=0)
        self.c.setFillColor(white)
        self.c.setFont("Helvetica-Bold", 8.5)
        self.c.drawString(MARGIN_L + 2.5 * mm, head_bottom + 2.4 * mm, headers[0])
        for i, h in enumerate(headers[1:]):
            cx = MARGIN_L + col_area + i * col_w + col_w / 2
            self.c.drawCentredString(cx, head_bottom + 2.4 * mm, h)
        self.y = head_bottom - 0.5 * mm

        for idx, item in enumerate(items):
            self.ensure(row_h + 2 * mm)
            row_bottom = self.y - row_h
            if idx % 2 == 0:
                self.c.setFillColor(LIGHT)
                self.c.rect(MARGIN_L, row_bottom, CONTENT_W, row_h, fill=1, stroke=0)
            self.c.setStrokeColor(HexColor("#C9E6D5"))
            self.c.setLineWidth(0.4)
            for i in range(6):
                gx = MARGIN_L + col_area + i * col_w
                self.c.line(gx, row_bottom + 0.8, gx, row_bottom + row_h - 0.8)
            self.c.setFillColor(INK)
            self.c.setFont("Helvetica", 8)
            self.c.drawString(MARGIN_L + 2.5 * mm, row_bottom + row_h / 2 - 2.2, item)
            for score in range(5, 0, -1):
                i = 5 - score
                cx = MARGIN_L + col_area + i * col_w + (col_w - box) / 2
                cy = row_bottom + (row_h - box) / 2
                self.c.acroForm.checkbox(
                    name=f"{prefix}_r{idx}_{score}",
                    tooltip=f"{item}: score {score} (tick one per row)",
                    x=cx,
                    y=cy,
                    size=box,
                    buttonStyle="check",
                    shape="square",
                    borderWidth=1.5,
                    borderColor=BRAND_DARK,
                    fillColor=HexColor("#FFFFFF"),
                    textColor=BRAND_DARK,
                    forceBorder=True,
                    checked=False,
                    fieldFlags="",
                )
            self.y = row_bottom
        self.y -= 3 * mm

    def save(self) -> None:
        self.footer()
        self.c.save()
        finalize_fillable_pdf(self.path)


def build_consultant_pdf(path: Path) -> None:
    pdf = FormPDF(path, "FairBanks Consultant/SHO Cardiology Camp Evaluation")
    pdf.header_full(
        "Consultant & SHO Post-Camp Evaluation Questionnaire",
        "FairBanks Reach Programme  ·  Confidential  ·  About 10–12 minutes",
    )
    pdf.intro(
        [
            "Dear Consultant / SHO, thank you for supporting our Community Cardiology Camp. "
            "Your expertise helped community members access specialist cardiovascular screening, "
            "diagnosis and medical advice.",
            "Please share honest feedback to improve future FairBanks Reach Programme camps and "
            "strengthen partnership and sponsorship proposals for long-term community cardiovascular care.",
        ]
    )

    pdf.section("SECTION A", "General information")
    pdf.text_line("Name (optional):", "c_name")
    pdf.text_line("Institution / hospital:", "c_institution")
    pdf.date_field("Camp date:", "c_camp")
    pdf.text_line("Camp venue / community:", "c_venue")
    pdf.label("Designation (select all that apply):")
    pdf.check_row(CONSULTANT_DESIGNATIONS, cols=2)
    pdf.text_line("If Other, specify:", "c_des_other")
    pdf.text_line("Years of practice (optional):", "c_years", field_w=30 * mm)
    pdf.date_field("Date you completed this form:", "c_completed")

    pdf.section("SECTION B", "Overall rating")
    pdf.rating_table(CONSULTANT_RATINGS, "c")

    pdf.section("SECTION C", "Professional feedback")
    pdf.multiline("1. How would you describe your overall experience at this camp?", "c_exp", height=12 * mm)
    pdf.multiline("2. What impressed you most?", "c_highlights", height=12 * mm)
    pdf.multiline("3. What challenges did you face?", "c_challenges", height=12 * mm)
    pdf.multiline("4. What should we change for the next cardiology camp?", "c_recommend", height=12 * mm)
    pdf.multiline(
        "5. Any clinical patterns or high-need findings to plan for next time?",
        "c_patterns",
        height=11 * mm,
    )

    pdf.section("SECTION D", "Patient follow-up and referral")
    pdf.multiline("a) How can FairBanks improve post-camp follow-up for CVD / risk patients?", "c_fu_a", height=12 * mm)
    pdf.multiline("b) How can we strengthen referral pathways for complex cardiac cases?", "c_fu_b", height=12 * mm)
    pdf.multiline("c) How can we improve medication adherence and continuity of care?", "c_fu_c", height=12 * mm)
    pdf.multiline(
        "d) How can CHWs, VHTs and digital tools help monitor patients after the camp?",
        "c_fu_d",
        height=12 * mm,
    )

    pdf.section("SECTION E", "Programme sustainability")
    pdf.multiline("1. Which strengths would appeal most to partners and sponsors?", "c_strengths", height=12 * mm)
    pdf.label("2. Priority services for future outreaches (select all that apply):")
    pdf.check_row(PRIORITY_SERVICES + [("Other", "svc_other")], cols=3)
    pdf.text_line("If Other, specify:", "c_svc_other_txt")
    pdf.multiline(
        "3. What would make a community cardiovascular programme sustainable over 2–3 years?",
        "c_sustain",
        height=12 * mm,
    )

    pdf.section("SECTION F", "Future collaboration")
    pdf.label("Would you join future FairBanks Reach Programmes?")
    pdf.radio_yes_no_maybe("c_future")
    pdf.label("If yes or maybe, in what capacity? (select all that apply)")
    pdf.check_row(COLLAB_CAPACITY + [("Other", "cap_other")], cols=2)
    pdf.text_line("If Other, specify:", "c_cap_other_txt")
    pdf.text_line("Preferred contact (optional):", "c_contact")

    pdf.section("SECTION G", "Closing remarks")
    pdf.multiline(
        "Additional comments, lessons learned or suggestions for the Board and Management:",
        "c_closing",
        height=22 * mm,
    )

    # Keep thank-you on the same page as closing remarks when possible
    if pdf.y < MARGIN_B + 16 * mm:
        pdf.footer()
        pdf.c.showPage()
        pdf.page += 1
        pdf.y = PAGE_H - MARGIN_T
        pdf.header_mini()
    pdf.c.setFillColor(BRAND)
    pdf.c.setFont("Helvetica-Bold", 9)
    pdf.c.drawCentredString(PAGE_W / 2, pdf.y, "Thank you for strengthening community cardiovascular care.")
    pdf.y -= 5 * mm
    pdf.c.setFillColor(MUTED)
    pdf.c.setFont("Helvetica", 8)
    pdf.c.drawCentredString(PAGE_W / 2, pdf.y, f"Submit to {SUBMIT_EMAIL}  ·  {SLOGAN}")

    pdf.save()


def build_staff_pdf(path: Path) -> None:
    pdf = FormPDF(path, "FairBanks Internal Staff Cardiology Camp Evaluation")
    pdf.header_full(
        "Internal Staff Post-Camp Evaluation Questionnaire",
        "FairBanks Reach Programme  ·  Confidential internal use  ·  About 8–10 minutes",
    )
    pdf.intro(
        [
            "Dear Colleague, thank you for serving at the Community Cardiology Camp. Your work at "
            "registration, clinical stations, pharmacy, laboratory, logistics and community support "
            "made the outreach possible.",
            "This internal form is for learning and planning - not individual performance appraisal. "
            "Please be frank so we can improve the next FairBanks Reach Programme camp.",
        ]
    )

    pdf.section("SECTION A", "About you and your role")
    pdf.text_line("Name (optional):", "s_name")
    pdf.text_line("Department / unit:", "s_dept")
    pdf.date_field("Camp date:", "s_camp")
    pdf.text_line("Camp venue / community:", "s_venue")
    pdf.label("Main role during the camp (select one or more):")
    pdf.check_row(STAFF_ROLES, cols=2)
    pdf.text_line("If Other, specify:", "s_role_other")
    pdf.label("Was this your first FairBanks community camp?")
    pdf.radio_choice_row("s_first", [("Yes", "Yes"), ("No", "No")])
    pdf.date_field("Date you completed this form:", "s_completed")

    pdf.section("SECTION B", "Operations rating")
    pdf.rating_table(STAFF_RATINGS, "s")

    pdf.section("SECTION C", "What worked and what did not")
    h = 11 * mm
    pdf.multiline("1. What worked well at your station or team?", "s_worked", height=h)
    pdf.multiline("2. What caused delays, confusion or bottlenecks?", "s_delays", height=h)
    pdf.multiline("3. Were supplies, forms, devices or medicines missing or short? List them.", "s_supplies", height=h)
    pdf.multiline("4. How clear was briefing and task assignment before and during the camp?", "s_briefing", height=h)
    pdf.multiline("5. How would you improve registration, triage and patient flow next time?", "s_flow", height=h)

    pdf.section("SECTION D", "Teamwork, welfare and safety")
    pdf.multiline("1. How was teamwork between clinical, support and community teams?", "s_team", height=h)
    pdf.multiline("2. Did you have enough staff at your station? If not, what was missing?", "s_staffing", height=h)
    pdf.multiline("3. Any safety, infection-prevention or patient-privacy concerns?", "s_safety", height=h)
    pdf.multiline("4. Comments on breaks, meals, transport or staff welfare:", "s_welfare", height=h)

    pdf.section("SECTION E", "Community and continuity")
    pdf.multiline(
        "1. What community needs or patient patterns stood out?",
        "s_patterns",
        height=h,
    )
    pdf.multiline(
        "2. How can CHWs / VHTs and FairBanks better support follow-up after the camp?",
        "s_followup",
        height=h,
    )
    pdf.label("3. Priority services for future camps (select all that apply):")
    pdf.check_row(PRIORITY_SERVICES + [("Other", "s_svc_other")], cols=3)
    pdf.text_line("If Other, specify:", "s_svc_other_txt")

    pdf.section("SECTION F", "Future camps")
    pdf.label("Would you like to serve again at future FairBanks Reach camps?")
    pdf.radio_yes_no_maybe("s_future")
    pdf.multiline("1. What training or tools would help you at the next camp?", "s_training", height=h)
    pdf.multiline("2. One change that would most improve the next camp:", "s_one_change", height=h)

    pdf.section("SECTION G", "Closing remarks")
    pdf.multiline(
        "Any other comments for Management (logistics, partnerships, recognition, ideas):",
        "s_closing",
        height=18 * mm,
    )

    if pdf.y < MARGIN_B + 16 * mm:
        pdf.footer()
        pdf.c.showPage()
        pdf.page += 1
        pdf.y = PAGE_H - MARGIN_T
        pdf.header_mini()
    pdf.c.setFillColor(BRAND)
    pdf.c.setFont("Helvetica-Bold", 9)
    pdf.c.drawCentredString(PAGE_W / 2, pdf.y, "Thank you - your feedback improves every FairBanks Reach camp.")
    pdf.y -= 5 * mm
    pdf.c.setFillColor(MUTED)
    pdf.c.setFont("Helvetica", 8)
    pdf.c.drawCentredString(PAGE_W / 2, pdf.y, f"Submit to {SUBMIT_EMAIL}  ·  {SLOGAN}")

    pdf.save()


def main() -> None:
    # Clean leftover from earlier typo helper if present — rebuild all outputs
    files = {
        "consultant_sho_evaluation.docx": build_consultant_docx,
        "internal_staff_evaluation.docx": build_staff_docx,
        "consultant_sho_evaluation_fillable.pdf": build_consultant_pdf,
        "internal_staff_evaluation_fillable.pdf": build_staff_pdf,
    }
    for name, builder in files.items():
        out = OUT_DIR / name
        builder(out)
        print(f"Wrote {out}")


if __name__ == "__main__":
    main()
