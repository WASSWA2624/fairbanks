"""Generate a visually polished CV for Tyra Rebecca Nalukwago (DOCX + PDF).
Leaves the original TYRA REBECCA NALUKWAGO.docx untouched.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    Frame,
    KeepTogether,
    NextPageTemplate,
    PageTemplate,
    BaseDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent
DOCX_OUT = ROOT / "TYRA_REBECCA_NALUKWAGO_CV.docx"
PDF_OUT = ROOT / "TYRA_REBECCA_NALUKWAGO_CV.pdf"
FONTS = Path(r"C:\Windows\Fonts")

# Palette — deep teal, calm and professional
TEAL = HexColor("#0F3D4C")
TEAL_MID = HexColor("#1A5A6E")
TEAL_SOFT = HexColor("#E8F1F4")
INK = HexColor("#2D3748")
MUTED = HexColor("#5A6A75")
RULE = HexColor("#B8CDD4")
GOLD = HexColor("#C4A35A")  # thin accent only

TEAL_RGB = RGBColor(0x0F, 0x3D, 0x4C)
TEAL_MID_RGB = RGBColor(0x1A, 0x5A, 0x6E)
TEAL_SOFT_RGB = RGBColor(0xE8, 0xF1, 0xF4)
INK_RGB = RGBColor(0x2D, 0x37, 0x48)
MUTED_RGB = RGBColor(0x5A, 0x6A, 0x75)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
GOLD_RGB = RGBColor(0xC4, 0xA3, 0x5A)

PROFILE = (
    "Accomplished Land Economist and Graduate Valuation Surveyor with over five years "
    "of professional experience in property valuation, land acquisition, compensation "
    "assessment, Resettlement Action Plans (RAPs), and property management. Demonstrated "
    "expertise in supporting large-scale infrastructure, energy, mining, water, and "
    "transport projects across Uganda for government agencies, financial institutions, "
    "and private-sector clients."
)
PROFILE_2 = (
    "Highly experienced in valuation methodologies, stakeholder engagement, community "
    "sensitisation, mortgage valuations, market analysis, and preparation of comprehensive "
    "valuation reports. A Graduate Member of the Institute of Surveyors of Uganda (ISU), "
    "committed to delivering professional, ethical, and data-driven valuation services "
    "that support sustainable development and investment."
)

HIGHLIGHTS = [
    "5+ years of professional valuation and land economics experience",
    "Graduate Member, Institute of Surveyors of Uganda (ISU)",
    "Extensive experience in Resettlement Action Plans (RAPs)",
    "Property valuation for mortgage, compensation and acquisition",
    "Worked on major national infrastructure, energy and mining projects",
    "Valuation assignments for leading commercial banks in Uganda",
    "Strong stakeholder engagement and community mobilisation skills",
]

EXPERIENCE = [
    {
        "title": "Valuation Surveyor",
        "org": "COVAE (U) Ltd",
        "dates": "Aug 2023 - Present",
        "bullets": [
            "Lead valuation assignments for mortgage, compensation and acquisition purposes.",
            "Prepare valuation methodologies and comprehensive valuation reports.",
            "Conduct market research and neighbourhood analysis.",
            "Support Resettlement Action Plan implementation for public infrastructure projects.",
        ],
    },
    {
        "title": "Valuer",
        "org": "GMT Consults Ltd",
        "dates": "Sept 2021 - Jul 2023",
        "bullets": [
            "Delivered RAP valuation assignments for road, water and electricity transmission projects.",
            "Coordinated stakeholder consultations, community engagement and grievance management.",
            "Managed valuation data collection, verification and reporting.",
        ],
    },
    {
        "title": "Valuer",
        "org": "Ministry of Energy & Mineral Development",
        "dates": "May 2022 - Jul 2022",
        "bullets": [
            "Negotiated wayleaves for rural electrification projects.",
            "Managed compensation assessments and stakeholder engagement.",
            "Supported electricity transmission corridor development.",
        ],
    },
    {
        "title": "Wayleave Liaison Officer",
        "org": "INTEC-GOPA International Energy Consultants",
        "dates": "May 2021 - Sept 2021",
        "bullets": [
            "Coordinated landowner negotiations and compensation processes.",
            "Facilitated access agreements and community liaison for power distribution projects.",
        ],
    },
    {
        "title": "Valuation Surveyor",
        "org": "COVAE (U) Ltd",
        "dates": "May 2018 - Apr 2020",
        "bullets": [
            "Conducted mortgage, compensation and property valuations.",
            "Prepared valuation reports for commercial banks, private clients and government projects.",
        ],
    },
]

EARLIER = [
    "Intern Valuer - Kampala Capital City Authority",
    "Intern Valuer - Peak Partners Property Valuers & Surveyors",
]

PROJECTS = [
    "Bagkara Mining Land Acquisition Project",
    "Buliisa Service Station Compensation Valuation",
    "Singo Military Training School Compensation Project",
    "Small Hydro Power Dam Resettlement Projects",
    "Isingiro Water Works Project",
    "Muhanga-Kisizi-Rwashamaire Road Project",
    "Masaka-Mbarara 400KV Transmission Line",
    "Rural Electrification Projects across Uganda",
    "Mortgage valuation assignments for major commercial banks",
]

TECHNICAL = [
    "Property Valuation",
    "Resettlement Action Plans (RAP)",
    "Land Acquisition",
    "Compensation Assessment",
    "Property Management",
    "Mortgage Valuation",
    "Property Market Analysis",
    "Community Engagement",
    "Geographic Information Systems (GIS)",
    "AutoCAD",
    "ArchiCAD",
    "REDAH Informatics",
    "VEKTA",
    "Microsoft Office Suite",
]

COMPETENCIES = [
    "Property & Asset Valuation",
    "Land Acquisition & Compensation",
    "Resettlement Action Planning (RAP)",
    "Stakeholder Engagement",
    "Community Mobilisation",
    "Property Market Research",
    "Negotiation & Conflict Resolution",
    "Valuation Reporting",
    "Project Coordination",
    "Data Collection & Analysis",
]

TAGLINE = "Senior Valuation Surveyor  ·  Land Economist  ·  Property & Resettlement Specialist"


# ---------------------------------------------------------------------------
# Font registration (PDF)
# ---------------------------------------------------------------------------

def register_fonts():
    pdfmetrics.registerFont(TTFont("Georgia", str(FONTS / "georgia.ttf")))
    pdfmetrics.registerFont(TTFont("Georgia-Bold", str(FONTS / "georgiab.ttf")))
    pdfmetrics.registerFont(TTFont("Georgia-Italic", str(FONTS / "georgiai.ttf")))
    pdfmetrics.registerFont(TTFont("Calibri", str(FONTS / "calibri.ttf")))
    pdfmetrics.registerFont(TTFont("Calibri-Bold", str(FONTS / "calibrib.ttf")))
    pdfmetrics.registerFont(TTFont("Calibri-Italic", str(FONTS / "calibrii.ttf")))
    pdfmetrics.registerFont(TTFont("Calibri-Light", str(FONTS / "calibril.ttf")))


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _set_run_font(run, name="Calibri", size=10, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _para_spacing(p, before=0, after=6, line=1.12):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    _para_spacing(p, before=14, after=2, line=1.0)
    run = p.add_run(text.upper())
    _set_run_font(run, name="Calibri", size=11, bold=True, color=TEAL_RGB)

    # Gold accent underline via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "1A5A6E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    _para_spacing(p, before=0, after=1.5, line=1.12)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    _set_run_font(run, name="Calibri", size=size, color=MUTED_RGB)
    return p


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # ---- Header banner ----
    header = doc.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(header)
    cell = header.rows[0].cells[0]
    _shade_cell(cell, "0F3D4C")
    _set_cell_margins(cell, top=160, bottom=140, left=200, right=200)

    name_p = cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(name_p, before=0, after=4, line=1.0)
    r = name_p.add_run("TYRA REBECCA NALUKWAGO")
    _set_run_font(r, name="Georgia", size=26, bold=True, color=WHITE_RGB)

    tag_p = cell.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(tag_p, before=0, after=2, line=1.0)
    r = tag_p.add_run(TAGLINE)
    _set_run_font(r, name="Calibri", size=11.5, color=RGBColor(0xD4, 0xE4, 0xEA))

    # Gold hairline under banner
    gold = doc.add_table(rows=1, cols=1)
    _remove_table_borders(gold)
    gcell = gold.rows[0].cells[0]
    _shade_cell(gcell, "C4A35A")
    _set_cell_margins(gcell, top=0, bottom=0, left=0, right=0)
    gp = gcell.paragraphs[0]
    _para_spacing(gp, before=0, after=0, line=0.5)
    r = gp.add_run(" ")
    _set_run_font(r, size=3)

    # Content wrapper with side margins via nested table
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    bcell = body.rows[0].cells[0]
    _set_cell_margins(bcell, top=80, bottom=40, left=200, right=200)

    def add_p(before=0, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = bcell.add_paragraph()
        p.alignment = align
        _para_spacing(p, before=before, after=after, line=1.12)
        return p

    def add_heading(text):
        p = add_p(before=12, after=3)
        r = p.add_run(text.upper())
        _set_run_font(r, name="Calibri", size=13, bold=True, color=TEAL_RGB)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), "1A5A6E")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Clear the default empty paragraph
    bcell.paragraphs[0].clear()

    add_heading("Professional Profile")
    for text in (PROFILE, PROFILE_2):
        p = add_p(before=2, after=4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        _set_run_font(r, name="Calibri", size=11, color=INK_RGB)

    # Highlights in soft panel
    add_heading("Career Highlights")
    hl = bcell.add_table(rows=1, cols=1)
    _remove_table_borders(hl)
    hcell = hl.rows[0].cells[0]
    _shade_cell(hcell, "E8F1F4")
    _set_cell_margins(hcell, top=60, bottom=50, left=100, right=100)

    mid = (len(HIGHLIGHTS) + 1) // 2
    left, right = HIGHLIGHTS[:mid], HIGHLIGHTS[mid:]
    cols = hcell.add_table(rows=max(len(left), len(right)), cols=2)
    _remove_table_borders(cols)
    for i in range(max(len(left), len(right))):
        for ci, items in enumerate((left, right)):
            c = cols.rows[i].cells[ci]
            c.text = ""
            _set_cell_margins(c, top=20, bottom=20, left=40, right=60)
            p = c.paragraphs[0]
            _para_spacing(p, before=0, after=0, line=1.15)
            if i < len(items):
                r = p.add_run(f"•  {items[i]}")
                _set_run_font(r, name="Calibri", size=10.5, color=TEAL_MID_RGB)

    add_heading("Professional Experience")
    for job in EXPERIENCE:
        # Title + dates row
        row = bcell.add_table(rows=1, cols=2)
        _remove_table_borders(row)
        row.columns[0].width = Cm(11.5)
        row.columns[1].width = Cm(5.5)
        lc, rc = row.rows[0].cells
        _set_cell_margins(lc, top=40, bottom=0, left=0, right=40)
        _set_cell_margins(rc, top=40, bottom=0, left=0, right=0)
        lp = lc.paragraphs[0]
        _para_spacing(lp, before=0, after=0, line=1.05)
        r = lp.add_run(job["title"])
        _set_run_font(r, name="Calibri", size=12.5, bold=True, color=TEAL_RGB)
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _para_spacing(rp, before=0, after=0, line=1.05)
        r = rp.add_run(job["dates"])
        _set_run_font(r, name="Calibri", size=10.5, color=TEAL_MID_RGB)

        org = add_p(before=0, after=2)
        r = org.add_run(job["org"])
        _set_run_font(r, name="Calibri", size=11, italic=True, color=TEAL_MID_RGB)

        for b in job["bullets"]:
            p = add_p(before=0, after=1)
            p.paragraph_format.left_indent = Cm(0.35)
            r = p.add_run(f"•  {b}")
            _set_run_font(r, name="Calibri", size=11, color=MUTED_RGB)

    p = add_p(before=8, after=2)
    r = p.add_run("Earlier Experience")
    _set_run_font(r, name="Calibri", size=11.5, bold=True, color=TEAL_RGB)
    for item in EARLIER:
        p = add_p(before=0, after=1)
        p.paragraph_format.left_indent = Cm(0.35)
        r = p.add_run(f"•  {item}")
        _set_run_font(r, name="Calibri", size=11, color=MUTED_RGB)

    # Education | Membership
    add_heading("Education & Membership")
    em = bcell.add_table(rows=2, cols=2)
    _remove_table_borders(em)
    cells = [
        (0, 0, "Bachelor of Science (Hons) in Land Economics", True, TEAL_RGB),
        (1, 0, "Makerere University", False, TEAL_MID_RGB),
        (0, 1, "Graduate Member, ISU", True, TEAL_RGB),
        (1, 1, "Institute of Surveyors of Uganda", False, TEAL_MID_RGB),
    ]
    for ri, ci, text, bold, color in cells:
        c = em.rows[ri].cells[ci]
        c.text = ""
        _set_cell_margins(c, top=20, bottom=20, left=0, right=80)
        p = c.paragraphs[0]
        _para_spacing(p, before=0, after=0, line=1.1)
        r = p.add_run(text)
        _set_run_font(
            r, name="Calibri", size=11.5 if bold else 11, bold=bold, italic=not bold, color=color
        )

    add_heading("Key Projects")
    mid = (len(PROJECTS) + 1) // 2
    left, right = PROJECTS[:mid], PROJECTS[mid:]
    pt = bcell.add_table(rows=max(len(left), len(right)), cols=2)
    _remove_table_borders(pt)
    for i in range(max(len(left), len(right))):
        for ci, items in enumerate((left, right)):
            c = pt.rows[i].cells[ci]
            c.text = ""
            _set_cell_margins(c, top=15, bottom=15, left=0, right=60)
            p = c.paragraphs[0]
            _para_spacing(p, before=0, after=0, line=1.12)
            if i < len(items):
                r = p.add_run(f"•  {items[i]}")
                _set_run_font(r, name="Calibri", size=11, color=MUTED_RGB)

    add_heading("Technical Expertise")
    p = add_p(before=2, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("  ·  ".join(TECHNICAL))
    _set_run_font(r, name="Calibri", size=11, color=INK_RGB)

    add_heading("Core Competencies")
    mid = (len(COMPETENCIES) + 1) // 2
    left, right = COMPETENCIES[:mid], COMPETENCIES[mid:]
    ct = bcell.add_table(rows=max(len(left), len(right)), cols=2)
    _remove_table_borders(ct)
    for i in range(max(len(left), len(right))):
        for ci, items in enumerate((left, right)):
            c = ct.rows[i].cells[ci]
            c.text = ""
            _set_cell_margins(c, top=15, bottom=15, left=0, right=60)
            p = c.paragraphs[0]
            _para_spacing(p, before=0, after=0, line=1.12)
            if i < len(items):
                r = p.add_run(f"•  {items[i]}")
                _set_run_font(r, name="Calibri", size=11, color=MUTED_RGB)

    add_heading("Languages")
    p = add_p(before=2, after=6)
    r = p.add_run("English")
    _set_run_font(r, name="Calibri", size=11.5, bold=True, color=TEAL_RGB)
    r = p.add_run("  ·  Fluent")
    _set_run_font(r, name="Calibri", size=11.5, color=MUTED_RGB)
    r = p.add_run("          ")
    _set_run_font(r, size=11.5)
    r = p.add_run("Luganda")
    _set_run_font(r, name="Calibri", size=11.5, bold=True, color=TEAL_RGB)
    r = p.add_run("  ·  Native")
    _set_run_font(r, name="Calibri", size=11.5, color=MUTED_RGB)

    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


# ---------------------------------------------------------------------------
# PDF flowables & styles
# ---------------------------------------------------------------------------

class ColoredBox(Flowable):
    """Rounded-ish soft panel by drawing a filled rect behind child table."""

    def __init__(self, inner, width, pad=8, bg=TEAL_SOFT, radius=4):
        Flowable.__init__(self)
        self.inner = inner
        self.box_width = width
        self.pad = pad
        self.bg = bg
        self.radius = radius

    def wrap(self, availWidth, availHeight):
        w = self.box_width or availWidth
        iw, ih = self.inner.wrap(w - 2 * self.pad, availHeight - 2 * self.pad)
        self.width = w
        self.height = ih + 2 * self.pad
        self._ih = ih
        return self.width, self.height

    def draw(self):
        self.canv.setFillColor(self.bg)
        self.canv.roundRect(0, 0, self.width, self.height, self.radius, fill=1, stroke=0)
        self.inner.drawOn(self.canv, self.pad, self.pad)


class SectionRule(Flowable):
    """Section title with teal rule and short gold accent tick."""

    def __init__(self, title, styles, width):
        Flowable.__init__(self)
        self.title = title.upper()
        self.styles = styles
        self.rule_width = width

    def wrap(self, availWidth, availHeight):
        self.width = self.rule_width or availWidth
        self.height = 22
        return self.width, self.height

    def draw(self):
        c = self.canv
        c.setFillColor(TEAL)
        c.setFont("Calibri-Bold", 13)
        c.drawString(0, 7, self.title)
        # Main teal rule
        c.setStrokeColor(TEAL_MID)
        c.setLineWidth(0.9)
        c.line(0, 2.5, self.width, 2.5)
        # Short gold accent at start
        c.setStrokeColor(GOLD)
        c.setLineWidth(2.2)
        c.line(0, 2.5, 28, 2.5)



def _styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CVBody",
            fontName="Calibri",
            fontSize=10.5,
            leading=13.2,
            textColor=INK,
            alignment=TA_JUSTIFY,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVJobTitle",
            fontName="Calibri-Bold",
            fontSize=11.5,
            leading=14,
            textColor=TEAL,
            spaceBefore=0,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVDates",
            fontName="Calibri",
            fontSize=10.5,
            leading=13,
            textColor=TEAL_MID,
            alignment=TA_RIGHT,
            spaceBefore=0,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVOrg",
            fontName="Calibri-Italic",
            fontSize=10.5,
            leading=13,
            textColor=TEAL_MID,
            spaceBefore=0,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBullet",
            fontName="Calibri",
            fontSize=10.5,
            leading=13,
            textColor=MUTED,
            leftIndent=8,
            spaceAfter=0.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVCell",
            fontName="Calibri",
            fontSize=10.5,
            leading=13.2,
            textColor=MUTED,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVHighlight",
            fontName="Calibri",
            fontSize=10.2,
            leading=13,
            textColor=TEAL_MID,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVSubhead",
            fontName="Calibri-Bold",
            fontSize=11.5,
            leading=14,
            textColor=TEAL,
            spaceBefore=4,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVLang",
            fontName="Calibri",
            fontSize=11,
            leading=13.5,
            textColor=MUTED,
        )
    )
    return styles


def _two_col(items, styles, style_name="CVCell", marker="•  ", col_w=None):
    mid = (len(items) + 1) // 2
    left, right = items[:mid], items[mid:]
    rows = []
    for i in range(max(len(left), len(right))):
        l = f"{marker}{left[i]}" if i < len(left) else ""
        r = f"{marker}{right[i]}" if i < len(right) else ""
        rows.append(
            [Paragraph(l, styles[style_name]), Paragraph(r, styles[style_name])]
        )
    if col_w is None:
        col_w = [8.6 * cm, 8.6 * cm]
    t = Table(rows, colWidths=col_w)
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ]
        )
    )
    return t


def _draw_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(MUTED)
    canvas.setFont("Calibri", 9)
    page_w, _ = A4
    canvas.drawCentredString(
        page_w / 2, 0.9 * cm, f"Tyra Rebecca Nalukwago  ·  Page {doc.page}"
    )
    canvas.restoreState()


def _draw_first_page(canvas, doc):
    """Full-bleed teal header banner + footer."""
    page_w, page_h = A4
    banner_h = 3.7 * cm
    canvas.saveState()
    canvas.setFillColor(TEAL)
    canvas.rect(0, page_h - banner_h, page_w, banner_h, fill=1, stroke=0)
    # Gold accent line under banner
    canvas.setFillColor(GOLD)
    canvas.rect(0, page_h - banner_h - 2.5, page_w, 2.5, fill=1, stroke=0)
    # Name
    canvas.setFillColor(white)
    canvas.setFont("Georgia-Bold", 24)
    canvas.drawCentredString(page_w / 2, page_h - 1.65 * cm, "TYRA REBECCA NALUKWAGO")
    # Tagline
    canvas.setFillColor(HexColor("#D4E4EA"))
    canvas.setFont("Calibri", 11)
    canvas.drawCentredString(page_w / 2, page_h - 2.45 * cm, TAGLINE)
    canvas.restoreState()
    _draw_footer(canvas, doc)


def build_pdf():
    register_fonts()
    styles = _styles()
    page_w, page_h = A4
    content_w = page_w - 3.0 * cm
    banner_h = 3.7 * cm

    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.1 * cm,
        bottomMargin=1.4 * cm,
    )

    # Page 1: leave room under full-bleed banner
    frame1 = Frame(
        1.5 * cm,
        1.4 * cm,
        content_w,
        page_h - banner_h - 2.5 - 1.4 * cm - 0.25 * cm,
        id="first",
    )
    frame2 = Frame(
        1.5 * cm,
        1.4 * cm,
        content_w,
        page_h - 1.1 * cm - 1.4 * cm,
        id="later",
    )
    doc.addPageTemplates(
        [
            PageTemplate(id="first", frames=frame1, onPage=_draw_first_page),
            PageTemplate(id="later", frames=frame2, onPage=_draw_footer),
        ]
    )

    story = [NextPageTemplate("later"), Spacer(1, 4)]

    # Profile
    story.append(SectionRule("Professional Profile", styles, content_w))
    story.append(Spacer(1, 5))
    story.append(Paragraph(PROFILE, styles["CVBody"]))
    story.append(Paragraph(PROFILE_2, styles["CVBody"]))

    # Highlights panel
    story.append(SectionRule("Career Highlights", styles, content_w))
    story.append(Spacer(1, 5))
    hl_table = _two_col(
        HIGHLIGHTS,
        styles,
        style_name="CVHighlight",
        marker="•  ",
        col_w=[8.3 * cm, 8.3 * cm],
    )
    story.append(ColoredBox(hl_table, content_w, pad=10, bg=TEAL_SOFT, radius=5))
    story.append(Spacer(1, 6))

    # Experience
    story.append(SectionRule("Professional Experience", styles, content_w))
    story.append(Spacer(1, 2))
    for job in EXPERIENCE:
        title_row = Table(
            [
                [
                    Paragraph(job["title"], styles["CVJobTitle"]),
                    Paragraph(job["dates"], styles["CVDates"]),
                ]
            ],
            colWidths=[11.5 * cm, 5.7 * cm],
        )
        title_row.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        block = [
            title_row,
            Paragraph(job["org"], styles["CVOrg"]),
        ]
        for b in job["bullets"]:
            block.append(Paragraph(f"•  {b}", styles["CVBullet"]))
        # Keep last role + earlier experience together to avoid orphans
        if job is EXPERIENCE[-1]:
            block.append(Paragraph("Earlier Experience", styles["CVSubhead"]))
            for item in EARLIER:
                block.append(Paragraph(f"•  {item}", styles["CVBullet"]))
        story.append(KeepTogether(block))

    # Education & Membership
    story.append(SectionRule("Education & Membership", styles, content_w))
    story.append(Spacer(1, 5))
    edu_mem = Table(
        [
            [
                Paragraph(
                    "<font color='#0F3D4C'><b>Bachelor of Science (Hons) in Land Economics</b></font>"
                    "<br/><font color='#1A5A6E'><i>Makerere University</i></font>",
                    styles["CVCell"],
                ),
                Paragraph(
                    "<font color='#0F3D4C'><b>Graduate Member</b></font>"
                    "<br/><font color='#1A5A6E'><i>Institute of Surveyors of Uganda (ISU)</i></font>",
                    styles["CVCell"],
                ),
            ]
        ],
        colWidths=[8.6 * cm, 8.6 * cm],
    )
    edu_mem.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, -1), TEAL_SOFT),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.append(edu_mem)

    # Key Projects
    story.append(Spacer(1, 8))
    story.append(
        KeepTogether(
            [
                SectionRule("Key Projects", styles, content_w),
                Spacer(1, 6),
                _two_col(PROJECTS, styles),
            ]
        )
    )

    story.append(Spacer(1, 8))
    story.append(SectionRule("Technical Expertise", styles, content_w))
    story.append(Spacer(1, 6))
    story.append(Paragraph("  ·  ".join(TECHNICAL), styles["CVBody"]))

    story.append(Spacer(1, 8))
    comp_table = _two_col(COMPETENCIES, styles)
    story.append(
        KeepTogether(
            [
                SectionRule("Core Competencies", styles, content_w),
                Spacer(1, 6),
                ColoredBox(comp_table, content_w, pad=10, bg=TEAL_SOFT, radius=5),
            ]
        )
    )

    story.append(Spacer(1, 8))
    story.append(SectionRule("Languages", styles, content_w))
    story.append(Spacer(1, 6))
    lang_box = Table(
        [
            [
                Paragraph(
                    "<font color='#0F3D4C'><b>English</b></font>"
                    "<br/><font color='#5A6A75'>Fluent</font>",
                    styles["CVCell"],
                ),
                Paragraph(
                    "<font color='#0F3D4C'><b>Luganda</b></font>"
                    "<br/><font color='#5A6A75'>Native</font>",
                    styles["CVCell"],
                ),
            ]
        ],
        colWidths=[8.6 * cm, 8.6 * cm],
    )
    lang_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TEAL_SOFT),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(lang_box)

    doc.build(story)
    print(f"Wrote {PDF_OUT}")


def render_preview():
    import pypdfium2 as pdfium

    preview_dir = ROOT / "tmp_preview"
    preview_dir.mkdir(exist_ok=True)
    for old in preview_dir.glob("*.png"):
        old.unlink()
    pdf = pdfium.PdfDocument(str(PDF_OUT))
    print(f"PDF pages: {len(pdf)}")
    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=2.2)
        pil = bitmap.to_pil()
        out = preview_dir / f"cv_page_{i + 1}.png"
        pil.save(out)
        print(f"Preview: {out}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    render_preview()
