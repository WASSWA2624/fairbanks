"""
Build Medequip Healthcare Rwanda Biomedical Engineer CV + cover letter.
Sources: wilson root candidate documents; job post Job in Rwanda.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT = Path(__file__).resolve().parent
WILSON = OUT.parent
SIGNATURE_SRC = WILSON / "signature.jpeg"
SIGNATURE = OUT / "signature_clean.png"
NAVY = HexColor("#0F2C4C")
TEAL = HexColor("#1A6B5C")
GRAY = HexColor("#2F2F2F")
MUTED = HexColor("#5A5A5A")


class NumberedCanvas(pdf_canvas.Canvas):
    """Draw 'Page x of y' centered in the footer on every page."""

    def __init__(self, *args, **kwargs):
        pdf_canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_page_number(page_count)
            pdf_canvas.Canvas.showPage(self)
        pdf_canvas.Canvas.save(self)

    def _draw_page_number(self, page_count: int) -> None:
        page_num = self._pageNumber
        self.setFont("Helvetica", 9)
        self.setFillColor(MUTED)
        self.drawCentredString(
            A4[0] / 2.0,
            8 * mm,
            f"Page {page_num} of {page_count}",
        )


def add_docx_page_footer(doc: Document) -> None:
    """Add centered Word footer: Page X of Y."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # clear existing runs
    for run in list(p.runs):
        run._element.getparent().remove(run._element)

    def add_field(paragraph, instr: str) -> None:
        run = paragraph.add_run()
        fld_char_begin = run._r.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "begin"}
        )
        run._r.append(fld_char_begin)

        run2 = paragraph.add_run()
        instr_text = run2._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr_text.text = instr
        run2._r.append(instr_text)

        run3 = paragraph.add_run()
        fld_char_end = run3._r.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "end"}
        )
        run3._r.append(fld_char_end)

    r = p.add_run("Page ")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    add_field(p, " PAGE ")
    r = p.add_run(" of ")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    add_field(p, " NUMPAGES ")
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)


def prepare_signature() -> Path:
    """Clean signature: keep main ink only, white background, flush left crop."""
    from collections import deque
    from PIL import Image, ImageFilter, ImageOps

    if not SIGNATURE_SRC.exists():
        raise FileNotFoundError(f"Missing signature file: {SIGNATURE_SRC}")

    im = Image.open(SIGNATURE_SRC).convert("RGB")
    w, h = im.size
    im = im.crop((int(w * 0.04), int(h * 0.18), int(w * 0.96), int(h * 0.88)))
    im = im.filter(ImageFilter.MedianFilter(size=3))
    gray = ImageOps.grayscale(im)

    samples = [
        gray.getpixel((x, y))
        for y in range(gray.height // 3, 2 * gray.height // 3, 5)
        for x in range(gray.width // 3, 2 * gray.width // 3, 5)
    ]
    samples.sort()
    paper = samples[int(len(samples) * 0.75)]
    thresh = max(0, paper - 45)

    gp = gray.load()
    ink_set = {
        (x, y)
        for y in range(im.height)
        for x in range(im.width)
        if gp[x, y] < thresh
    }

    # Keep only the largest connected ink component (drops scan speckles)
    visited = set()
    best = []
    for seed in list(ink_set):
        if seed in visited:
            continue
        q = deque([seed])
        visited.add(seed)
        comp = [seed]
        while q:
            x, y = q.popleft()
            for dy in (-1, 0, 1):
                for dx in (-1, 0, 1):
                    nxt = (x + dx, y + dy)
                    if nxt in ink_set and nxt not in visited:
                        visited.add(nxt)
                        q.append(nxt)
                        comp.append(nxt)
        if len(comp) > len(best):
            best = comp

    out = Image.new("RGB", im.size, (255, 255, 255))
    op = out.load()
    for x, y in best:
        op[x, y] = (10, 10, 10)

    if best:
        xs = [p[0] for p in best]
        ys = [p[1] for p in best]
        # Flush left (2px), modest top/right/bottom padding
        box = (
            max(0, min(xs) - 2),
            max(0, min(ys) - 6),
            min(im.width, max(xs) + 8),
            min(im.height, max(ys) + 6),
        )
        out = out.crop(box)

    out.save(SIGNATURE)
    return SIGNATURE


def signature_flowable(width_mm: float = 40):
    from PIL import Image as PILImage

    prepare_signature()
    with PILImage.open(SIGNATURE) as im:
        w, h = im.size
    height_mm = width_mm * (h / w)
    img = RLImage(str(SIGNATURE), width=width_mm * mm, height=height_mm * mm)
    img.hAlign = "LEFT"
    return img


def add_signature_docx(doc: Document, width_in: float = 1.55) -> None:
    """Insert left-aligned signature tightly above the typed name."""
    prepare_signature()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.paragraph_format.space_before = Pt(1)
    sig.paragraph_format.space_after = Pt(0)
    sig.paragraph_format.left_indent = Inches(0)
    sig.paragraph_format.first_line_indent = Inches(0)
    run = sig.add_run()
    run.add_picture(str(SIGNATURE), width=Inches(width_in))


def signature_block_pdf(styles, name_style: str, title_style: str, width_mm: float = 38):
    """Signature flush-left above typed name and title."""
    img = signature_flowable(width_mm)
    data = [
        [img],
        [Paragraph("Wasswa Wilson", styles[name_style])],
        [Paragraph("Biomedical Engineer", styles[title_style])],
    ]
    table = Table(data, colWidths=[max(70 * mm, width_mm * mm + 5 * mm)])
    table.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (0, 0), 0),
                ("BOTTOMPADDING", (0, 0), (0, 0), 1),
                ("TOPPADDING", (0, 1), (0, 1), 2),
                ("BOTTOMPADDING", (0, 1), (0, 1), 0),
                ("TOPPADDING", (0, 2), (0, 2), 0),
                ("BOTTOMPADDING", (0, 2), (0, 2), 0),
            ]
        )
    )
    table.hAlign = "LEFT"
    return table


CONTACT = {
    "name": "WASSWA WILSON",
    "title": "Biomedical Engineer",
    "location": "Kampala, Uganda | Ready for Kigali",
    "email": "wasswawilson0001@gmail.com",
    "phone": "+256783230321",
    "languages": "English (fluent) · Luganda (native)",
    "years": "7+ years in biomedical equipment work",
}

SUMMARY = (
    "I am a biomedical engineer with over 7 years of hands-on hospital "
    "experience: installing, commissioning, calibrating, maintaining, and "
    "training users on medical equipment. I have carried out this work in "
    "Uganda, the Democratic Republic of Congo, Kenya, Tanzania, and "
    "Somaliland. At International Hospital Kampala I led equipment "
    "installation across the laboratory, theatre, ICU, and radiology "
    "departments - covering haematology analysers, chemistry analysers, blood "
    "bank machines, CT, X-ray, oxygen plant, and critical-care systems. Since "
    "March 2025 I have also worked on personal contracts for hospitals that "
    "call me in for installs and service, including lab, theatre, and "
    "radiology work in Congo (among them the hospital that serves the Kibali "
    "Gold Mine). I am used to travelling for field jobs, working independently "
    "on site, writing clear service and maintenance notes, and handing over "
    "equipment with trained clinical staff. I want to bring that regional "
    "field experience to Medequip's technical team and clients across Rwanda "
    "and beyond."
)

SKILLS = [
    (
        "What I do on site",
        "Install, calibrate, commission, prevent faults, fix breakdowns, "
        "and hand over equipment ready for clinical use",
    ),
    (
        "Equipment I know well",
        "Radiology (X-ray, CT, ultrasound, C-arm); theatre and ICU "
        "(ventilators, monitors, infusion pumps, theatre systems); "
        "laboratory (haematology analysers, chemistry analysers, blood bank "
        "machines); support plant (RO water, oxygen)",
    ),
    (
        "Countries worked",
        "Uganda, Democratic Republic of Congo, Kenya, Tanzania, and "
        "Somaliland - field installs, maintenance, and user training",
    ),
    (
        "Client support",
        "User training, procurement and commissioning support, service "
        "reports, maintenance logs, accreditation paperwork",
    ),
    (
        "How I work",
        "Clear spoken English, independent field work, cross-border travel, "
        "quick learner on new brands and models",
    ),
]

ROLES = [
    {
        "title": "Independent Biomedical Engineer (personal contracts)",
        "dates": "Mar 2025 - Present",
        "bullets": [
            "Install and maintain medical equipment for hospitals on a "
            "personal-contract basis after leaving GFF.",
            "Complete field assignments in Uganda, Congo, Kenya, Tanzania, "
            "and Somaliland - travel to site, install or repair, train users, "
            "and report back to the client.",
            "Completed lab, theatre, and radiology installations at several "
            "hospitals in the Democratic Republic of Congo, including the "
            "hospital that serves the Kibali Gold Mine.",
            "Lab installs and support cover haematology analysers, clinical "
            "chemistry analysers, and blood bank machines, with user handover "
            "after commissioning.",
            "Handle preventive and corrective maintenance, on-site "
            "troubleshooting, and clear service documentation after each job.",
        ],
    },
    {
        "title": "Biomedical Programs Manager, Gould Family Foundation (GFF)",
        "dates": "Aug 2024 - Feb 2025",
        "bullets": [
            "Ran biomedical programmes across several health facilities, from "
            "planning through to equipment readiness.",
            "Supported buying, installing, and commissioning of medical "
            "technologies for partner sites.",
            "Trained engineers and technicians on correct use and care of "
            "equipment.",
            "Kept work in line with healthcare technology and quality "
            "standards, including documentation for programme review.",
        ],
    },
    {
        "title": "Biomedical Manager, International Hospital Kampala (IHK)",
        "dates": "Jan 2020 - Jan 2024",
        "bullets": [
            "Managed day-to-day biomedical engineering for a major private "
            "hospital for four years.",
            "Led installation and commissioning of equipment across the IHK "
            "laboratory, theatre, ICU, and radiology departments.",
            "Lab leadership included haematology analysers, chemistry "
            "analysers, and blood bank machines; imaging included CT and "
            "X-ray; ICU and theatre systems were part of the same programme.",
            "Also led installs of oxygen plant and related support systems.",
            "Set up preventive maintenance routines to cut downtime and keep "
            "devices safe for clinical use.",
            "Supported procurement, lifecycle planning, service records, "
            "staff training, and COHSASA accreditation work.",
        ],
    },
    {
        "title": "Biomedical Engineer, Norvik Hospital Ltd",
        "dates": "Apr 2019 - Jan 2020",
        "bullets": [
            "Serviced diagnostic and patient-monitoring equipment day to day.",
            "Supported installation and calibration of imaging and ICU devices.",
            "Helped build and follow preventive maintenance schedules and "
            "service follow-up.",
        ],
    },
]

EARLIER = [
    (
        "Research Intern, Uganda Virus Research Institute",
        "Dec 2018 - Apr 2019",
        "Supported biomedical research projects and laboratory data work.",
    ),
    (
        "Teaching Assistant, Makerere University, College of Health Sciences",
        "Mar 2016 - Aug 2017",
        "Supported biomedical engineering teaching and laboratory practical "
        "sessions.",
    ),
]

ACHIEVEMENTS = [
    "7+ years of hands-on hospital biomedical work across Uganda, Congo, "
    "Kenya, Tanzania, and Somaliland.",
    "At IHK, led equipment installation across laboratory, theatre, ICU, and "
    "radiology departments.",
    "Installed and supported haematology, chemistry, and blood bank laboratory "
    "machines alongside theatre and radiology systems.",
    "Completed Congo hospital installs including the facility serving Kibali "
    "Gold Mine.",
    "Commissioned CT, X-ray, oxygen plant, and ICU systems in Uganda and "
    "supported COHSASA accreditation at IHK.",
    "Trained engineers, technicians, and clinical users so equipment stays "
    "useful after install day.",
]

EQUIPMENT = [
    ("Radiology / imaging", "CT scanners, X-ray machines, ultrasound systems, C-arms"),
    (
        "Theatre and ICU",
        "Theatre systems, ventilators, patient monitors, infusion pumps, "
        "related critical-care devices",
    ),
    (
        "Laboratory - haematology",
        "Haematology analysers and related blood-count laboratory systems",
    ),
    (
        "Laboratory - chemistry",
        "Clinical chemistry analysers and related chemistry lab systems",
    ),
    (
        "Laboratory - blood bank",
        "Blood bank machines and transfusion / blood-banking support equipment",
    ),
    ("Support plant", "Oxygen plants, RO water systems"),
    ("Maternity / neonatal", "Incubators, phototherapy units, foetal Dopplers, CTG"),
]

EDUCATION = [
    {
        "level": "Bachelor's degree (undergraduate)",
        "school": "Makerere University",
        "detail": "BSc Biomedical Engineering, 2012-2016.",
    },
    {
        "level": "Uganda Advanced Certificate of Education (A-Level)",
        "school": "Mengo Secondary School",
        "detail": (
            "UACE 2011. Principal subjects: Mathematics, Physics, Biology, "
            "and Chemistry."
        ),
    },
    {
        "level": "Uganda Certificate of Education (O-Level)",
        "school": "St. John's Wakiso Secondary School",
        "detail": "UCE 2009.",
    },
]

CERTIFICATIONS = [
    {
        "level": "Professional development certificate",
        "school": "University of Washington",
        "detail": "Leadership and Management in Health, December 2022.",
    },
    {
        "level": "Technical certificate",
        "school": "Greenbridge School of Open Technologies",
        "detail": "JAVA Programming - Level 1, Dec 2015 - Jan 2016.",
    },
    {
        "level": "Workplace safety training",
        "school": "Fire Technologies Limited (IHK)",
        "detail": (
            "Fire safety, prevention, firefighting, emergency scene "
            "management, and evacuation drills, April 2021."
        ),
    },
]

REFEREES = [
    "Eng. Richard Ssejongo - Biomedical Engineer, St. Francis Hospital Nsambya - "
    "+256 753 818 754 / +256 777 132 489",
    "Dr. Annet Khingi - Senior Radiologist and Administrator, Mengo Hospital - "
    "+256 772 592 771 / +256 701 592 771",
    "Racheal Nabukeera - Director, FairBanks Medical Centre - "
    "+256772849258 / nracheal017@gmail.com",
]

LETTER_DATE = "28 July 2026"
LETTER_BODY = [
    "Dear Hiring Manager,",
    (
        "I am applying for the Biomedical Engineer job at Medequip Healthcare "
        "Ltd in Kigali. Your team installs equipment, keeps it running, and "
        "trains hospital staff across Rwanda and beyond. That is the work I "
        "already do on the ground, and I want to do it with Medequip."
    ),
    (
        "I have over 7 years installing, commissioning, maintaining, and "
        "training users on hospital equipment. At Norvik Hospital I serviced "
        "diagnostic and monitoring systems day to day. At International "
        "Hospital Kampala I spent four years as Biomedical Manager - leading "
        "equipment installation across the laboratory, theatre, ICU, and "
        "radiology departments. That included haematology analysers, chemistry "
        "analysers, blood bank machines, CT, X-ray, oxygen plant, and ICU "
        "systems, plus preventive maintenance, staff training, and COHSASA "
        "accreditation support. I later managed biomedical programmes for the "
        "Gould Family Foundation across several facilities."
    ),
    (
        "Since March 2025 I have worked on personal contracts: installing and "
        "maintaining equipment for hospitals that call me in. That includes "
        "lab, theatre, and radiology jobs in Congo, among them the hospital "
        "that serves the Kibali Gold Mine. On the lab side I handle "
        "haematology analysers, chemistry analysers, and blood bank machines. "
        "I have done this kind of field work in Uganda, Congo, Kenya, "
        "Tanzania, and Somaliland - travel to site, install or repair, train "
        "users, and leave clear service notes."
    ),
    (
        "I hold a BSc in Biomedical Engineering from Makerere University, but "
        "what I bring Medequip is field experience: show up, install or "
        "repair, train the team, document the job, and travel when needed. I "
        "speak clear English, I am ready to relocate to Kigali, and I am keen "
        "to learn your product lines through international training."
    ),
    (
        "My CV is attached. Thank you for your time - I would be glad to talk."
    ),
]


def set_docx_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(0)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_heading_run(paragraph, text: str, size: float, bold: bool = True, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return run


def section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_heading_run(p, text.upper(), 11.5, True, RGBColor(0x0F, 0x2C, 0x4C))
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "1",
            qn("w:color"): "1A6B5C",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def body_para(doc: Document, text: str, space_after: int = 6) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.12
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)


def role_header(doc: Document, title: str, dates: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    add_heading_run(p, title, 11, True, RGBColor(0x1A, 0x1A, 0x1A))
    r = p.add_run(f"  |  {dates}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x5C)
    r.font.name = "Calibri"


def build_cv_docx(path: Path) -> None:
    doc = Document()
    set_docx_defaults(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    add_heading_run(name, CONTACT["name"], 20, True, RGBColor(0x0F, 0x2C, 0x4C))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    add_heading_run(title, CONTACT["title"], 12, True, RGBColor(0x1A, 0x6B, 0x5C))

    years = doc.add_paragraph()
    years.alignment = WD_ALIGN_PARAGRAPH.CENTER
    years.paragraph_format.space_after = Pt(4)
    add_heading_run(years, CONTACT["years"], 11, True, RGBColor(0x1A, 0x6B, 0x5C))

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1)
    r = contact.add_run(f"{CONTACT['email']}  ·  {CONTACT['phone']}")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    loc = doc.add_paragraph()
    loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc.paragraph_format.space_after = Pt(1)
    r = loc.add_run(CONTACT["location"])
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    langs = doc.add_paragraph()
    langs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    langs.paragraph_format.space_after = Pt(4)
    r = langs.add_run(CONTACT["languages"])
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    section_title(doc, "Profile")
    body_para(doc, SUMMARY, space_after=4)

    section_title(doc, "Skills for this job")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        add_heading_run(p, f"{label}: ", 11, True, RGBColor(0x0F, 0x2C, 0x4C))
        r = p.add_run(value)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)

    section_title(doc, "Work experience")
    for role in ROLES:
        role_header(doc, role["title"], role["dates"])
        for b in role["bullets"]:
            bullet(doc, b)

    section_title(doc, "Earlier roles")
    for title, dates, note in EARLIER:
        role_header(doc, title, dates)
        bullet(doc, note)

    section_title(doc, "Selected highlights")
    for item in ACHIEVEMENTS:
        bullet(doc, item)

    section_title(doc, "Equipment experience")
    for label, value in EQUIPMENT:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_heading_run(p, f"{label}: ", 11, True, RGBColor(0x0F, 0x2C, 0x4C))
        r = p.add_run(value)
        r.font.size = Pt(11)

    section_title(doc, "Education and academic levels")
    for item in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        add_heading_run(p, item["level"], 10.5, True, RGBColor(0x1A, 0x6B, 0x5C))
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        add_heading_run(p2, item["school"], 11, True)
        body_para(doc, item["detail"], space_after=2)

    section_title(doc, "Certificates and short courses")
    for item in CERTIFICATIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        add_heading_run(p, item["level"], 10.5, True, RGBColor(0x1A, 0x6B, 0x5C))
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        add_heading_run(p2, item["school"], 11, True)
        body_para(doc, item["detail"], space_after=2)

    section_title(doc, "Referees")
    for ref in REFEREES:
        body_para(doc, ref, space_after=3)

    if SIGNATURE_SRC.exists():
        section_title(doc, "Signature")
        add_signature_docx(doc, width_in=1.55)
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_p.paragraph_format.space_before = Pt(2)
        name_p.paragraph_format.space_after = Pt(1)
        add_heading_run(name_p, "Wasswa Wilson", 11, True)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(0)
        r = title_p.add_run("Biomedical Engineer")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    add_docx_page_footer(doc)
    doc.save(path)


def build_letter_docx(path: Path) -> None:
    doc = Document()
    set_docx_defaults(doc)

    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(2)
    add_heading_run(header, "WASSWA WILSON", 18, True, RGBColor(0x0F, 0x2C, 0x4C))

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    add_heading_run(
        sub,
        "Biomedical Engineer · 7+ years experience",
        11.5,
        True,
        RGBColor(0x1A, 0x6B, 0x5C),
    )

    for line in ["Kampala, Uganda", CONTACT["email"], CONTACT["phone"]]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    pPr = rule._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "1",
            qn("w:color"): "1A6B5C",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)

    date_p = doc.add_paragraph(LETTER_DATE)
    date_p.paragraph_format.space_after = Pt(10)

    for line in [
        "Hiring Manager",
        "Medequip Healthcare Ltd",
        "Kigali, Rwanda",
        "To: admin@medequipltd.com; finance.rwanda@medequipltd.com",
        "Cc: joel.ngabo@medequipltd.com; melad@medequipltd.com",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(11)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    subj = doc.add_paragraph()
    subj.paragraph_format.space_after = Pt(10)
    add_heading_run(
        subj,
        "Subject: Application for Biomedical Engineer",
        12,
        True,
        RGBColor(0x0F, 0x2C, 0x4C),
    )

    for para in LETTER_BODY:
        if para == "Dear Hiring Manager,":
            body_para(doc, para, space_after=4)
        else:
            body_para(doc, para, space_after=9)

    body_para(doc, "Yours sincerely,", space_after=2)
    if SIGNATURE_SRC.exists():
        add_signature_docx(doc, width_in=1.65)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_before = Pt(2)
    name_p.paragraph_format.space_after = Pt(1)
    add_heading_run(name_p, "Wasswa Wilson", 10.5, True)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    r = title_p.add_run("Biomedical Engineer")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    add_docx_page_footer(doc)
    doc.save(path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CVName",
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=3,
            leading=24,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVTitle",
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=TEAL,
            alignment=TA_CENTER,
            spaceAfter=2,
            leading=15,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVYears",
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=TEAL,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVContact",
            fontName="Helvetica",
            fontSize=10,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=1,
            leading=13,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVSection",
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=2,
            leading=13,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBody",
            fontName="Helvetica",
            fontSize=10.5,
            textColor=GRAY,
            alignment=TA_JUSTIFY,
            leading=14,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVRole",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            textColor=GRAY,
            spaceBefore=5,
            spaceAfter=1,
            leading=13,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBullet",
            fontName="Helvetica",
            fontSize=10.5,
            textColor=GRAY,
            leftIndent=11,
            leading=13.5,
            spaceAfter=1.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVLabel",
            fontName="Helvetica",
            fontSize=10.5,
            textColor=GRAY,
            leading=13.5,
            spaceAfter=2.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterHead",
            fontName="Helvetica-Bold",
            fontSize=18,
            textColor=NAVY,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterSub",
            fontName="Helvetica-Bold",
            fontSize=11.5,
            textColor=TEAL,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterLine",
            fontName="Helvetica",
            fontSize=11,
            textColor=GRAY,
            spaceAfter=1,
            leading=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterBody",
            fontName="Helvetica",
            fontSize=11.5,
            textColor=GRAY,
            alignment=TA_JUSTIFY,
            leading=16,
            spaceAfter=9,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterLeft",
            fontName="Helvetica",
            fontSize=11.5,
            textColor=GRAY,
            alignment=TA_LEFT,
            leading=15,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVSignName",
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=GRAY,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=1,
            leading=13,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVSignTitle",
            fontName="Helvetica",
            fontSize=10.5,
            textColor=MUTED,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=2,
            leading=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterSignName",
            fontName="Helvetica-Bold",
            fontSize=11.5,
            textColor=GRAY,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=1,
            leading=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterSignTitle",
            fontName="Helvetica",
            fontSize=11,
            textColor=MUTED,
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=2,
            leading=13,
        )
    )
    return styles


def hr():
    return HRFlowable(
        width="100%",
        thickness=1,
        color=TEAL,
        spaceBefore=0,
        spaceAfter=5,
    )


def build_cv_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=11 * mm,
        bottomMargin=12 * mm,
    )
    email_phone = (
        f"{CONTACT['email']}&nbsp;&nbsp;&#183;&nbsp;&nbsp;{CONTACT['phone']}"
    )
    story = [
        Paragraph(CONTACT["name"], styles["CVName"]),
        Paragraph(CONTACT["title"], styles["CVTitle"]),
        Paragraph(CONTACT["years"], styles["CVYears"]),
        Paragraph(email_phone, styles["CVContact"]),
        Paragraph(CONTACT["location"], styles["CVContact"]),
        Paragraph(CONTACT["languages"], styles["CVContact"]),
        Spacer(1, 4),
        Paragraph("PROFILE", styles["CVSection"]),
        hr(),
        Paragraph(SUMMARY, styles["CVBody"]),
        Paragraph("SKILLS FOR THIS JOB", styles["CVSection"]),
        hr(),
    ]
    for label, value in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["CVLabel"]))

    story += [Paragraph("WORK EXPERIENCE", styles["CVSection"]), hr()]
    for role in ROLES:
        block = [
            Paragraph(
                f'{role["title"]} '
                f'<font color="#1A6B5C"><i>| {role["dates"]}</i></font>',
                styles["CVRole"],
            ),
        ]
        for b in role["bullets"]:
            block.append(Paragraph(f"&#8226; {b}", styles["CVBullet"]))
        story.append(KeepTogether(block))

    story += [Paragraph("EARLIER ROLES", styles["CVSection"]), hr()]
    for title, dates, note in EARLIER:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f'{title} <font color="#1A6B5C"><i>| {dates}</i></font>',
                        styles["CVRole"],
                    ),
                    Paragraph(f"&#8226; {note}", styles["CVBullet"]),
                ]
            )
        )

    story += [Paragraph("SELECTED HIGHLIGHTS", styles["CVSection"]), hr()]
    for item in ACHIEVEMENTS:
        story.append(Paragraph(f"&#8226; {item}", styles["CVBullet"]))

    story += [Paragraph("EQUIPMENT EXPERIENCE", styles["CVSection"]), hr()]
    for label, value in EQUIPMENT:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["CVLabel"]))

    story += [Paragraph("EDUCATION AND ACADEMIC LEVELS", styles["CVSection"]), hr()]
    for item in EDUCATION:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f'<font color="#1A6B5C"><b>{item["level"]}</b></font>',
                        styles["CVLabel"],
                    ),
                    Paragraph(item["school"], styles["CVRole"]),
                    Paragraph(item["detail"], styles["CVBody"]),
                ]
            )
        )

    story += [Paragraph("CERTIFICATES AND SHORT COURSES", styles["CVSection"]), hr()]
    for item in CERTIFICATIONS:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f'<font color="#1A6B5C"><b>{item["level"]}</b></font>',
                        styles["CVLabel"],
                    ),
                    Paragraph(item["school"], styles["CVRole"]),
                    Paragraph(item["detail"], styles["CVBody"]),
                ]
            )
        )

    story += [Paragraph("REFEREES", styles["CVSection"]), hr()]
    for ref in REFEREES:
        story.append(Paragraph(ref, styles["CVLabel"]))

    if SIGNATURE_SRC.exists():
        story.append(Spacer(1, 4))
        story.append(
            KeepTogether(
                [
                    signature_block_pdf(
                        styles, "CVSignName", "CVSignTitle", width_mm=32
                    )
                ]
            )
        )

    doc.build(story, canvasmaker=NumberedCanvas)


def build_letter_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph("WASSWA WILSON", styles["LetterHead"]),
        Paragraph("Biomedical Engineer · 7+ years experience", styles["LetterSub"]),
        Paragraph("Kampala, Uganda", styles["LetterLine"]),
        Paragraph(CONTACT["email"], styles["LetterLine"]),
        Paragraph(CONTACT["phone"], styles["LetterLine"]),
        Spacer(1, 4),
        HRFlowable(
            width="100%", thickness=1.2, color=TEAL, spaceBefore=2, spaceAfter=10
        ),
        Paragraph(LETTER_DATE, styles["LetterLine"]),
        Spacer(1, 8),
        Paragraph("Hiring Manager", styles["LetterLine"]),
        Paragraph("Medequip Healthcare Ltd", styles["LetterLine"]),
        Paragraph("Kigali, Rwanda", styles["LetterLine"]),
        Paragraph(
            "To: admin@medequipltd.com; finance.rwanda@medequipltd.com",
            styles["LetterLine"],
        ),
        Paragraph(
            "Cc: joel.ngabo@medequipltd.com; melad@medequipltd.com",
            styles["LetterLine"],
        ),
        Spacer(1, 10),
        Paragraph(
            "<b>Subject: Application for Biomedical Engineer</b>",
            styles["LetterLeft"],
        ),
        Spacer(1, 4),
    ]
    for para in LETTER_BODY:
        style = (
            styles["LetterLeft"]
            if para == "Dear Hiring Manager,"
            else styles["LetterBody"]
        )
        story.append(Paragraph(para, style))

    story.append(Paragraph("Yours sincerely,", styles["LetterLeft"]))
    if SIGNATURE_SRC.exists():
        story.append(Spacer(1, 2))
        story.append(
            signature_block_pdf(
                styles, "LetterSignName", "LetterSignTitle", width_mm=38
            )
        )
    else:
        story.append(Paragraph("Wasswa Wilson", styles["LetterSignName"]))
        story.append(Paragraph("Biomedical Engineer", styles["LetterSignTitle"]))
    doc.build(story, canvasmaker=NumberedCanvas)


def main() -> None:
    cv_docx = OUT / "Wasswa_Wilson_CV_Medequip_Biomedical_Engineer.docx"
    cv_pdf = OUT / "Wasswa_Wilson_CV_Medequip_Biomedical_Engineer.pdf"
    letter_docx = OUT / "Wasswa_Wilson_Cover_Letter_Medequip_Biomedical_Engineer.docx"
    letter_pdf = OUT / "Wasswa_Wilson_Cover_Letter_Medequip_Biomedical_Engineer.pdf"

    build_cv_docx(cv_docx)
    build_letter_docx(letter_docx)
    build_cv_pdf(cv_pdf)
    build_letter_pdf(letter_pdf)

    print("Wrote:")
    for p in (cv_docx, cv_pdf, letter_docx, letter_pdf):
        print(f"  {p}")


if __name__ == "__main__":
    main()
