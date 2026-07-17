"""
Photo-rich DOCX + PDF of refined_project_idea.md for AWIEF Pitch n Grow 2026.
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm, mm
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
    PageBreak, KeepTogether, ListFlowable, ListItem, HRFlowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

AWIEF = Path(__file__).resolve().parent
REPO = AWIEF.parents[1]
ASSETS = REPO / "assets"
OUT_DOC = AWIEF / "documents" / "awief_word.docx"
OUT_PDF = AWIEF / "documents" / "awief_pdf.pdf"

NAVY = "0A1F2E"
TEAL = "0D6E6E"
TEAL_L = "14A3A3"
ACCENT = "C45C26"
SLATE = "1E2F38"
MUTED = "3A4A54"
CREAM = "F7F5F0"
LINE = "D0DCDC"

# Curated real-photo gallery placements (conceptual diagrams kept where they teach the product).
PHOTOS = {
    "logo": "fairbanks_logo.jpeg",
    "facility_sign": "facility_exterior_sign.jpeg",
    "facility_entrance": "facility_exterior_entrance_01.jpg",
    "facility_entrance_2": "facility_exterior_entrance_02.jpg",
    "pharmacy_branded": "pharmacy_exterior_01.jpg",
    "pharmacy_interior": "pharmacy_interior_02.jpg",
    "staff_team": "staff_team_reception.jpeg",
    "mission_wall": "reception_mission_wall.jpeg",
    "waiting_branded": "waiting_room_reception_01.jpeg",
    "problem_clinic": "clinic_reception_desk_01.jpg",
    "problem_pharmacy": "pharmacy_exterior_01.jpg",
    "dashboard": "dashboard_demo.png",
    "architecture": "data_flow_iso_labeled.png",
    "deep_tech": "deep_tech_collage.png",
    "maternal": "bloom_maternal_health_participant_01.jpg",
    "gis": "gis_hotspots.png",
    "mobile": "outreach_mobile_phone_demo_01.jpg",
    "outreach_hero": "outreach_facilitator_canopy_01.jpg",
    "outreach_bp": "outreach_bp_screening.jpeg",
    "outreach_outdoor": "outreach_outdoor_clinic.jpeg",
    "training": "indoor_training_staff_presenting_01.jpg",
    "pharmacy_digital": "pharmacy_staff_laptop_01.jpg",
    "mothers_wait": "waiting_room_mothers_01.jpeg",
    "gericare": "gericare_wheelchair_assist.jpeg",
    "impact": "waiting_room_mothers_02.jpeg",
    "mvp_capture": "outreach_registration_form_01.jpg",
    "conclusion": "outreach_audience_full_group_01.jpg",
    "facility_building": "facility_exterior_building_01.jpg",
    "staff_field": "staff_outreach_conversation_01.jpg",
}


def asset(name: str) -> Path:
    p = ASSETS / name
    if not p.exists():
        raise FileNotFoundError(p)
    return p


def photo(key: str) -> Path:
    return asset(PHOTOS[key])


def _pil_size(path: Path):
    from PIL import Image as PILImage
    with PILImage.open(path) as im:
        return im.size


def fit_width(path: Path, max_width_in: float, max_height_in: float = 3.6) -> float:
    """Return an image width in inches that respects max width and height."""
    iw, ih = _pil_size(path)
    aspect = ih / float(iw)
    w = max_width_in
    h = w * aspect
    if h > max_height_in:
        w = max_height_in / aspect
    return w


def optimized_image(path: Path, max_px: int = 1600) -> Path:
    """Downscale large photos for DOCX/PDF embedding; leave diagrams untouched if already small."""
    from PIL import Image as PILImage

    cache_dir = REPO / "tmp" / "doc_assets"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out = cache_dir / f"{path.stem}_opt{path.suffix.lower()}"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out

    with PILImage.open(path) as im:
        im = im.convert("RGB") if im.mode not in ("RGB", "L") else im
        iw, ih = im.size
        scale = min(1.0, max_px / float(max(iw, ih)))
        if scale < 1.0:
            im = im.resize((int(iw * scale), int(ih * scale)), PILImage.Resampling.LANCZOS)
        save_kw = {"optimize": True}
        if out.suffix.lower() in (".jpg", ".jpeg"):
            save_kw["quality"] = 85
            if out.suffix.lower() == ".jpeg":
                out = cache_dir / f"{path.stem}_opt.jpg"
            im.save(out, format="JPEG", **save_kw)
        else:
            im.save(out, **save_kw)
    return out


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size=11, bold=False, color=SLATE, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, hex_color):
    tc = cell._tePr if False else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_border(cell, color=LINE, sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(borders)


def add_para(doc, text, size=11, bold=False, color=SLATE, space_after=8, space_before=0,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    run.italic = italic
    return p


def add_heading_custom(doc, text, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: NAVY, 2: TEAL, 3: SLATE}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, SLATE))
    return p


def add_image(doc, path, width_in=6.3, caption=None, max_height_in=3.8):
    path = optimized_image(Path(path))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    w = fit_width(path, width_in, max_height_in=max_height_in)
    run.add_picture(str(path), width=Inches(w))
    if caption:
        add_para(doc, caption, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
                 italic=True, space_after=12, space_before=2)


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.clear()
        run = p.add_run(item)
        set_run_font(run, size=size, color=SLATE)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color="FFFFFF")
        shade_cell(cell, TEAL)
        set_cell_border(cell, TEAL)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10, color=SLATE)
            if r_i % 2:
                shade_cell(cell, CREAM)
            set_cell_border(cell)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_two_images(doc, path_a, path_b, cap_a="", cap_b="", width=3.1, max_height_in=2.6):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (path, cap) in enumerate([(path_a, cap_a), (path_b, cap_b)]):
        path = optimized_image(Path(path))
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        w = fit_width(path, width, max_height_in=max_height_in)
        run.add_picture(str(path), width=Inches(w))
        cell2 = table.rows[1].cells[i]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(cap)
        set_run_font(run2, size=9, color=MUTED)
        run2.italic = True
    doc.add_paragraph()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Cover
    add_para(doc, "AWIEF Pitch n Grow 2026", size=12, bold=True, color=TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Grant & Pitch Application Write-Up", size=11, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add_para(doc, "FairBanks Community Health Intelligence Platform (FCHIP)", size=26, bold=True,
             color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Your health, our mission.", size=13, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, italic=True)
    add_para(doc, "Deep Roots. Digital Futures.", size=13, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(
        doc,
        "HealthTech  |  Artificial Intelligence & Machine Learning  |  Uganda -> East Africa -> Pan-African scale",
        size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    add_para(doc, "Application deadline: 20 July 2026", size=10, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_image(
        doc, photo("pharmacy_interior"), width_in=5.8, max_height_in=2.8,
        caption="FairBanks Pharmacy — Your health, our mission.",
    )
    add_two_images(
        doc,
        photo("facility_entrance"),
        photo("outreach_hero"),
        "Branded facility entrance — operating validation site",
        "Community outreach — last-mile engagement that feeds FCHIP",
        max_height_in=2.4,
    )

    add_para(
        doc,
        "Programme: https://awief.untap.us/pitch-n-grow2026",
        size=10, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
    )

    # 1. Executive Summary
    add_heading_custom(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "FairBanks is evolving from a community-rooted medical centre into Africa's leading "
        "Community Health Intelligence Company. Through the FairBanks Community Health "
        "Intelligence Platform (FCHIP), we are building a deep-technology system that connects "
        "community health workers, clinics, outreach programmes, and public-health "
        "decision-makers into one intelligent network.",
    )
    add_para(
        doc,
        "Today, primary healthcare in underserved African communities is largely reactive: "
        "facilities treat patients only after illness appears, with little visibility into what "
        "is happening in surrounding villages until outbreaks, complications, or stock-outs "
        "occur. FCHIP transforms this model by continuously ingesting community-generated health "
        "data and using artificial intelligence, machine learning, predictive analytics, and GIS "
        "mapping to forecast disease risk, maternal complications, chronic-disease hotspots, "
        "child-health threats, and medicine demand before crises escalate.",
    )
    add_para(
        doc,
        "FairBanks is uniquely positioned to build and validate this platform. We already operate "
        "a functioning medical centre, run active community outreach across Kampala-area "
        "communities (including Bukoto, Kyebando, Kisaasi, Kamwokya, and Kikaaya), work with "
        "Village Health Teams (VHTs) and community health workers (CHWs), and deliver maternal "
        "and child health, geriatric care (Gericare), chronic-disease screening, corporate and "
        "school health programmes, digital health records, and research partnerships.",
    )
    add_para(
        doc,
        "Vision: To build Africa's leading Community Health Intelligence Platform that harnesses "
        "artificial intelligence, community-generated data, and predictive analytics to transform "
        "primary healthcare from reactive treatment to proactive prevention.",
        bold=True,
    )
    add_para(
        doc,
        "Funding ask (indicative): Support to develop and pilot FCHIP MVP - mobile data collection "
        "for CHWs/VHTs, cloud sync, predictive analytics engine, and facility/district dashboards "
        "- validated in FairBanks' live community health ecosystem before district and regional "
        "scale-up.",
        italic=True,
    )

    # 2. Problem
    add_heading_custom(doc, "2. Problem Statement", 1)
    add_two_images(
        doc,
        photo("problem_clinic"),
        photo("problem_pharmacy"),
        "Facility reception - care still starts after patients arrive",
        "On-site FairBanks Pharmacy — branded medicine supply point",
    )
    add_heading_custom(doc, "2.1 The reactive healthcare gap", 2)
    add_para(
        doc,
        "Across underserved African communities, health systems face a structural information gap.",
    )
    add_table(
        doc,
        ["Current reality", "Consequence"],
        [
            ["Facilities wait for patients to become sick", "Late detection of outbreaks and complications"],
            ["Care ends when the patient leaves the facility", "No continuous view of community-level health trends"],
            ["Outreach data sits in paper registers or siloed systems", "District offices and NGOs lack real-time intelligence"],
            ["Medicine ordering is based on historical guesswork", "Stock-outs during seasonal disease surges"],
            ["High-risk pregnancies and NCDs are identified late", "Preventable maternal morbidity and stroke/diabetes burden"],
        ],
        col_widths=[3.2, 3.2],
    )
    add_heading_custom(doc, "2.2 Why existing approaches fall short", 2)
    add_para(
        doc,
        "Many digital health tools in Africa focus on single use cases - appointment booking, EMR "
        "digitisation, or SMS reminders. These are valuable but do not constitute deep technology "
        "or population-level intelligence. They rarely combine:",
    )
    add_bullets(doc, [
        "Multi-source community data (VHTs, schools, churches, pharmacies, outreach events)",
        "AI/ML-driven predictive modelling",
        "Geospatial disease mapping",
        "Real-time decision support for frontline workers and policy makers",
    ])
    add_para(
        doc,
        "The result is fragmented data, delayed response, and missed opportunities for prevention "
        "- exactly the challenge AWIEF's 2026 theme, Deep Roots. Digital Futures., calls founders "
        "to address with technology rooted in local context.",
    )

    # 3. Solution
    add_heading_custom(doc, "3. Solution: FairBanks Community Health Intelligence Platform (FCHIP)", 1)
    add_image(doc, photo("dashboard"), width_in=6.4,
              caption="FCHIP concept: facility and district health intelligence dashboards")
    add_heading_custom(doc, "3.1 Venture positioning", 2)
    add_table(
        doc,
        ["Instead of", "We present"],
        [[
            '"A private medical centre"',
            "A Community Health Intelligence Company using AI, predictive analytics, and "
            "community-generated data to improve health outcomes in underserved African communities",
        ]],
        col_widths=[2.4, 4.0],
    )
    add_para(
        doc,
        "FCHIP is FairBanks' deep-technology platform - connecting communities, healthcare "
        "providers, governments, and development partners through one intelligent system.",
    )

    add_heading_custom(doc, "3.2 What FCHIP does", 2)
    add_para(doc, "FCHIP continuously receives structured health information from:")
    add_bullets(doc, [
        "Community Health Workers and Village Health Teams (household visits, symptom checks, referrals)",
        "Patients and caregivers (via mobile applications and self-reporting)",
        "FairBanks medical centre and partner clinics (encounters, vitals, diagnoses, digital health records)",
        "Pharmacies, drug shops, and medicine outlets (dispensing patterns, stock levels, common complaints)",
        "Schools, churches, mosques, and community outreach events (screening camps, health talks)",
        "Maternal and child health touchpoints (ANC/PNC visits, immunisation posts, nutrition programmes)",
        "Corporate and workplace wellness programmes (BP, BMI, glucose, occupational health checks)",
        "Laboratories and point-of-care testing (malaria RDTs, HIV, Hb, glucose, and related results)",
        "Geriatric / Gericare and chronic-disease follow-up cohorts (hypertension, diabetes, elderly care)",
        "Public health and administrative systems (district HMIS/DHIS2 aggregates, where integration is approved)",
        "Environmental and contextual signals (rainfall/seasonality, flooding, sanitation reports)",
        "Research partners and programme M&E teams (anonymised field datasets)",
    ])
    add_para(
        doc,
        "The platform processes this data through an AI and analytics engine to predict - not "
        "merely record - health events and resource needs, then delivers actionable alerts and "
        "dashboards to the right users at the right time.",
    )

    add_heading_custom(doc, "3.3 Platform architecture", 2)
    add_image(doc, photo("architecture"), width_in=6.4,
              caption="Data sources -> Capture -> FCHIP intelligence -> Action in the community")
    add_para(
        doc,
        "FCHIP sits between field data capture and decision-makers. Data flows in from the "
        "community, is processed by intelligence modules (AI/ML, predictive analytics, GIS, "
        "clinical decision support), and returns as alerts, dashboards, and referrals to CHWs, "
        "health facilities, district offices, NGOs, and partners.",
    )
    add_table(
        doc,
        ["Layer", "Role"],
        [
            ["Data sources", "Who generates health signals in the community and at facilities"],
            ["Capture", "How those signals are entered and synced (mobile/offline + facility systems)"],
            ["FCHIP core", "How signals become predictions, maps, and clinical guidance"],
            ["Consumers", "Who acts - CHWs (field), facilities (care), districts/partners (programmes)"],
        ],
        col_widths=[1.8, 4.6],
    )
    add_para(
        doc,
        "FCHIP becomes the digital nervous system for community healthcare - moving data from the "
        "last mile to decision-makers in near real time, and closing the loop with actionable "
        "alerts back to the field.",
        bold=True,
    )

    # 4. Deep tech
    add_heading_custom(doc, "4. Deep Technology Core", 1)
    add_image(doc, photo("deep_tech"), width_in=6.4,
              caption="Integrated deep-tech stack: AI, ML, GIS, mobile edge, cloud, NLP")
    add_para(
        doc,
        "AWIEF 2026 requires ventures built on substantial research and engineering, not "
        "lightweight digital services. FCHIP meets this standard through an integrated deep-tech stack:",
    )
    add_table(
        doc,
        ["Technology", "Function in FCHIP"],
        [
            ["Artificial Intelligence", "Disease-risk prediction, outbreak early warning, maternal and NCD risk scoring"],
            ["Machine Learning", "Pattern learning from historical community health data, seasonal trends, and outreach outcomes"],
            ["GIS Mapping", "Geospatial visualisation of disease distribution, hotspots, and resource gaps"],
            ["Mobile Data Collection", "Offline-capable CHW/VHT apps for structured capture at household and community level"],
            ["Cloud Computing", "Secure synchronisation across facilities, partners, and administrative levels"],
            ["Analytics Dashboard", "Real-time trend monitoring for clinicians, programme managers, and district officials"],
            ["Natural Language Processing", "Local-language symptom reporting and summarisation where appropriate"],
        ],
        col_widths=[2.2, 4.2],
    )

    # 5. Use cases
    add_heading_custom(doc, "5. Predictive Use Cases", 1)
    add_two_images(
        doc,
        photo("maternal"),
        photo("gis"),
        "Maternal and child health - community programme signals",
        "Disease surveillance - GIS hotspot early warning",
    )
    add_image(doc, photo("mobile"), width_in=5.8,
              caption="Mobile-first field engagement - phones as last-mile data tools for FCHIP")

    use_cases = [
        ("5.1 Disease surveillance",
         "Signal: Three neighbouring villages report increasing fever cases via VHT mobile entries. "
         "Prediction: Possible malaria outbreak within 14 days. "
         "Action: District office and FairBanks outreach teams deploy targeted testing, bed-net "
         "distribution, and pharmacy pre-stocking."),
        ("5.2 Maternal health",
         "Signal: Pregnant mothers' home-visit data (BP, haemoglobin proxies, visit adherence). "
         "Prediction: High-risk pregnancy, missed antenatal care, anaemia risk, possible pre-eclampsia. "
         "Action: CHWs receive alerts; mothers are referred before complications occur."),
        ("5.3 Non-communicable diseases (NCDs)",
         "Signal: Blood pressure readings from community screening and corporate/school programmes. "
         "Prediction: Communities with rising hypertension, diabetes hotspots, elevated stroke risk. "
         "Action: Targeted screening campaigns and lifestyle intervention planning."),
        ("5.4 Child health",
         "Signal: Growth monitoring, immunisation records, diarrhoea incidence. "
         "Prediction: Rising malnutrition, low immunisation coverage, diarrhoeal disease clusters. "
         "Action: Community flags trigger nutrition support and immunisation drives."),
        ("5.5 Medicine demand forecasting",
         "Signal: Disease trends, rainfall patterns, historical consumption, outreach schedules. "
         "Prediction: Medicine demand by facility and community before stock-outs. "
         "Action: Procurement and pharmacy replenishment aligned to forecasted need."),
    ]
    for title, body in use_cases:
        add_heading_custom(doc, title, 3)
        add_para(doc, body)

    # 6. Market
    add_heading_custom(doc, "6. Target Market & Customers", 1)
    add_table(
        doc,
        ["Customer segment", "Value delivered"],
        [
            ["Medical centres & clinics", "Better follow-up, outreach planning, and population health visibility"],
            ["District health offices", "Population-level disease intelligence and early warning"],
            ["NGOs & development partners", "Real-time M&E, impact evidence, programme optimisation"],
            ["Community health workers / VHTs", "Mobile tools, structured workflows, decision support"],
            ["Insurance companies", "Prevention-focused population health insights"],
            ["Ministries of health", "National and sub-national planning, outbreak preparedness"],
            ["Research institutions", "Ethical, anonymised datasets for community health research"],
        ],
        col_widths=[2.6, 3.8],
    )
    add_para(
        doc,
        "Primary market (Phase 1): Uganda - starting with FairBanks' established community "
        "catchment and Kampala metropolitan outreach zones. Expansion markets (Phase 2-3): East "
        "Africa, then additional African countries with similar primary healthcare and CHW/VHT "
        "infrastructure.",
    )

    # 7. Business model
    add_heading_custom(doc, "7. Business Model & Revenue Streams", 1)
    add_para(doc, "FCHIP is designed for a diversified, scalable revenue model:")
    add_bullets(doc, [
        "Subscription licences for clinics and hospitals",
        "District health office deployments (SaaS + implementation)",
        "NGO programme monitoring contracts",
        "Ministry of Health national/sub-national implementations",
        "Custom analytics and reporting for partners",
        "Research collaborations with universities",
        "API integrations for digital health ecosystem partners",
        "Training and certification for community health workers on platform use",
    ])

    # 8. Traction
    add_heading_custom(doc, "8. Traction, Foundation & Competitive Advantage", 1)
    add_two_images(
        doc,
        photo("facility_sign"),
        photo("facility_entrance"),
        "FairBanks Medical Centre signage — brand identity in the community",
        "Facility entrance — live operating foundation for FCHIP validation",
        max_height_in=2.5,
    )
    add_image(doc, photo("outreach_hero"), width_in=6.4,
              caption="Live FairBanks Community Reach - canopy outreach in Kampala communities")
    add_heading_custom(doc, "8.1 Existing FairBanks ecosystem", 2)
    add_bullets(doc, [
        "FairBanks Community Reach Programme",
        "CHW and VHT engagement in Bukoto, Kyebando, Kisaasi, Kamwokya, Kikaaya, and surrounding communities",
        "Maternal and child health initiatives",
        "Gericare (geriatric care programme)",
        "Chronic disease screening",
        "Corporate and school health programmes",
        "Digital health records (existing digitisation foundation)",
        "Research and community partnerships",
        "Functioning medical centre with direct patient and community access",
    ])
    add_heading_custom(doc, "8.1.1 Evidence from the field", 3)
    add_para(
        doc,
        "Real FairBanks operations already generate the community signals FCHIP will turn into "
        "predictions - outreach screening, maternal programmes, pharmacy dispensing, facility "
        "encounters, and staff digital workflows.",
        space_after=6,
    )
    add_two_images(
        doc,
        photo("staff_team"),
        photo("mission_wall"),
        "FairBanks staff in branded uniforms at reception",
        "Mission, vision & values wall — FairBanks Medical Centre identity",
        max_height_in=2.6,
    )
    add_two_images(
        doc,
        photo("pharmacy_branded"),
        photo("waiting_branded"),
        "FairBanks Pharmacy exterior — Medical Centre branding",
        "Waiting area with FairBanks mission branding visible",
        max_height_in=2.6,
    )
    add_two_images(
        doc,
        photo("outreach_bp"),
        photo("outreach_outdoor"),
        "BP screening and registration under outreach canopies",
        "Outdoor community clinic - structured field capture opportunity",
    )
    add_two_images(
        doc,
        photo("training"),
        photo("pharmacy_digital"),
        "Staff training and product briefings for community programmes",
        "Pharmacy digital workflow - stock and dispensing signals for FCHIP",
    )
    add_two_images(
        doc,
        photo("mothers_wait"),
        photo("gericare"),
        "Maternal and child health touchpoints at the facility",
        "Gericare - continuous support for older community members",
        max_height_in=2.9,
    )
    add_two_images(
        doc,
        photo("facility_building"),
        photo("staff_field"),
        "FairBanks Medical Centre - operating validation site",
        "Staff-community conversation - trust that enables data quality",
    )
    add_heading_custom(doc, "8.2 Why FairBanks wins", 2)
    add_table(
        doc,
        ["Typical startup", "Typical health facility", "FairBanks"],
        [[
            "Technology without field access",
            "Patients without technology",
            "Both: live healthcare operation + deep-tech build capability",
        ], [
            "Imported solutions",
            "Manual, reactive workflows",
            "Context-rooted AI built from African community health realities",
        ], [
            "Pilot without validation site",
            "Data without intelligence layer",
            "Design -> pilot -> validate -> refine in a real setting before scale",
        ]],
        col_widths=[2.1, 2.1, 2.2],
    )
    add_heading_custom(doc, "8.3 Recommended AWIEF track", 2)
    add_para(
        doc,
        "Startup Track (recommended): Strong fit - registered operating medical centre, "
        "revenue-generating services, active community programmes, building deep-tech platform "
        "on proven field operations.",
        bold=True,
    )

    # 9. Impact
    add_heading_custom(doc, "9. Social Impact & Development Alignment", 1)
    add_image(doc, photo("impact"), width_in=5.2, max_height_in=4.2,
              caption="Impact focus: earlier intervention for mothers, children, and communities")
    add_para(
        doc,
        "FCHIP shifts primary healthcare from sick-care to predictive, community-centred "
        "prevention - directly benefiting underserved populations who depend on CHWs, outreach, "
        "and under-resourced facilities.",
    )
    add_bullets(doc, [
        "Earlier outbreak detection and reduced morbidity from communicable diseases",
        "Fewer maternal and neonatal complications through risk-based alerting",
        "Reduced NCD burden via hotspot-targeted screening",
        "Improved child nutrition and immunisation coverage through community-level monitoring",
        "Reduced medicine stock-outs and more efficient public-health resource allocation",
        "Evidence base for NGOs, donors, and governments to invest in what works",
    ])
    add_table(
        doc,
        ["SDG", "FCHIP contribution"],
        [
            ["SDG 3 - Good Health and Well-Being", "Predictive community health, maternal/child/NCD focus"],
            ["SDG 5 - Gender Equality", "Women-led venture; maternal health intelligence"],
            ["SDG 9 - Industry, Innovation and Infrastructure", "Deep-tech health infrastructure for last-mile Africa"],
            ["SDG 10 - Reduced Inequalities", "Intelligence for underserved communities"],
            ["SDG 17 - Partnerships", "Multi-stakeholder health ecosystem integration"],
        ],
        col_widths=[3.0, 3.4],
    )

    # 10. MVP
    add_heading_custom(doc, "10. MVP Scope & Product Roadmap", 1)
    add_two_images(
        doc,
        photo("mvp_capture"),
        photo("gis"),
        "MVP: structured field registration and mobile capture",
        "MVP: GIS risk visualisation layer",
    )
    add_heading_custom(doc, "10.1 Minimum Viable Product (MVP)", 2)
    add_para(doc, "Goal: Prove predictive community health intelligence in FairBanks' live catchment.")
    add_table(
        doc,
        ["MVP component", "Description"],
        [
            ["CHW/VHT mobile app", "Offline data capture: symptoms, vitals, maternal/child indicators, household visits"],
            ["Cloud sync & data pipeline", "Secure ingestion, validation, and storage from field to platform"],
            ["Analytics dashboard", "Facility-level view of trends, alerts, and outreach priorities"],
            ["AI Module v1", "Rule-based + ML-assisted risk scoring for 2-3 priority use cases"],
            ["GIS layer v1", "Map visualisation of reported cases and risk zones by village/parish"],
            ["Pilot integration", "Connect existing FairBanks digital records and outreach workflows"],
        ],
        col_widths=[2.2, 4.2],
    )
    add_heading_custom(doc, "10.2 Roadmap", 2)
    add_table(
        doc,
        ["Phase", "Timeline", "Milestones"],
        [
            ["Phase 1 - Pilot", "Months 0-12", "MVP with FairBanks CHWs/VHTs; validate 3 use cases"],
            ["Phase 2 - District scale", "Months 12-24", "Partner clinics; Kampala district structures; NGO M&E; APIs"],
            ["Phase 3 - National & regional", "Months 24-36", "Multi-district Uganda; East Africa; advanced ML & demand forecasting"],
            ["Phase 4 - Platform expansion", "Year 3+", "CDS, local-language NLP, research modules"],
        ],
        col_widths=[2.0, 1.5, 2.9],
    )

    # 11. Alignment
    add_heading_custom(doc, "11. AWIEF Pitch n Grow 2026 Alignment", 1)
    add_table(
        doc,
        ["AWIEF expectation", "FairBanks / FCHIP response"],
        [
            ["Deep knowledge of local context", "Years of outreach in named Kampala communities; CHW/VHT relationships"],
            ["Deep understanding of communities", "Platform designed around how African communities report, move, and access care"],
            ["Deep technology, not superficial digital", "AI + ML + GIS + mobile edge + cloud + NLP integrated platform"],
            ["African problems, African-built solutions", "Technology developed from FairBanks' field operations"],
        ],
        col_widths=[2.6, 3.8],
    )
    add_bullets(doc, [
        "Sector: HealthTech (priority innovation sector)",
        "Deep-tech core: Artificial Intelligence & Machine Learning",
        "Market focus: African communities (Uganda pilot -> continental scale)",
        "Founder eligibility: Woman founder or co-founder (presenting founder must be female)",
        "Commitment: Available for AWIEF 2026 Conference live finals, Cape Town, 10-11 November 2026",
    ])

    # 12. Application copy
    add_heading_custom(doc, "12. Application-Ready Copy Blocks", 1)
    add_heading_custom(doc, "12.1 One-line pitch", 3)
    add_para(
        doc,
        "FairBanks Community Health Intelligence Platform (FCHIP) is a deep-tech platform that uses "
        "AI and predictive analytics to transform African primary healthcare from reactive "
        "treatment to proactive, community-level prevention.",
        italic=True,
    )
    add_heading_custom(doc, "12.2 Unique value proposition", 3)
    add_para(
        doc,
        "Only venture combining a live medical centre and community outreach engine with a "
        "deep-tech predictive health platform - allowing real-world design, pilot, and validation "
        "before district and national scale.",
        bold=True,
    )

    # 13. Risks
    add_heading_custom(doc, "13. Risks & Mitigation", 1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Data quality from field capture", "Structured mobile forms, validation rules, CHW training and certification"],
            ["Low CHW digital literacy", "Simple UX, offline mode, local-language support, supervised rollout"],
            ["Privacy and consent", "Ethical data governance, anonymisation, Uganda Data Protection compliance"],
            ["Model accuracy in early phase", "Start with rule-based + limited ML; validate against FairBanks clinical outcomes"],
            ["Adoption by public sector", "Pilot evidence; NGO and district co-design; align with MoH community health strategy"],
        ],
        col_widths=[2.4, 4.0],
    )

    # 14. Conclusion
    add_heading_custom(doc, "14. Conclusion", 1)
    add_image(doc, photo("conclusion"), width_in=6.4,
              caption="From community signals to life-saving predictions - rooted in FairBanks outreach")
    add_para(
        doc,
        "FairBanks Community Health Intelligence Platform represents a strategic evolution from community "
        "medical practice to community health intelligence - a venture that matches AWIEF Pitch n "
        "Grow 2026's call for women-led, deep-technology solutions rooted in African reality.",
    )
    add_para(
        doc,
        "With an operating medical centre, established outreach across Kampala communities, "
        "trusted CHW/VHT relationships, digital health foundations, and a clear AI-powered product "
        "roadmap, FairBanks is ready to build, pilot, and scale FCHIP as the platform that connects "
        "community-generated data to life-saving predictions.",
    )
    add_para(doc, "Next immediate steps:", bold=True, space_before=8)
    add_bullets(doc, [
        "Finalise MVP scope and AI priorities for Phase 1 (surveillance, maternal risk, NCD hotspots)",
        "Complete AWIEF application and 10-12 slide pitch deck",
        "Record 3-minute technology demo video",
        "Launch FairBanks pilot with CHW/VHT cohort in primary catchment communities",
    ])
    add_para(
        doc,
        "Document prepared for AWIEF Pitch n Grow 2026 submission. Source: refined_project_idea.md",
        size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, italic=True,
    )

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT_DOC))
        print(f"DOCX: {OUT_DOC}")
        return OUT_DOC
    except PermissionError:
        alt_dir = REPO / "tmp"
        alt_dir.mkdir(parents=True, exist_ok=True)
        alt = alt_dir / (OUT_DOC.stem + "_unlocked" + OUT_DOC.suffix)
        doc.save(str(alt))
        print(f"DOCX locked; saved as: {alt}")
        return alt


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def build_pdf():
    navy = HexColor("#" + NAVY)
    teal = HexColor("#" + TEAL)
    accent = HexColor("#" + ACCENT)
    slate = HexColor("#" + SLATE)
    muted = HexColor("#" + MUTED)
    cream = HexColor("#" + CREAM)
    line = HexColor("#" + LINE)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CoverTitle", fontName="Helvetica-Bold", fontSize=22, leading=26,
        textColor=navy, alignment=TA_CENTER, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="CoverSub", fontName="Helvetica-Bold", fontSize=13, leading=17,
        textColor=teal, alignment=TA_CENTER, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H1Custom", fontName="Helvetica-Bold", fontSize=16, leading=20,
        textColor=navy, spaceBefore=16, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H2Custom", fontName="Helvetica-Bold", fontSize=12, leading=15,
        textColor=teal, spaceBefore=10, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="H3Custom", fontName="Helvetica-Bold", fontSize=11, leading=14,
        textColor=slate, spaceBefore=8, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="BodyCustom", fontName="Helvetica", fontSize=10, leading=14,
        textColor=slate, alignment=TA_JUSTIFY, spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="BodyBold", fontName="Helvetica-Bold", fontSize=10, leading=14,
        textColor=slate, alignment=TA_JUSTIFY, spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="Caption", fontName="Helvetica-Oblique", fontSize=8, leading=10,
        textColor=muted, alignment=TA_CENTER, spaceAfter=10, spaceBefore=2,
    ))
    styles.add(ParagraphStyle(
        name="BulletText", fontName="Helvetica", fontSize=10, leading=13,
        textColor=slate, leftIndent=12, spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name="Meta", fontName="Helvetica", fontSize=9, leading=12,
        textColor=muted, alignment=TA_CENTER, spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="CellHead", fontName="Helvetica-Bold", fontSize=8, leading=10,
        textColor=white,
    ))
    styles.add(ParagraphStyle(
        name="CellBody", fontName="Helvetica", fontSize=8, leading=10,
        textColor=slate,
    ))

    story = []
    page_w = A4[0] - 1.6 * inch

    def img(name, w=page_w, caption=None, max_h=3.4 * inch):
        raw = asset(name) if str(name).endswith((".png", ".jpg", ".jpeg")) else photo(name)
        path = str(optimized_image(raw))
        from PIL import Image as PILImage
        with PILImage.open(path) as pi:
            iw, ih = pi.size
        aspect = ih / float(iw)
        h = w * aspect
        if h > max_h:
            h = max_h
            w = h / aspect
        im = Image(path, width=w, height=h)
        block = [im]
        if caption:
            block.append(Paragraph(caption, styles["Caption"]))
        else:
            block.append(Spacer(1, 6))
        return KeepTogether(block)

    def h1(t):
        story.append(Paragraph(t, styles["H1Custom"]))
        story.append(HRFlowable(width="100%", thickness=1.2, color=teal, spaceAfter=8))

    def h2(t):
        story.append(Paragraph(t, styles["H2Custom"]))

    def h3(t):
        story.append(Paragraph(t, styles["H3Custom"]))

    def body(t, bold=False):
        story.append(Paragraph(t, styles["BodyBold"] if bold else styles["BodyCustom"]))

    def bullets(items):
        for it in items:
            story.append(Paragraph(f"•  {it}", styles["BulletText"]))
        story.append(Spacer(1, 4))

    def table(headers, rows, col_widths=None):
        data = [[Paragraph(h, styles["CellHead"]) for h in headers]]
        for row in rows:
            data.append([Paragraph(str(c), styles["CellBody"]) for c in row])
        if not col_widths:
            col_widths = [page_w / len(headers)] * len(headers)
        t = Table(data, colWidths=col_widths, repeatRows=1)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), teal),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("GRID", (0, 0), (-1, -1), 0.4, line),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, cream]),
        ]
        t.setStyle(TableStyle(style_cmds))
        story.append(t)
        story.append(Spacer(1, 10))

    def two_imgs(a, b, ca="", cb="", max_h=2.2 * inch):
        wa = (page_w - 10) / 2
        from PIL import Image as PILImage

        def resolve(name):
            raw = asset(name) if str(name).endswith((".png", ".jpg", ".jpeg")) else photo(name)
            return str(optimized_image(raw))

        def make(path_key, cap):
            p = resolve(path_key)
            with PILImage.open(p) as pi:
                iw, ih = pi.size
            aspect = ih / float(iw)
            h = min(wa * aspect, max_h)
            w = h / aspect if wa * aspect > max_h else wa
            elems = [Image(p, width=w, height=h)]
            if cap:
                elems.append(Paragraph(cap, styles["Caption"]))
            return elems

        left = make(a, ca)
        right = make(b, cb)
        # pad to equal length
        while len(left) < len(right):
            left.append(Spacer(1, 1))
        while len(right) < len(left):
            right.append(Spacer(1, 1))
        t = Table([[left, right]], colWidths=[page_w / 2, page_w / 2])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(t)
        story.append(Spacer(1, 6))

    # Cover
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("AWIEF Pitch n Grow 2026", styles["Meta"]))
    story.append(Paragraph("Grant &amp; Pitch Application Write-Up", styles["Meta"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "FairBanks Community Health Intelligence Platform (FCHIP)", styles["CoverTitle"]
    ))
    story.append(Paragraph(
        '<font color="#C45C26"><b><i>Your health, our mission.</i></b></font>',
        styles["Meta"],
    ))
    story.append(Paragraph(
        '<font color="#C45C26"><b>Deep Roots. Digital Futures.</b></font>',
        styles["Meta"],
    ))
    story.append(Paragraph(
        "HealthTech  |  AI &amp; Machine Learning  |  Uganda → East Africa → Pan-African scale",
        styles["Meta"],
    ))
    story.append(Paragraph("Application deadline: 20 July 2026", styles["Meta"]))
    story.append(Spacer(1, 12))
    story.append(img(
        "pharmacy_interior", w=page_w * 0.85, max_h=2.4 * inch,
        caption="FairBanks Pharmacy — Your health, our mission.",
    ))
    two_imgs(
        "facility_entrance", "outreach_hero",
        "Branded facility entrance — operating validation site",
        "Community outreach — last-mile engagement that feeds FCHIP",
        max_h=2.1 * inch,
    )
    story.append(Paragraph(
        '<link href="https://awief.untap.us/pitch-n-grow2026">'
        "https://awief.untap.us/pitch-n-grow2026</link>",
        styles["Meta"],
    ))
    story.append(PageBreak())

    # 1
    h1("1. Executive Summary")
    body(
        "FairBanks is evolving from a community-rooted medical centre into Africa's leading "
        "Community Health Intelligence Company. Through the FairBanks Community Health "
        "Intelligence Platform (FCHIP), we are building a deep-technology system that connects "
        "community health workers, clinics, outreach programmes, and public-health "
        "decision-makers into one intelligent network."
    )
    body(
        "Today, primary healthcare in underserved African communities is largely reactive: "
        "facilities treat patients only after illness appears, with little visibility into what "
        "is happening in surrounding villages until outbreaks, complications, or stock-outs "
        "occur. FCHIP transforms this model by continuously ingesting community-generated health "
        "data and using artificial intelligence, machine learning, predictive analytics, and GIS "
        "mapping to forecast disease risk, maternal complications, chronic-disease hotspots, "
        "child-health threats, and medicine demand before crises escalate."
    )
    body(
        "FairBanks already operates a functioning medical centre and active community outreach "
        "across Kampala-area communities (Bukoto, Kyebando, Kisaasi, Kamwokya, and Kikaaya), "
        "working with VHTs and CHWs on maternal and child health, Gericare, chronic-disease "
        "screening, corporate and school programmes, digital health records, and research "
        "partnerships. FCHIP is the missing technology layer that unifies these activities."
    )
    body(
        "Vision: To build Africa's leading Community Health Intelligence Platform that harnesses "
        "artificial intelligence, community-generated data, and predictive analytics to transform "
        "primary healthcare from reactive treatment to proactive prevention.",
        bold=True,
    )
    body(
        "Funding ask (indicative): Support to develop and pilot FCHIP MVP - mobile data collection "
        "for CHWs/VHTs, cloud sync, predictive analytics engine, and facility/district dashboards "
        "- validated in FairBanks' live community health ecosystem before district and regional "
        "scale-up."
    )

    # 2
    h1("2. Problem Statement")
    two_imgs(
        "problem_clinic", "problem_pharmacy",
        "Facility reception - care still starts after patients arrive",
        "On-site FairBanks Pharmacy — branded medicine supply point",
    )
    h2("2.1 The reactive healthcare gap")
    body("Across underserved African communities, health systems face a structural information gap.")
    table(
        ["Current reality", "Consequence"],
        [
            ["Facilities wait for patients to become sick", "Late detection of outbreaks and complications"],
            ["Care ends when the patient leaves the facility", "No continuous view of community-level health trends"],
            ["Outreach data sits in paper registers or siloed systems", "District offices and NGOs lack real-time intelligence"],
            ["Medicine ordering is based on historical guesswork", "Stock-outs during seasonal disease surges"],
            ["High-risk pregnancies and NCDs are identified late", "Preventable maternal morbidity and stroke/diabetes burden"],
        ],
        [page_w * 0.5, page_w * 0.5],
    )
    h2("2.2 Why existing approaches fall short")
    body(
        "Many digital health tools focus on single use cases - booking, EMR digitisation, or SMS "
        "reminders. Valuable, but not deep technology or population-level intelligence. They rarely combine:"
    )
    bullets([
        "Multi-source community data (VHTs, schools, churches, pharmacies, outreach)",
        "AI/ML-driven predictive modelling",
        "Geospatial disease mapping",
        "Real-time decision support for frontline workers and policy makers",
    ])
    body(
        "The result is fragmented data and missed prevention - exactly the challenge AWIEF's 2026 "
        "theme, Deep Roots. Digital Futures., calls founders to address."
    )

    # 3
    h1("3. Solution: FCHIP")
    story.append(img("dashboard", caption="FCHIP concept: facility and district intelligence dashboards"))
    h2("3.1 Venture positioning")
    table(
        ["Instead of", "We present"],
        [[
            '"A private medical centre"',
            "A Community Health Intelligence Company using AI, predictive analytics, and "
            "community-generated data to improve health outcomes in underserved African communities",
        ]],
        [page_w * 0.35, page_w * 0.65],
    )
    body(
        "FCHIP is FairBanks' deep-technology platform - connecting communities, healthcare "
        "providers, governments, and development partners."
    )
    h2("3.2 What FCHIP does")
    body("FCHIP continuously receives structured health information from:")
    bullets([
        "CHWs / VHTs (household visits, symptom checks, referrals)",
        "Patients and caregivers (mobile apps and self-reporting)",
        "FairBanks medical centre and partner clinics",
        "Pharmacies, drug shops, and medicine outlets",
        "Schools, churches, mosques, and community outreach events",
        "Maternal and child health touchpoints (ANC/PNC, immunisation, nutrition)",
        "Corporate and workplace wellness programmes",
        "Laboratories and point-of-care testing",
        "Gericare and chronic-disease follow-up cohorts",
        "Public health systems (HMIS/DHIS2 where approved)",
        "Environmental signals (rainfall, flooding, sanitation)",
        "Research partners and NGO M&E datasets",
    ])
    body(
        "The platform processes this data through an AI and analytics engine to predict - not "
        "merely record - health events and resource needs, then delivers actionable alerts and "
        "dashboards to the right users at the right time."
    )
    h2("3.3 Platform architecture")
    story.append(img("architecture",
                     caption="Data sources → Capture → FCHIP intelligence → Action"))
    body(
        "FCHIP sits between field data capture and decision-makers. Intelligence modules (AI/ML, "
        "predictive analytics, GIS, clinical decision support) return alerts and dashboards to "
        "CHWs, facilities, districts, NGOs, and partners."
    )
    table(
        ["Layer", "Role"],
        [
            ["Data sources", "Who generates health signals in the community and at facilities"],
            ["Capture", "How signals are entered and synced (mobile/offline + facility systems)"],
            ["FCHIP core", "How signals become predictions, maps, and clinical guidance"],
            ["Consumers", "Who acts - CHWs, facilities, districts/partners"],
        ],
        [page_w * 0.28, page_w * 0.72],
    )

    # 4
    h1("4. Deep Technology Core")
    story.append(img("deep_tech", max_h=2.6 * inch,
                     caption="Integrated stack: AI, ML, GIS, mobile edge, cloud, NLP"))
    body(
        "AWIEF 2026 requires ventures built on substantial research and engineering, not "
        "lightweight digital services. FCHIP meets this standard through:"
    )
    table(
        ["Technology", "Function in FCHIP"],
        [
            ["Artificial Intelligence", "Disease-risk prediction, outbreak early warning, maternal and NCD risk scoring"],
            ["Machine Learning", "Pattern learning from community health data, seasonal trends, outreach outcomes"],
            ["GIS Mapping", "Geospatial visualisation of disease distribution, hotspots, resource gaps"],
            ["Mobile Data Collection", "Offline-capable CHW/VHT apps for household and community capture"],
            ["Cloud Computing", "Secure synchronisation across facilities, partners, administrative levels"],
            ["Analytics Dashboard", "Real-time monitoring for clinicians, programme managers, district officials"],
            ["Natural Language Processing", "Local-language symptom reporting and summarisation"],
        ],
        [page_w * 0.32, page_w * 0.68],
    )

    # 5
    story.append(PageBreak())
    h1("5. Predictive Use Cases")
    two_imgs(
        "maternal", "gis",
        "Maternal and child health - community programme signals",
        "Disease surveillance - GIS hotspot early warning",
    )
    story.append(img("mobile", w=page_w * 0.9, max_h=2.8 * inch,
                     caption="Mobile-first field engagement - phones as last-mile data tools for FCHIP"))
    for title, text in [
        ("5.1 Disease surveillance",
         "Signal: Neighbouring villages report rising fever via VHT entries. Prediction: Possible "
         "malaria outbreak within 14 days. Action: Targeted testing, bed-nets, pharmacy pre-stocking."),
        ("5.2 Maternal health",
         "Signal: Home-visit BP, haemoglobin proxies, ANC adherence. Prediction: High-risk pregnancy "
         "and complication risk. Action: CHW alerts and early referral."),
        ("5.3 Non-communicable diseases",
         "Signal: Community and workplace BP/glucose. Prediction: Hypertension and diabetes hotspots. "
         "Action: Targeted screening and lifestyle interventions."),
        ("5.4 Child health",
         "Signal: Growth monitoring, immunisation, diarrhoea incidence. Prediction: Malnutrition and "
         "coverage gaps. Action: Nutrition support and immunisation drives."),
        ("5.5 Medicine demand forecasting",
         "Signal: Disease trends, rainfall, historical consumption. Prediction: Facility demand before "
         "stock-outs. Action: Procurement aligned to forecasted need."),
    ]:
        h3(title)
        body(text)

    # 6-7
    h1("6. Target Market &amp; Customers")
    table(
        ["Customer segment", "Value delivered"],
        [
            ["Medical centres & clinics", "Follow-up, outreach planning, population health visibility"],
            ["District health offices", "Disease intelligence and early warning"],
            ["NGOs & development partners", "Real-time M&E, impact evidence, programme optimisation"],
            ["CHWs / VHTs", "Mobile tools, structured workflows, decision support"],
            ["Insurance companies", "Prevention-focused population health insights"],
            ["Ministries of health", "Planning and outbreak preparedness"],
            ["Research institutions", "Ethical, anonymised community health datasets"],
        ],
        [page_w * 0.38, page_w * 0.62],
    )
    body(
        "Primary market (Phase 1): Uganda - FairBanks catchment and Kampala metro. Expansion "
        "(Phase 2-3): East Africa, then broader African markets with CHW/VHT infrastructure."
    )

    h1("7. Business Model &amp; Revenue Streams")
    bullets([
        "Subscription licences for clinics and hospitals",
        "District health office deployments (SaaS + implementation)",
        "NGO programme monitoring contracts",
        "Ministry of Health national/sub-national implementations",
        "Custom analytics and reporting for partners",
        "Research collaborations with universities",
        "API integrations for digital health partners",
        "CHW training and certification on platform use",
    ])

    # 8
    h1("8. Traction, Foundation &amp; Competitive Advantage")
    two_imgs(
        "facility_sign", "facility_entrance",
        "FairBanks Medical Centre signage — brand identity in the community",
        "Facility entrance — live operating foundation for FCHIP validation",
        max_h=2.2 * inch,
    )
    story.append(img("outreach_hero",
                     caption="Live FairBanks Community Reach - canopy outreach in Kampala communities"))
    h2("8.1 Existing FairBanks ecosystem")
    bullets([
        "FairBanks Community Reach Programme",
        "CHW/VHT engagement: Bukoto, Kyebando, Kisaasi, Kamwokya, Kikaaya+",
        "Maternal and child health initiatives",
        "Gericare (geriatric care) and chronic disease screening",
        "Corporate and school health programmes",
        "Digital health records foundation",
        "Research and community partnerships",
        "Functioning medical centre with patient and community access",
    ])
    h3("8.1.1 Evidence from the field")
    body(
        "Real FairBanks operations already generate the community signals FCHIP will turn into "
        "predictions - outreach screening, maternal programmes, pharmacy dispensing, facility "
        "encounters, and staff digital workflows."
    )
    two_imgs(
        "staff_team", "mission_wall",
        "FairBanks staff in branded uniforms at reception",
        "Mission, vision &amp; values wall — FairBanks Medical Centre identity",
        max_h=2.4 * inch,
    )
    two_imgs(
        "pharmacy_branded", "waiting_branded",
        "FairBanks Pharmacy exterior — Medical Centre branding",
        "Waiting area with FairBanks mission branding visible",
        max_h=2.4 * inch,
    )
    two_imgs(
        "outreach_bp", "outreach_outdoor",
        "BP screening and registration under outreach canopies",
        "Outdoor community clinic - structured field capture opportunity",
    )
    two_imgs(
        "training", "pharmacy_digital",
        "Staff training and product briefings for community programmes",
        "Pharmacy digital workflow - stock and dispensing signals for FCHIP",
    )
    two_imgs(
        "mothers_wait", "gericare",
        "Maternal and child health touchpoints at the facility",
        "Gericare - continuous support for older community members",
        max_h=2.6 * inch,
    )
    two_imgs(
        "facility_building", "staff_field",
        "FairBanks Medical Centre - operating validation site",
        "Staff-community conversation - trust that enables data quality",
    )
    h2("8.2 Why FairBanks wins")
    table(
        ["Typical startup", "Typical facility", "FairBanks"],
        [
            ["Technology without field access", "Patients without technology",
             "Both: live healthcare + deep-tech build"],
            ["Imported solutions", "Manual, reactive workflows",
             "Context-rooted AI from African realities"],
            ["Pilot without validation site", "Data without intelligence",
             "Design → pilot → validate → refine before scale"],
        ],
        [page_w / 3, page_w / 3, page_w / 3],
    )
    body(
        "Recommended AWIEF track: Startup Track - operating medical centre, revenue-generating "
        "services, active community programmes, deep-tech platform on proven field operations.",
        bold=True,
    )

    # 9
    impact_block = [
        Paragraph("9. Social Impact &amp; SDG Alignment", styles["H1Custom"]),
        HRFlowable(width="100%", thickness=1.2, color=teal, spaceAfter=8),
        img("impact", w=page_w * 0.55, max_h=3.2 * inch,
            caption="Earlier intervention for mothers, children, and communities"),
        Paragraph(
            "FCHIP shifts primary healthcare from sick-care to predictive, community-centred prevention.",
            styles["BodyCustom"],
        ),
    ]
    story.append(KeepTogether(impact_block))
    bullets([
        "Earlier outbreak detection; lower communicable morbidity",
        "Fewer maternal and neonatal complications",
        "Reduced NCD burden via hotspot-targeted screening",
        "Improved child nutrition and immunisation coverage",
        "Fewer medicine stock-outs; smarter resource allocation",
        "Evidence base for NGOs, donors, and governments",
    ])
    table(
        ["SDG", "FCHIP contribution"],
        [
            ["SDG 3 - Good Health and Well-Being", "Predictive community health; maternal/child/NCD focus"],
            ["SDG 5 - Gender Equality", "Women-led venture; maternal health intelligence"],
            ["SDG 9 - Industry, Innovation & Infrastructure", "Deep-tech health infrastructure for last-mile Africa"],
            ["SDG 10 - Reduced Inequalities", "Intelligence for underserved communities"],
            ["SDG 17 - Partnerships", "Multi-stakeholder health ecosystem integration"],
        ],
        [page_w * 0.42, page_w * 0.58],
    )

    # 10
    h1("10. MVP Scope &amp; Product Roadmap")
    two_imgs(
        "mvp_capture", "gis",
        "MVP: structured field registration and mobile capture",
        "MVP: GIS risk visualisation layer",
    )
    h2("10.1 Minimum Viable Product")
    body("Goal: Prove predictive community health intelligence in FairBanks' live catchment.")
    table(
        ["MVP component", "Description"],
        [
            ["CHW/VHT mobile app", "Offline capture: symptoms, vitals, maternal/child, household visits"],
            ["Cloud sync & data pipeline", "Secure ingestion, validation, and storage"],
            ["Analytics dashboard", "Facility trends, alerts, outreach priorities"],
            ["AI Module v1", "Rule-based + ML risk scoring for 2-3 priority use cases"],
            ["GIS layer v1", "Cases and risk zones by village/parish"],
            ["Pilot integration", "Connect FairBanks records and outreach workflows"],
        ],
        [page_w * 0.32, page_w * 0.68],
    )
    h2("10.2 Roadmap")
    table(
        ["Phase", "Timeline", "Milestones"],
        [
            ["Phase 1 - Pilot", "0-12 mo", "MVP with FairBanks CHWs; validate 3 use cases"],
            ["Phase 2 - District", "12-24 mo", "Partner clinics; district structures; NGO M&E; APIs"],
            ["Phase 3 - Regional", "24-36 mo", "Multi-district UG; East Africa; advanced ML"],
            ["Phase 4 - Expansion", "Year 3+", "CDS, local-language NLP, research modules"],
        ],
        [page_w * 0.28, page_w * 0.2, page_w * 0.52],
    )

    # 11-14
    h1("11. AWIEF Pitch n Grow 2026 Alignment")
    table(
        ["AWIEF expectation", "FairBanks / FCHIP response"],
        [
            ["Deep knowledge of local context", "Years of Kampala outreach; CHW/VHT relationships"],
            ["Deep understanding of communities", "Designed around how African communities access care"],
            ["Deep technology, not superficial digital", "AI + ML + GIS + mobile + cloud + NLP"],
            ["African problems, African-built solutions", "Built from FairBanks field operations"],
        ],
        [page_w * 0.4, page_w * 0.6],
    )
    bullets([
        "Sector: HealthTech | Deep-tech core: AI & Machine Learning",
        "Market: African communities (Uganda pilot → continental scale)",
        "Eligibility: Woman founder/co-founder presenting",
        "Commitment: AWIEF 2026 Conference finals, Cape Town, 10-11 November 2026",
    ])

    h1("12. Application-Ready Pitch Lines")
    h3("One-line pitch")
    body(
        "FairBanks Community Health Intelligence Platform (FCHIP) is a deep-tech platform that uses "
        "AI and predictive analytics to transform African primary healthcare from reactive "
        "treatment to proactive, community-level prevention."
    )
    h3("Unique value proposition")
    body(
        "Only venture combining a live medical centre and community outreach engine with a "
        "deep-tech predictive health platform - enabling real-world design, pilot, and validation "
        "before district and national scale.",
        bold=True,
    )

    h1("13. Risks &amp; Mitigation")
    table(
        ["Risk", "Mitigation"],
        [
            ["Data quality from field capture", "Structured forms, validation, CHW training"],
            ["Low CHW digital literacy", "Simple UX, offline mode, local language, supervised rollout"],
            ["Privacy and consent", "Ethical governance, anonymisation, Uganda Data Protection compliance"],
            ["Model accuracy early phase", "Rule-based + limited ML; validate against clinical outcomes"],
            ["Public-sector adoption", "Pilot evidence; district co-design; MoH alignment"],
        ],
        [page_w * 0.35, page_w * 0.65],
    )

    story.append(PageBreak())
    conclusion_block = [
        Paragraph("14. Conclusion", styles["H1Custom"]),
        HRFlowable(width="100%", thickness=1.2, color=teal, spaceAfter=8),
        img("conclusion", max_h=2.9 * inch,
            caption="From community signals to life-saving predictions - rooted in FairBanks outreach"),
        Paragraph(
            "FairBanks Community Health Intelligence Platform represents a strategic evolution from community "
            "medical practice to community health intelligence - matching AWIEF Pitch n Grow 2026's "
            "call for women-led, deep-technology solutions rooted in African reality.",
            styles["BodyCustom"],
        ),
        Paragraph(
            "With an operating medical centre, established Kampala outreach, trusted CHW/VHT "
            "relationships, digital health foundations, and a clear AI-powered roadmap, FairBanks is "
            "ready to build, pilot, and scale FCHIP.",
            styles["BodyCustom"],
        ),
        Paragraph("Next immediate steps:", styles["BodyBold"]),
    ]
    story.append(KeepTogether(conclusion_block))
    bullets([
        "Finalise MVP scope and AI priorities (surveillance, maternal risk, NCD hotspots)",
        "Complete AWIEF application and 10-12 slide pitch deck",
        "Record 3-minute technology demo video",
        "Launch FairBanks pilot with CHW/VHT cohort in primary catchment communities",
    ])
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Document prepared for AWIEF Pitch n Grow 2026 submission. Source: refined_project_idea.md",
        styles["Meta"],
    ))

    def add_page_number(canvas, doc_):
        canvas.saveState()
        canvas.setFillColor(navy)
        canvas.rect(0, A4[1] - 10, A4[0], 10, fill=1, stroke=0)
        canvas.setFillColor(teal)
        canvas.rect(0, A4[1] - 12, A4[0], 2, fill=1, stroke=0)
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.8 * inch, 0.45 * inch,
                          "FairBanks FCHIP  |  AWIEF Pitch n Grow 2026")
        canvas.drawRightString(A4[0] - 0.8 * inch, 0.45 * inch, f"{doc_.page}")
        canvas.restoreState()

    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    def write_pdf(path):
        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=0.8 * inch,
            rightMargin=0.8 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
            title="FairBanks FCHIP - AWIEF Pitch n Grow 2026",
            author="FairBanks Community Health Intelligence Platform",
        )
        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)

    try:
        write_pdf(OUT_PDF)
        print(f"PDF: {OUT_PDF}")
        return OUT_PDF
    except PermissionError:
        alt_dir = REPO / "tmp"
        alt_dir.mkdir(parents=True, exist_ok=True)
        alt = alt_dir / (OUT_PDF.stem + "_unlocked" + OUT_PDF.suffix)
        write_pdf(alt)
        print(f"PDF locked; saved as: {alt}")
        return alt


if __name__ == "__main__":
    # Pillow needed for PDF image sizing
    try:
        from PIL import Image as _
    except ImportError:
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pillow", "-q"])
    build_docx()
    build_pdf()
