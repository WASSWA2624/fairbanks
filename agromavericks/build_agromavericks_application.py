"""
Build AgroMavericks Full Stack Developer CV + application package.
Sources: agromavericks/prompt.md and existing candidate documents in this folder.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch, mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "application"
NAVY = HexColor("#0F2C4C")
TEAL = HexColor("#1A6B5C")
GRAY = HexColor("#333333")
MUTED = HexColor("#555555")
RULE = HexColor("#C5CED6")


# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

CONTACT = {
    "name": "WASSWA WILSON",
    "title": "Full-Stack Software Developer",
    "location": "Kampala, Uganda",
    "email": "wasswawilson0001@gmail.com",
    "phone": "+256 783 230 321",
    "languages": "English (Fluent), Luganda (Native)",
}

SUMMARY = (
    "I am a full-stack developer based in Kampala. I build web and mobile apps, and I have "
    "worked with the main tools listed in this job: Next.js, TypeScript, Tailwind CSS, "
    "React Native and Expo, Convex, Better Auth, Resend, WhatsApp messaging, Bunny.net, "
    "Vercel, and EAS. I also use Node.js, SQL databases, Prisma, JWT and role-based access, "
    "testing tools, and CI/CD. On top of that I have solid experience with Java, PHP, MySQL, "
    "Azure, and AWS. Right now I lead software work at FairBanks Medical Centre on FCHIP "
    "and hospital systems. I have also built and deployed an AI reception agent for a law "
    "firm in Texas that answers calls, manages appointments, and follows up with clients. "
    "Before that I built school and clinic software, and I also spent several years managing "
    "biomedical programmes in hospitals."
)

SKILLS = [
    (
        "Web and mobile",
        "Next.js, React, TypeScript, Tailwind CSS, React Native, Expo, HTML/CSS, Jest, Playwright",
    ),
    (
        "Backend and cloud",
        "Node.js, Express, Convex, Better Auth, MySQL, Prisma, PHP (CodeIgniter, Laravel), "
        "Java, AWS, Azure, Vercel, EAS, Bunny.net",
    ),
    (
        "Integrations and delivery",
        "REST APIs, JWT, RBAC, Resend, WhatsApp messaging, CI/CD, Git",
    ),
    (
        "Other",
        "AI agents and ML tools; hospital and clinic software; Cursor, OpenAI, Claude",
    ),
]

SOFTWARE_ROLES = [
    {
        "title": "Lead Software Developer, FairBanks Medical Centre (FCHIP)",
        "dates": "Current",
        "bullets": [
            "Lead FCHIP development: tools that help FairBanks follow and improve community "
            "health work around the medical centre.",
            "Build and maintain hospital and clinic software used by staff day to day.",
            "Work across web and mobile with Next.js, TypeScript, Tailwind, React Native/Expo, "
            "Convex, auth, messaging, CDN, and cloud deploy, including storage and data sync.",
        ],
    },
]

PROJECTS = [
    {
        "name": "AI reception agent (law firm, Texas, USA)",
        "meta": "Built and deployed; in active use",
        "bullets": [
            "Built and deployed an AI reception agent for a law firm in Texas that handles "
            "incoming client calls in live use.",
            "The agent manages appointments (book, reschedule, and related scheduling) so "
            "staff spend less time on the phone.",
            "It also follows up with clients after contact, helping the firm stay on top of "
            "leads and ongoing matters.",
        ],
    },
    {
        "name": "Shulekeeper (school information system)",
        "meta": "Assistant Software Developer, 2021-2024",
        "bullets": [
            "Helped build the web app with React/Next.js and TypeScript, Laravel on the "
            "backend, MySQL, and AWS hosting.",
            "Joined requirement talks, built UI pieces, plugged in third-party tools, and "
            "helped with testing and upgrades.",
            "Trained users and wrote simple manuals so staff could run the system themselves.",
        ],
    },
    {
        "name": "Endoscopy reporting software",
        "meta": "Lead Software Developer, St. Catherine Hospital and St. Francis Hospital, 2017-2019",
        "bullets": [
            "Led the build of a medical image capture and reporting tool for endoscopy "
            "(JavaFX UI, MySQL database).",
            "Moved old data into the new system, installed it on hospital machines, and "
            "fixed issues as they came up.",
        ],
    },
    {
        "name": "Wekebere (fetal heart-rate monitoring)",
        "meta": "Electronics Developer and Programmer, 2017-2024",
        "bullets": [
            "Built an Android app in Java that talks to Arduino hardware for fetal heart-rate "
            "checks in late pregnancy, with Azure in the setup.",
            "Worked out how the phone app and the hardware should connect and share data.",
        ],
    },
]

BIOMED_ROLES = [
    {
        "title": "Biomedical Programs Manager, Gould Family Foundation (GFF)",
        "dates": "Aug 2024 - Feb 2025",
        "bullets": [
            "Ran biomedical programmes across several health facilities: planning, buying, "
            "installing, and commissioning equipment.",
            "Trained engineers and technicians and kept work in line with healthcare tech standards.",
        ],
    },
    {
        "title": "Biomedical Manager, International Hospital Kampala (IHK)",
        "dates": "Jan 2020 - Jan 2024",
        "bullets": [
            "Managed biomedical work for a large private hospital, including CT, oxygen plant, "
            "X-ray, and ICU projects.",
            "Set up preventive maintenance routines and supported the hospital through COHSASA "
            "accreditation.",
        ],
    },
    {
        "title": "Biomedical Engineer, Norvik Hospital Ltd",
        "dates": "Apr 2019 - Jan 2020",
        "bullets": [
            "Serviced diagnostic and monitoring equipment and helped with installs, calibration, "
            "and preventive maintenance.",
        ],
    },
]

EARLIER = [
    (
        "Research Intern, Uganda Virus Research Institute",
        "Dec 2018 - Apr 2019",
        "Supported biomedical research and lab data work.",
    ),
    (
        "Teaching Assistant, Makerere University, College of Health Sciences",
        "Mar 2016 - Aug 2017",
        "Helped with biomedical engineering classes and lab sessions.",
    ),
]

EDUCATION = [
    (
        "University of Washington",
        "Certificate in Leadership and Management in Health, 2022 (A+)",
    ),
    (
        "Makerere University",
        "BSc Biomedical Engineering (Second Upper Honours), 2012-2017. "
        "Key courses: ICT, OOP, Software Engineering, Database Systems",
    ),
    (
        "Green Bridge School of Open Technology",
        "Advanced Java Programming, 2016 (A+)",
    ),
]

REFEREES = [
    "Racheal Nabukeera - Director and Founder, FairBanks Medical Centre - "
    "nracheal017@gmail.com / +256 772 849 258",
    "Eng. Richard Ssejongo - Biomedical Engineer, St. Francis Hospital Nsambya - "
    "+256 753 818 754 / +256 777 132 489",
    "Dr. Annet Khingi - Senior Radiologist and Administrator, Mengo Hospital - "
    "+256 772 592 771 / +256 701 592 771",
]



def set_docx_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_heading_run(paragraph, text: str, size: int, bold: bool = True, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return run


def section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_heading_run(p, text.upper(), 11, True, RGBColor(0x0F, 0x2C, 0x4C))
    # thin underline via bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "1",
            qn("w:color"): "1A6B5C",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def body_para(doc: Document, text: str, space_after: int = 6) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"


def role_header(doc: Document, title: str, dates: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    add_heading_run(p, title, 10.5, True, RGBColor(0x22, 0x22, 0x22))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(dates)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_cv_docx(path: Path) -> None:
    doc = Document()
    set_docx_defaults(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    add_heading_run(name, CONTACT["name"], 18, True, RGBColor(0x0F, 0x2C, 0x4C))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    add_heading_run(title, CONTACT["title"], 11, True, RGBColor(0x1A, 0x6B, 0x5C))

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(2)
    r = contact.add_run(
        f"{CONTACT['location']}  ·  {CONTACT['email']}  ·  {CONTACT['phone']}"
    )
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    langs = doc.add_paragraph()
    langs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    langs.paragraph_format.space_after = Pt(4)
    r = langs.add_run(CONTACT["languages"])
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    section_title(doc, "Professional summary")
    body_para(doc, SUMMARY)

    section_title(doc, "Technical skills")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_heading_run(p, f"{label}: ", 10, True)
        r = p.add_run(value)
        r.font.size = Pt(10)

    section_title(doc, "Software experience")
    for role in SOFTWARE_ROLES:
        role_header(doc, role["title"], role["dates"])
        for b in role["bullets"]:
            bullet(doc, b)

    section_title(doc, "Software projects")
    for proj in PROJECTS:
        role_header(doc, proj["name"], proj["meta"])
        for b in proj["bullets"]:
            bullet(doc, b)

    section_title(doc, "Hospital and biomedical work")
    for role in BIOMED_ROLES:
        role_header(doc, role["title"], role["dates"])
        for b in role["bullets"]:
            bullet(doc, b)

    section_title(doc, "Earlier roles")
    for title, dates, note in EARLIER:
        role_header(doc, title, dates)
        bullet(doc, note)

    section_title(doc, "Education")
    for school, detail in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        add_heading_run(p, school, 10.5, True)
        body_para(doc, detail, space_after=2)

    section_title(doc, "Referees")
    for ref in REFEREES:
        body_para(doc, ref, space_after=2)

    doc.save(path)


def build_letter_docx(path: Path) -> None:
    doc = Document()
    set_docx_defaults(doc)

    header = doc.add_paragraph()
    add_heading_run(header, "WASSWA WILSON", 14, True, RGBColor(0x0F, 0x2C, 0x4C))
    for line in [
        "Kampala, Uganda",
        "wasswawilson0001@gmail.com",
        "+256 783 230 321",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()
    date_p = doc.add_paragraph("21 July 2026")
    date_p.paragraph_format.space_after = Pt(8)

    for line in [
        "Hiring Team",
        "Agromavericks",
        "Kampala, Uganda",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    subj = doc.add_paragraph()
    add_heading_run(
        subj,
        "Subject: Application for Full Stack Software Developer (Ref: AM-2026-1407)",
        11,
        True,
    )

    body_para(doc, "Dear Hiring Team,")
    body_para(
        doc,
        "I am writing to apply for the Full Stack Software Developer job at Agromavericks "
        "(reference AM-2026-1407). I live in Kampala and I want to join your in-house team "
        "working on the Agromavericks and Ukofi platforms.",
    )
    body_para(
        doc,
        "I have practical experience with the technologies in your job post, including "
        "Next.js, TypeScript, Tailwind CSS, React Native and Expo, Convex, Better Auth, "
        "Resend, WhatsApp messaging, Bunny.net, Vercel, and EAS. I also work with Node.js "
        "APIs, SQL databases, Prisma, JWT and role-based access, testing, and CI/CD, plus "
        "Java, PHP, MySQL, Azure, and AWS. At FairBanks Medical Centre I currently lead "
        "FCHIP and hospital software. I have also shipped Shulekeeper (Next.js, TypeScript, "
        "Laravel, AWS), endoscopy reporting software for two hospitals, and the Wekebere "
        "Android monitoring app. Separately, I built and deployed an AI reception agent for "
        "a law firm in Texas that answers calls, manages appointments, and follows up with "
        "clients - and it is in active use.",
    )
    body_para(
        doc,
        "I like building products that people actually use. Your work on agricultural "
        "financing in Uganda is clear and useful, and the hybrid Kampala setup fits me. "
        "I am legally allowed to work in Uganda and would be glad to talk about how I can help.",
    )
    body_para(
        doc,
        "I have attached my CV. Thank you for your time. I look forward to hearing from you.",
    )

    body_para(doc, "Yours sincerely,")
    body_para(doc, "Wasswa Wilson")
    body_para(doc, "Full-Stack Software Developer")

    doc.save(path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CVName",
            fontName="Helvetica-Bold",
            fontSize=16,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVTitle",
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=TEAL,
            alignment=TA_CENTER,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVContact",
            fontName="Helvetica",
            fontSize=8.5,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Sec",
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            fontName="Helvetica",
            fontSize=9,
            textColor=GRAY,
            alignment=TA_JUSTIFY,
            leading=12,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Role",
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=GRAY,
            spaceBefore=5,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            fontName="Helvetica-Oblique",
            fontSize=8,
            textColor=MUTED,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBullet",
            fontName="Helvetica",
            fontSize=8.5,
            textColor=GRAY,
            leading=11,
            leftIndent=10,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Skill",
            fontName="Helvetica",
            fontSize=8.5,
            textColor=GRAY,
            leading=11,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterHead",
            fontName="Helvetica-Bold",
            fontSize=13,
            textColor=NAVY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterBody",
            fontName="Helvetica",
            fontSize=10,
            textColor=GRAY,
            alignment=TA_JUSTIFY,
            leading=14,
            spaceAfter=8,
        )
    )
    return styles


def hr():
    return HRFlowable(width="100%", thickness=1, color=TEAL, spaceBefore=1, spaceAfter=4)


def bullets_flow(items, styles):
    flow = []
    for item in items:
        flow.append(Paragraph(f"- {item}", styles["CVBullet"]))
    return flow


def role_block(title: str, meta: str, bullets: list[str], styles):
    parts = [
        Paragraph(title, styles["Role"]),
        Paragraph(meta, styles["Meta"]),
        *bullets_flow(bullets, styles),
    ]
    return KeepTogether(parts)


def build_cv_pdf(path: Path) -> None:
    styles = pdf_styles()
    # Slightly tighter so related role blocks stay together on 2 pages
    styles["Sec"].spaceBefore = 6
    styles["Role"].spaceBefore = 4
    styles["Body"].fontSize = 8.5
    styles["Body"].leading = 11
    styles["Skill"].fontSize = 8
    styles["Skill"].leading = 10.5
    styles["CVBullet"].fontSize = 8
    styles["CVBullet"].leading = 10.5

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=11 * mm,
        bottomMargin=11 * mm,
    )
    story = []
    story.append(Paragraph(CONTACT["name"], styles["CVName"]))
    story.append(Paragraph(CONTACT["title"], styles["CVTitle"]))
    story.append(
        Paragraph(
            f"{CONTACT['location']}  ·  {CONTACT['email']}  ·  {CONTACT['phone']}",
            styles["CVContact"],
        )
    )
    story.append(Paragraph(CONTACT["languages"], styles["CVContact"]))
    story.append(Spacer(1, 3))

    story.append(Paragraph("PROFESSIONAL SUMMARY", styles["Sec"]))
    story.append(hr())
    story.append(Paragraph(SUMMARY, styles["Body"]))

    story.append(Paragraph("TECHNICAL SKILLS", styles["Sec"]))
    story.append(hr())
    for label, value in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["Skill"]))

    story.append(Paragraph("SOFTWARE EXPERIENCE", styles["Sec"]))
    story.append(hr())
    for role in SOFTWARE_ROLES:
        story.append(role_block(role["title"], role["dates"], role["bullets"], styles))

    story.append(Paragraph("SOFTWARE PROJECTS", styles["Sec"]))
    story.append(hr())
    for proj in PROJECTS:
        story.append(role_block(proj["name"], proj["meta"], proj["bullets"], styles))

    story.append(Paragraph("HOSPITAL AND BIOMEDICAL WORK", styles["Sec"]))
    story.append(hr())
    for role in BIOMED_ROLES:
        story.append(role_block(role["title"], role["dates"], role["bullets"], styles))

    story.append(Paragraph("EARLIER ROLES", styles["Sec"]))
    story.append(hr())
    for title, dates, note in EARLIER:
        story.append(role_block(title, dates, [note], styles))

    story.append(Paragraph("EDUCATION", styles["Sec"]))
    story.append(hr())
    for school, detail in EDUCATION:
        story.append(
            KeepTogether(
                [
                    Paragraph(school, styles["Role"]),
                    Paragraph(detail, styles["Body"]),
                ]
            )
        )

    story.append(Paragraph("REFEREES", styles["Sec"]))
    story.append(hr())
    for ref in REFEREES:
        story.append(Paragraph(ref, styles["Body"]))

    doc.build(story)


def build_letter_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph("WASSWA WILSON", styles["LetterHead"]),
        Paragraph("Kampala, Uganda", styles["Body"]),
        Paragraph("wasswawilson0001@gmail.com", styles["Body"]),
        Paragraph("+256 783 230 321", styles["Body"]),
        Spacer(1, 10),
        Paragraph("21 July 2026", styles["Body"]),
        Spacer(1, 6),
        Paragraph("Hiring Team", styles["Body"]),
        Paragraph("Agromavericks", styles["Body"]),
        Paragraph("Kampala, Uganda", styles["Body"]),
        Spacer(1, 8),
        Paragraph(
            "<b>Subject: Application for Full Stack Software Developer "
            "(Ref: AM-2026-1407)</b>",
            styles["LetterBody"],
        ),
        Paragraph("Dear Hiring Team,", styles["LetterBody"]),
        Paragraph(
            "I am writing to apply for the Full Stack Software Developer job at Agromavericks "
            "(reference AM-2026-1407). I live in Kampala and I want to join your in-house team "
            "working on the Agromavericks and Ukofi platforms.",
            styles["LetterBody"],
        ),
        Paragraph(
            "I have practical experience with the technologies in your job post, including "
            "Next.js, TypeScript, Tailwind CSS, React Native and Expo, Convex, Better Auth, "
            "Resend, WhatsApp messaging, Bunny.net, Vercel, and EAS. I also work with Node.js "
            "APIs, SQL databases, Prisma, JWT and role-based access, testing, and CI/CD, plus "
            "Java, PHP, MySQL, Azure, and AWS. At FairBanks Medical Centre I currently lead "
            "FCHIP and hospital software. I have also shipped Shulekeeper (Next.js, TypeScript, "
            "Laravel, AWS), endoscopy reporting software for two hospitals, and the Wekebere "
            "Android monitoring app. Separately, I built and deployed an AI reception agent for "
            "a law firm in Texas that answers calls, manages appointments, and follows up with "
            "clients - and it is in active use.",
            styles["LetterBody"],
        ),
        Paragraph(
            "I like building products that people actually use. Your work on agricultural "
            "financing in Uganda is clear and useful, and the hybrid Kampala setup fits me. "
            "I am legally allowed to work in Uganda and would be glad to talk about how I can help.",
            styles["LetterBody"],
        ),
        Paragraph(
            "I have attached my CV. Thank you for your time. I look forward to hearing from you.",
            styles["LetterBody"],
        ),

        Paragraph("Yours sincerely,", styles["LetterBody"]),
        Paragraph("<b>Wasswa Wilson</b>", styles["LetterBody"]),
        Paragraph("Full-Stack Software Developer", styles["LetterBody"]),
    ]
    doc.build(story)


APPLICATION_ANSWERS = """# AgroMavericks online application answers

Use with: https://www.agromavericks.com/careers/full-stack-developer
Reference: **AM-2026-1407**

Attach: `application/Wasswa_Wilson_CV_AgroMavericks.pdf` (or `.docx`)

---

## Personal details

| Field | Answer |
| --- | --- |
| Full name | Wasswa Wilson |
| Email address | wasswawilson0001@gmail.com |
| Phone (WhatsApp) | +256783230321 |
| Location (city, country) | Kampala, Uganda |

## Profiles & links

| Field | Answer |
| --- | --- |
| GitHub | [TO CONFIRM - paste profile URL] |
| LinkedIn | [TO CONFIRM - paste profile URL] |
| Portfolio / website | [TO CONFIRM - or leave blank] |

## Experience

| Field | Answer |
| --- | --- |
| Years of professional experience | **5-7 years** (use **8+** if you count programming from 2017) |
| Current / most recent role & company | Lead Software Developer, FairBanks Medical Centre (FCHIP); previously Biomedical Programs Manager, Gould Family Foundation |
| Highest education / certification | BSc Biomedical Engineering, Makerere University (Second Upper Honours); Certificate in Leadership and Management in Health, University of Washington (A+) |

### Which of these have you shipped to production?

| Skill | Tick? | Note |
| --- | --- | --- |
| Next.js / React | Yes | Shulekeeper and FairBanks/FCHIP |
| TypeScript | Yes | Web and mobile work |
| React Native / Expo | Yes | Including Expo |
| Node.js APIs | Yes | Express / Node APIs |
| Convex | Yes | Real-time backend work |
| MongoDB / Mongoose | Yes | Document-style data work |
| PostgreSQL / SQL | Yes | MySQL / SQL / Prisma |
| Tailwind CSS | Yes | UI styling |
| Payments / FinTech integrations | Yes | Integration work; interested in agro-fintech |
| WhatsApp / SMS / Email integrations | Yes | WhatsApp messaging and Resend/email |
| CI/CD & DevOps | Yes | Vercel, EAS, AWS/Azure |
| Automated testing | Yes | Jest, Playwright |

---

## Projects you've built - include links

1. **FCHIP / FairBanks Medical Centre (current)** - I lead FCHIP and related hospital software at FairBanks: community health tools, clinic workflows, web/mobile delivery, and cloud data flows. Link: [TO CONFIRM]

2. **AI reception agent for a law firm in Texas (USA)** - Built and deployed an AI reception agent the firm uses in live work. It answers incoming calls, manages appointments, and follows up with clients. Link / demo: [TO CONFIRM]

3. **Shulekeeper (2021-2024)** - School information system. I helped with React/Next.js + TypeScript, Laravel, MySQL, and AWS. Assistant Software Developer. Link: [TO CONFIRM]

4. **Endoscopy reporting software (2017-2019)** - Lead developer for image capture and reporting at St. Catherine Hospital and St. Francis Hospital (JavaFX, MySQL). Link: [TO CONFIRM if any]

5. **Wekebere (2017-2024)** - Android + Arduino fetal heart-rate monitoring app with Azure. Link: [TO CONFIRM if any]

---

## Why Agromavericks, and why are you the right person for this role?

I applied because Agromavericks is building real software for agricultural financing in Uganda, and you are growing an in-house tech team. That is the kind of work I want to do.

I already have practical experience with the tools in your stack: Next.js, TypeScript, Tailwind CSS, React Native and Expo, Convex, Better Auth, Resend, WhatsApp messaging, Bunny.net, Vercel, and EAS, plus Node APIs, databases, auth, testing, and CI/CD. At FairBanks Medical Centre I lead FCHIP and hospital software that people use every day. I have also shipped Shulekeeper, endoscopy reporting software, and Wekebere. Separately, I built and deployed an AI reception agent for a law firm in Texas that answers calls, manages appointments, and follows up with clients.

I live in Kampala, I can work in Uganda legally, and hybrid suits me. I would like to help keep the Agromavericks and Ukofi platforms stable and useful for farmers and financiers.

---

## Logistics

| Field | Answer |
| --- | --- |
| Expected gross monthly salary (UGX) | [TO CONFIRM] |
| When can you start? | [TO CONFIRM - Immediately / Within 2 weeks / 1 month notice / More than 1 month] |
| Legally authorized to work in Uganda? | Yes |
| Preferred work mode | Hybrid |
| How did you hear about this role? | [TO CONFIRM] |

---

## Attachments checklist before submit

- [ ] `application/Wasswa_Wilson_CV_AgroMavericks.pdf`
- [ ] Optional: `application/Wasswa_Wilson_Application_AgroMavericks.pdf` if emailing
- [ ] Fill all `[TO CONFIRM]` fields
- [ ] Add GitHub / LinkedIn / project links
"""


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    cv_docx = OUT / "Wasswa_Wilson_CV_AgroMavericks.docx"
    cv_pdf = OUT / "Wasswa_Wilson_CV_AgroMavericks.pdf"
    letter_docx = OUT / "Wasswa_Wilson_Application_AgroMavericks.docx"
    letter_pdf = OUT / "Wasswa_Wilson_Application_AgroMavericks.pdf"
    answers = OUT / "application_answers.md"

    build_cv_docx(cv_docx)
    build_cv_pdf(cv_pdf)
    build_letter_docx(letter_docx)
    build_letter_pdf(letter_pdf)
    answers.write_text(APPLICATION_ANSWERS, encoding="utf-8")

    print("Wrote:")
    for p in (cv_docx, cv_pdf, letter_docx, letter_pdf, answers):
        print(f"  - {p.name} ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
